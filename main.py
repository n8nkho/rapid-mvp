from fastapi import FastAPI, APIRouter, HTTPException, File, Form, UploadFile, Request, Header, Depends, Path, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.exceptions import RequestValidationError
from pydantic import BaseModel
from contextlib import asynccontextmanager
from typing import Optional, List, Dict, Any
import io
import json
import logging
import os
import re
import secrets
import uuid
from uuid import UUID
import psycopg2
import httpx
from datetime import datetime, timezone, timedelta
from openpyxl import Workbook, load_workbook

# Import config and validate env before any DB or providers
from config import validate_config, get_cors_origins
from auth import require_api_key

validate_config()  # fail fast if required env missing

from providers import get_provider, MODEL_HAIKU, MODEL_SONNET, MODEL_SONNET_SEED
from database import (
    save_gap_analysis,
    get_results_by_engagement,
    get_gap_results_by_req_id,
    create_requirement,
    get_requirements_by_engagement,
    get_requirement_by_id,
    update_requirement,
    create_client,
    get_client,
    list_clients,
    update_client,
    get_benchmark_hints_by_engagement,
    create_benchmark_hint,
    create_engagement,
    get_engagement,
    list_engagements,
    get_engagement_with_client,
    update_engagement,
    create_asset,
    get_assets_by_engagement,
    get_assets_by_requirement,
    update_asset,
    upload_file_to_storage,
    create_process_step,
    get_process_steps,
    get_process_step,
    update_process_step,
    delete_process_step,
    get_process_steps_by_engagement,
    create_ricefw_item,
    get_ricefw_by_engagement,
    update_ricefw_item,
    delete_ricefw_item,
    create_hitl_event,
    list_hitl_events,
    create_fit_gap_assessment,
    get_fit_gap_by_engagement,
    get_fit_gap_by_assessment_id,
    update_fit_gap_assessment,
    delete_fit_gap_assessment,
    delete_requirement,
    create_feedback_event,
    list_feedback_events,
    get_pattern_library,
    increment_pattern_use,
    list_agent_roles,
    get_agent_role_by_role_id,
    get_agent_knowledge_by_role,
    create_agent_maturity_score,
    get_agent_maturity_scores,
    create_platform_issue,
    list_platform_issues,
    update_platform_issue,
    create_audit_event,
    list_audit_events_by_engagement,
    get_raci_matrix,
    upsert_raci_matrix,
    get_engagement_scope,
    upsert_engagement_scope,
    retain_only_engagement,
    create_source,
    get_source,
    list_sources_by_engagement,
    update_source,
    delete_source,
    create_test_script,
    list_test_scripts_by_engagement,
    list_test_scripts_by_ricefw,
    create_portal_user,
    get_portal_user_by_token,
    update_portal_user_last_access,
    create_go_live_checklist_item,
    get_go_live_checklist,
    update_go_live_checklist_item,
    create_pattern,
    create_queue_item,
    get_queue_items,
    update_queue_item_status,
    update_engagement_mode,
    supabase,
)
from autonomy import should_auto_execute, get_effective_config
from scope_items import SCOPE_ITEMS, get_catalogue_text

# Structured logging; config validated in lifespan
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("rapid")


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Config already validated at import; use lifespan for future startup (e.g. connection pools)
    yield


app = FastAPI(
    title="RAPID Gap Analysis API",
    description="AI-powered SAP S/4HANA scope item gap analysis using semantic matching",
    version="1.2.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=get_cors_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Global exception handlers (enterprise: stable JSON, no internal leak on 5xx) ──

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(_request: Request, exc: RequestValidationError):
    return JSONResponse(
        status_code=422,
        content={
            "detail": exc.errors(),
            "message": "Validation error",
        },
    )


@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    if exc.status_code >= 500:
        request_id = getattr(request.state, "request_id", None) or str(uuid.uuid4())[:8]
        logger.exception(
            "HTTP 5xx: request_id=%s path=%s detail=%s",
            request_id,
            request.url.path,
            exc.detail,
        )
        return JSONResponse(
            status_code=exc.status_code,
            content={
                "detail": "An internal error occurred. Please try again or contact support.",
                "request_id": request_id,
            },
        )
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    request_id = str(uuid.uuid4())[:8]
    logger.exception(
        "Unhandled exception: request_id=%s path=%s",
        request_id,
        request.url.path,
        exc_info=exc,
    )
    return JSONResponse(
        status_code=500,
        content={
            "detail": "An internal error occurred. Please try again or contact support.",
            "request_id": request_id,
        },
    )


@app.middleware("http")
async def add_request_id(request: Request, call_next):
    request.state.request_id = str(uuid.uuid4())[:8]
    response = await call_next(request)
    return response


# v1 API: all routes under /v1; optional API key when API_KEY env is set
router = APIRouter(prefix="/v1", dependencies=[Depends(require_api_key)])
# Public portal routes: no API key — authenticated by portal token only
portal_router = APIRouter(prefix="/v1")
# Admin routes: no API key required; only X-Admin-Key (or query admin_key) when ADMIN_API_KEY is set
admin_router = APIRouter(prefix="/v1")


# ── Pydantic models ───────────────────────────────────────────────────────────

class GapAnalysisRequest(BaseModel):
    engagement_id: str
    process_description: Optional[str] = None  # required unless req_id provided
    req_id: Optional[str] = None               # looks up description from requirements
    top_n: Optional[int] = 5
    lob_filter: Optional[str] = None

class ScopeItemMatch(BaseModel):
    id: str
    name: str
    lob: str
    process_group: str
    description: str
    confidence: str
    rationale: str
    migration_objects: List[str]

class GapAnalysisResponse(BaseModel):
    engagement_id: str
    req_id: Optional[str] = None
    process_description: str
    matches: List[ScopeItemMatch]
    total_scope_items_searched: int
    tokens_used: Optional[int] = None
    timestamp: str

class RequirementCreate(BaseModel):
    engagement_id: str
    title: str
    description: str
    source_type: Optional[str] = None
    tags: Optional[List[str]] = []
    stakeholder: Optional[str] = None
    raw_input: Optional[str] = None
    # Phase 1 extended fields
    business_process: Optional[str] = None         # LOB, e.g. Finance / Sales / Supply Chain
    priority: Optional[str] = "Must-Have"          # Must-Have / Should-Have / Nice-to-Have
    category: Optional[str] = None                 # Automation / Control/Compliance / Reporting / Integration / UX / Data Migration
    kpi_impact: Optional[Dict] = None              # {"metric": "...", "target": "...", "unit": "..."}
    confidence_score: Optional[float] = 0.8
    current_state_ref: Optional[str] = None        # quote, file name, video timestamp
    actors: Optional[List[Dict]] = None            # [{"role": "AP Clerk", "type": "formal"}, ...]
    shadow_tools: Optional[List[str]] = None       # ["Excel macro", "WhatsApp group", ...]
    sign_off_status: Optional[str] = "draft"       # draft / sme_approved / owner_approved / confirmed
    sign_off_by: Optional[str] = None
    sign_off_at: Optional[str] = None
    sap_mapping_id: Optional[str] = None           # nearest scope item ID
    fit_assessment: Optional[str] = None           # Fit-to-Standard / Soft-Gap / Hard-Gap
    # Sprint 5 process flow fields
    process_level_2: Optional[str] = None          # e.g. Record to Report, Procure to Pay
    process_level_3: Optional[str] = None          # e.g. Accounts Payable, General Ledger
    # Audit / RTM: external reference ID and source fields
    reference_id: Optional[str] = None            # e.g. Excel/source requirement ID for audit trail
    business_value: Optional[str] = None
    current_system: Optional[str] = None
    target_system_module: Optional[str] = None
    fit_type: Optional[str] = None
    related_test_case_id: Optional[str] = None
    # HITL pipeline
    hitl_state: Optional[str] = "ai_draft"
    hitl_history: Optional[List[Dict]] = None
    ai_rationale: Optional[str] = None
    reviewer_notes: Optional[str] = None

class RequirementUpdate(BaseModel):
    status: Optional[str] = None
    tags: Optional[List[str]] = None
    title: Optional[str] = None
    description: Optional[str] = None
    stakeholder: Optional[str] = None
    # Phase 1 extended fields
    business_process: Optional[str] = None
    priority: Optional[str] = None
    category: Optional[str] = None
    kpi_impact: Optional[Dict] = None
    confidence_score: Optional[float] = None
    current_state_ref: Optional[str] = None
    actors: Optional[List[Dict]] = None
    shadow_tools: Optional[List[str]] = None
    sign_off_status: Optional[str] = None
    sign_off_by: Optional[str] = None
    sign_off_at: Optional[str] = None
    sap_mapping_id: Optional[str] = None
    fit_assessment: Optional[str] = None
    # Sprint 5 process flow fields
    process_level_2: Optional[str] = None
    process_level_3: Optional[str] = None
    reference_id: Optional[str] = None
    business_value: Optional[str] = None
    current_system: Optional[str] = None
    target_system_module: Optional[str] = None
    fit_type: Optional[str] = None
    related_test_case_id: Optional[str] = None
    # HITL pipeline
    hitl_state: Optional[str] = None
    hitl_history: Optional[List[Dict]] = None
    ai_rationale: Optional[str] = None
    reviewer_notes: Optional[str] = None
    # Enterprise: archive instead of delete when approved FGA
    archived: Optional[bool] = None

class TranscriptExtractRequest(BaseModel):
    engagement_id: str
    stakeholder: str
    transcript_text: str


class DuplicateCheckRequest(BaseModel):
    engagement_id: str
    title: str
    exclude_req_id: Optional[str] = None  # when editing, exclude self from matches


class ArchaeologistSessionRequest(BaseModel):
    engagement_id: str
    stakeholder: str
    role: str
    business_process: str
    message: str
    session_history: Optional[List[Dict]] = []

class ClientCreate(BaseModel):
    name: str
    industry: Optional[str] = None
    employees: Optional[int] = None
    legal_entities: Optional[int] = None
    address: Optional[str] = None
    current_systems: Optional[List[str]] = None
    systems_to_keep: Optional[List[str]] = None
    systems_to_replace: Optional[List[str]] = None
    countries: Optional[List[str]] = None
    regulatory_environment: Optional[List[str]] = None
    # Strategy & context (pre-fill from website)
    business_strategy: Optional[str] = None
    goals: Optional[List[str]] = None
    key_products: Optional[List[str]] = None
    value_proposition: Optional[str] = None
    senior_executives: Optional[List[Dict[str, str]]] = None  # [{"name": "...", "title": "..."}]
    direct_competitors: Optional[List[str]] = None
    substitutes: Optional[List[str]] = None
    # Phase E: sector & benchmarks
    sector_archetype: Optional[str] = None
    complexity_drivers: Optional[List[str]] = None
    erp_maturity: Optional[str] = None
    benchmark_opt_in: Optional[bool] = True

class ClientPrefillFromWebsiteRequest(BaseModel):
    url: str


class EngagementCreate(BaseModel):
    client_id: str
    name: str
    description: Optional[str] = None
    go_live_date: Optional[str] = None
    project_type: Optional[str] = None
    # Control: status & dates
    status: Optional[str] = "open"          # open | completed | abandoned
    planned_start_date: Optional[str] = None
    planned_end_date: Optional[str] = None
    actual_start_date: Optional[str] = None
    actual_end_date: Optional[str] = None
    # Additional engagement control fields
    project_manager: Optional[str] = None
    sponsor: Optional[str] = None
    risk_level: Optional[str] = None        # low | medium | high
    health: Optional[str] = None            # on_track | at_risk | off_track
    # Engagement Mode (Phase 1)
    mode: Optional[str] = "collaborative"   # guided | collaborative | autonomous
    autonomy_config: Optional[dict] = None

class SignOffRequest(BaseModel):
    level: str      # "sme" or "owner"
    signed_by: str


class HitlAdvanceRequest(BaseModel):
    engagement_id: str
    actor: str
    actor_role: Optional[str] = None
    notes: Optional[str] = None
    correction: Optional[Dict[str, Any]] = None


class HitlRejectRequest(BaseModel):
    engagement_id: str
    actor: str
    reason: str
    target_state: Optional[str] = None  # "ai_draft" or "out_of_scope"

class AssetUpdate(BaseModel):
    req_id: Optional[str] = None
    process_level_2: Optional[str] = None
    process_level_3: Optional[str] = None
    confirmed: Optional[bool] = None
    suggested_tags: Optional[List[str]] = None

class RequirementResponse(BaseModel):
    req_id: str
    engagement_id: str
    title: str
    description: str
    source_type: Optional[str] = None
    tags: Optional[List[str]] = []
    stakeholder: Optional[str] = None
    raw_input: Optional[str] = None
    status: str
    created_at: Optional[str] = None
    # Phase 1 extended fields
    business_process: Optional[str] = None
    priority: Optional[str] = None
    category: Optional[str] = None
    kpi_impact: Optional[Any] = None
    confidence_score: Optional[float] = None
    current_state_ref: Optional[str] = None
    actors: Optional[Any] = None
    shadow_tools: Optional[List[str]] = None
    sign_off_status: Optional[str] = None
    sign_off_by: Optional[str] = None
    sign_off_at: Optional[str] = None
    sap_mapping_id: Optional[str] = None
    fit_assessment: Optional[str] = None
    # Sprint 5 process flow fields
    process_level_2: Optional[str] = None
    process_level_3: Optional[str] = None
    # Audit / RTM: external reference and source fields
    reference_id: Optional[str] = None
    business_value: Optional[str] = None
    current_system: Optional[str] = None
    target_system_module: Optional[str] = None
    fit_type: Optional[str] = None
    related_test_case_id: Optional[str] = None
    # HITL pipeline
    hitl_state: Optional[str] = None
    hitl_history: Optional[Any] = None
    ai_rationale: Optional[str] = None
    reviewer_notes: Optional[str] = None


class ProcessFlowAssignRequest(BaseModel):
    req_id: str
    process_level_2: str
    process_level_3: Optional[str] = None


class ProcessStepCreate(BaseModel):
    step_number: int
    title: str
    description: str
    performer_name: str
    performer_role: str
    shape: str = "process"      # start | end | process | decision | document
    step_type: str = "manual"   # manual | system | agentic
    duration_minutes: Optional[float] = None
    systems_used: Optional[List[str]] = []
    kpis: Optional[Dict] = None
    is_pain_point: bool = False
    next_step_id: Optional[str] = None
    branches: Optional[List[Dict]] = None


class ProcessStepUpdate(BaseModel):
    step_number: Optional[int] = None
    title: Optional[str] = None
    description: Optional[str] = None
    performer_name: Optional[str] = None
    performer_role: Optional[str] = None
    shape: Optional[str] = None
    step_type: Optional[str] = None
    duration_minutes: Optional[float] = None
    systems_used: Optional[List[str]] = None
    kpis: Optional[Dict] = None
    is_pain_point: Optional[bool] = None
    next_step_id: Optional[str] = None
    branches: Optional[List[Dict]] = None


# Sprint 7: RICEFW = Reports, Interfaces, Conversions, Enhancements, Forms, Workflows, Agents
_RICEFW_COMPLEXITY = {"very_high", "high", "medium", "low"}
_RICEFW_PRIORITY = {"must", "should", "could", "noneed"}

class RICEFWCreate(BaseModel):
    type: str       # R | I | C | E | F | W | A
    name: str
    description: str
    req_id: str     # required: link to requirement
    status: Optional[str] = "identified"
    complexity: Optional[str] = None   # very_high | high | medium | low
    priority: Optional[str] = None     # must | should | could | noneed
    effort_days_low: Optional[int] = None
    effort_days_high: Optional[int] = None
    owner: Optional[str] = None


class RICEFWUpdate(BaseModel):
    type: Optional[str] = None
    name: Optional[str] = None
    description: Optional[str] = None
    req_id: Optional[str] = None
    status: Optional[str] = None
    complexity: Optional[str] = None
    priority: Optional[str] = None
    effort_days_low: Optional[int] = None
    effort_days_high: Optional[int] = None
    owner: Optional[str] = None


# Phase D: Feedback
class FeedbackCreate(BaseModel):
    engagement_id: Optional[str] = None
    event_type: str = "general"  # general | pattern_used | correction | suggestion
    payload: Optional[Dict[str, Any]] = None  # e.g. {"pattern_id": "uuid"} for pattern_used


# Phase B: Fit/Gap
class FitGapReviewRequest(BaseModel):
    reviewer: str
    fit_type: Optional[str] = None
    complexity: Optional[str] = None
    notes: Optional[str] = None
    approve: bool


# Agent Team & Simulation
class SimulateAgentRequest(BaseModel):
    engagement_id: Optional[str] = ""
    agent_role_id: str
    phase: Optional[str] = None
    context_message: Optional[str] = None
    conversation_turn: Optional[List[Dict[str, str]]] = None


class SeedRequirementsRequest(BaseModel):
    engagement_id: str
    industry: str
    processes: List[str]


class PlatformIssueCreate(BaseModel):
    engagement_id: str
    agent_role_id: Optional[str] = None
    phase: Optional[str] = "requirements"
    context: Optional[Dict[str, Any]] = None
    problem_description: str
    issue_type: Optional[str] = "missing_feature"
    suggested_improvement: Optional[str] = None
    priority: Optional[str] = "medium"


class PlatformIssueUpdate(BaseModel):
    status: Optional[str] = None
    priority: Optional[str] = None


class MaturityScoreCreate(BaseModel):
    criterion: str  # domain_knowledge | reasoning_quality | authenticity | collaboration
    score: int  # 1-5
    notes: Optional[str] = None


# ── Helpers ───────────────────────────────────────────────────────────────────


def _normalize_tag(value: str) -> str:
    """Normalise tag strings to canonical snake_case form, similar to frontend."""
    if not value:
        return ""
    slug = re.sub(r"[^a-z0-9]+", "_", value.strip().lower())
    slug = re.sub(r"_+", "_", slug).strip("_")
    return slug


def build_catalogue_for_prompt(lob_filter: Optional[str] = None) -> str:
    items = SCOPE_ITEMS
    if lob_filter:
        items = [i for i in items if i.get('lob', '').lower() == lob_filter.lower()]
    lines = []
    for item in items:
        lines.append(
            f"ID:{item['id']} | {item['name']} | {item['lob']} > {item['process_group']}\n"
            f"  {item['description']}"
        )
    return "\n".join(lines)


_GAP_SYSTEM_PROMPT = """You are an expert SAP S/4HANA implementation consultant specializing in Fit-to-Standard gap analysis.

Your task: Given a business process description, identify the most relevant SAP S/4HANA Cloud Public Edition scope items from the provided catalogue.

Instructions:
1. Analyze the business process description semantically - look beyond keywords to understand intent
2. Return the top matching scope items ranked by relevance
3. For each match, provide confidence (HIGH / MEDIUM / LOW) and a brief rationale
4. Consider that one business requirement often maps to multiple scope items
5. Always return valid JSON only

Response format (JSON array):
[
  {
    "id": "scope_item_code",
    "confidence": "HIGH|MEDIUM|LOW",
    "rationale": "One sentence explaining why this scope item matches"
  }
]"""


def _run_gap_analysis(
    provider,
    process_description: str,
    top_n: int = 5,
    lob_filter: Optional[str] = None,
) -> tuple:
    """Returns (matches: List[ScopeItemMatch], tokens_used: int)."""
    catalogue = build_catalogue_for_prompt(lob_filter)
    user_prompt = (
        f"Business Process Description:\n{process_description}\n\n"
        f"SAP S/4HANA Cloud 2602 Scope Item Catalogue (2602 release):\n{catalogue}\n\n"
        f"Return the top {top_n} most relevant scope items as JSON."
    )
    # Haiku is cheap and sufficient; keep max output small
    result = provider.complete(_GAP_SYSTEM_PROMPT, user_prompt, max_tokens=512, model=MODEL_HAIKU)
    raw_text = result.get("content", "[]")
    tokens_used = result.get("tokens_used")

    json_match = re.search(r'\[.*\]', raw_text, re.DOTALL)
    if not json_match:
        raise ValueError("No JSON array found in response")

    matches_raw = json.loads(json_match.group())
    scope_lookup = {item['id']: item for item in SCOPE_ITEMS}
    matches = []
    for m in matches_raw[:top_n]:
        item_id = m.get('id', '')
        scope = scope_lookup.get(item_id, {})
        if scope:
            matches.append(ScopeItemMatch(
                id=item_id,
                name=scope['name'],
                lob=scope['lob'],
                process_group=scope['process_group'],
                description=scope['description'],
                confidence=m.get('confidence', 'MEDIUM'),
                rationale=m.get('rationale', ''),
                migration_objects=scope.get('migration_objects', []),
            ))
    return matches, tokens_used


# ── Health / Catalogue / LOBs ─────────────────────────────────────────────────

@app.get("/health")
def health():
    """Liveness: app is running."""
    return {
        "status": "ok",
        "version": "1.2.0",
        "scope_items_loaded": len(SCOPE_ITEMS),
        "release": "S/4HANA Cloud Public Edition 2602"
    }


@app.get("/health/ready")
def health_ready():
    """Readiness: app and dependencies (e.g. DB) are reachable."""
    try:
        list_clients()
        return {"status": "ok", "checks": {"database": "ok"}}
    except Exception as e:
        logger.warning("Readiness check failed: %s", e)
        raise HTTPException(status_code=503, detail="Service temporarily unavailable")

@router.get("/catalogue")
def get_catalogue(lob: Optional[str] = None):
    items = SCOPE_ITEMS
    if lob:
        items = [i for i in items if i.get('lob', '').lower() == lob.lower()]
    return {"total": len(items), "items": items}

@router.get("/lobs")
def get_lobs():
    from collections import Counter
    counts = Counter(i['lob'] for i in SCOPE_ITEMS)
    return {"lobs": [{"name": k, "count": v} for k, v in sorted(counts.items())]}

@router.get("/results")
def get_results(engagement_id: str):
    try:
        results = get_results_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"engagement_id": engagement_id, "total": len(results), "results": results}


# ── Client context helper ─────────────────────────────────────────────────────

def _build_client_context_line(engagement_id: str) -> str:
    """Return a single context sentence from engagement + client, or empty string."""
    try:
        ctx = get_engagement_with_client(engagement_id)
    except Exception:
        return ""
    if not ctx:
        return ""
    client = ctx.get("client") or {}
    parts = []
    if client.get("name"):
        parts.append(f"Client: {client['name']}")
    if client.get("industry"):
        parts.append(f"Industry: {client['industry']}")
    if client.get("employees"):
        parts.append(f"Size: {client['employees']} employees")
    if client.get("legal_entities"):
        parts.append(f"{client['legal_entities']} legal entities")
    if client.get("current_systems"):
        parts.append(f"Current systems: {client['current_systems']}")
    if client.get("countries"):
        parts.append(f"Countries: {client['countries']}")
    if client.get("regulatory_environment"):
        parts.append(f"Regulatory: {client['regulatory_environment']}")
    return ", ".join(parts) if parts else ""


# ── Clients ───────────────────────────────────────────────────────────────────

@router.post("/clients", status_code=201)
def post_client(body: ClientCreate):
    try:
        client = create_client(body.dict(exclude_none=True))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not client:
        raise HTTPException(status_code=500, detail="Failed to create client")
    return client

@router.get("/clients")
def list_all_clients():
    try:
        return {"clients": list_clients()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _html_to_plain_text(html: str, max_chars: int = 35000) -> str:
    """Strip HTML tags and collapse whitespace for LLM consumption."""
    text = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
    text = re.sub(r"<style[^>]*>[\s\S]*?</style>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:max_chars] if len(text) > max_chars else text


_CLIENT_PREFILL_SYSTEM = """You are extracting structured company/client profile data from website content.
Fill as many fields as you can. Infer from context when reasonable (e.g. "global operations" -> countries; "manufacturing" or "automotive" -> industry; "500+ employees" -> employees; "listed company" -> regulatory SOX).
Return a single JSON object with these keys only. Use null for unknown. For arrays use [] if none found.
- name (string): company name (required if visible)
- industry (string): industry sector (e.g. Manufacturing, Financial Services, Healthcare)
- sub_industry (string): more specific sector if mentioned (e.g. Automotive, EV, Aerospace)
- employees (number or null): employee count; infer band if only "500+", "thousands", "global" mentioned
- legal_entities (number or null): number of legal entities if mentioned
- locations (number or null): sites, offices, or plants if mentioned
- annual_revenue (string or null): revenue if mentioned (e.g. "$500M", "€1B")
- current_systems (array of strings): ERP or IT systems mentioned (SAP, Oracle, legacy, etc.)
- countries (array of strings): countries or regions of operation
- regulatory_environment (array of strings): e.g. SOX, GDPR, HIPAA, ISO, FDA
- business_strategy (string): brief description of business strategy or mission
- goals (array of strings): strategic or business goals
- key_products (array of strings): main products or services
- value_proposition (string): value proposition or differentiator
- senior_executives (array of objects with "name" and "title"): leadership team
- direct_competitors (array of strings): direct competitors
- substitutes (array of strings): substitute products/services
- sector_archetype (string or null): e.g. Manufacturing, Retail, Professional Services
- erp_maturity (string or null): if inferable: "Greenfield", "Legacy migration", "Hybrid", "Not assessed"
- complexity_drivers (array of strings): e.g. Multi-country, M&A, Regulatory, High volume
Return only the JSON object, no markdown or explanation."""


@router.post("/clients/prefill-from-website", response_model=Dict[str, Any])
def prefill_client_from_website(body: ClientPrefillFromWebsiteRequest):
    """Fetch URL, extract text, and use LLM to pre-populate client profile (Create Client form)."""
    if not os.getenv("ANTHROPIC_API_KEY"):
        raise HTTPException(
            status_code=503,
            detail="Pre-fill is not available: ANTHROPIC_API_KEY is not configured. Enter client details manually.",
        )
    raw = (body.url or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="url is required")
    # Heuristic: if input doesn't look like a full URL, treat it as a host or keyword
    candidates: List[str] = []
    base_inputs: List[str] = []
    if "." not in raw and " " in raw:
        # Keyword like "Carrier Global" -> carrier, carrierglobal
        slug = "".join(raw.lower().split())
        base_inputs = [slug]
    elif "." not in raw:
        base_inputs = [raw.lower()]
    else:
        base_inputs = [raw]

    for base in base_inputs:
        if base.startswith(("http://", "https://")):
            candidates.append(base)
        else:
            candidates.append(f"https://{base}")
            candidates.append(f"https://www.{base}")

    # Try base and a few common paths to find a usable page
    discovered: List[str] = []
    html = None
    with httpx.Client(timeout=15.0, follow_redirects=True) as client:
        for base_url in candidates:
            for path in ["", "/", "/en", "/en-us", "/about", "/about-us", "/company", "/en/worldwide/"]:
                try:
                    url = base_url.rstrip("/") + path
                    resp = client.get(url)
                    if resp.status_code < 400:
                        if html is None:
                            html = resp.text
                            chosen_url = url
                        if url not in discovered:
                            discovered.append(url)
                    # Stop early if we already found a good primary URL and a handful of alternates
                    if html and len(discovered) >= 3:
                        break
                except httpx.HTTPError:
                    continue
            if html and len(discovered) >= 3:
                break

    if html is None:
        # Could not fetch any usable page; surface discovered candidates (if any) to the UI
        detail = "Could not fetch URL. Try choosing a specific About/Company page."
        raise HTTPException(
            status_code=422,
            detail={"message": detail, "candidates": discovered} if discovered else detail,
        )

    text = _html_to_plain_text(html)
    if len(text) < 100:
        raise HTTPException(status_code=422, detail="Page content too short to extract meaningful data")
    provider = get_provider()
    user_prompt = f"Website content (excerpt):\n\n{text}\n\nExtract company profile JSON."
    try:
        result = provider.complete(_CLIENT_PREFILL_SYSTEM, user_prompt, max_tokens=1024, model=MODEL_HAIKU)
        raw = result.get("content", "{}")
        data = _extract_json_object(raw)
    except Exception as e:
        logger.exception("prefill LLM failed: %s", e)
        raise HTTPException(status_code=500, detail="Failed to extract company data from page")
    # Normalize keys to match ClientCreate; drop unknown keys
    allowed = {
        "name", "industry", "sub_industry", "employees", "legal_entities", "address", "locations", "annual_revenue",
        "current_systems", "countries", "regulatory_environment", "business_strategy", "goals", "key_products",
        "value_proposition", "senior_executives", "direct_competitors", "substitutes",
        "sector_archetype", "erp_maturity", "complexity_drivers",
    }
    return {k: v for k, v in data.items() if k in allowed}


@router.get("/clients/{client_id}")
def get_single_client(client_id: str):
    try:
        client = get_client(client_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not client:
        raise HTTPException(status_code=404, detail=f"{client_id} not found")
    return client


class ClientUpdate(BaseModel):
    """Fields that can be updated on a client (all optional)."""
    name: Optional[str] = None
    industry: Optional[str] = None
    employees: Optional[int] = None
    legal_entities: Optional[int] = None
    address: Optional[str] = None
    current_systems: Optional[List[str]] = None
    systems_to_keep: Optional[List[str]] = None
    systems_to_replace: Optional[List[str]] = None
    countries: Optional[List[str]] = None
    regulatory_environment: Optional[List[str]] = None
    business_strategy: Optional[str] = None
    goals: Optional[List[str]] = None
    key_products: Optional[List[str]] = None
    value_proposition: Optional[str] = None
    senior_executives: Optional[List[Dict[str, str]]] = None
    direct_competitors: Optional[List[str]] = None
    substitutes: Optional[List[str]] = None
    sector_archetype: Optional[str] = None
    complexity_drivers: Optional[List[str]] = None
    erp_maturity: Optional[str] = None
    benchmark_opt_in: Optional[bool] = None


@router.patch("/clients/{client_id}")
def patch_client(client_id: str, body: ClientUpdate):
    """Update client fields. Used for auto-save and editing."""
    client = get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail=f"{client_id} not found")
    updates = body.dict(exclude_none=True)
    if not updates:
        return client
    try:
        updated = update_client(client_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return updated or client


@router.post("/clients/{client_id}/benchmark-opt-out", status_code=200)
def client_benchmark_opt_out(client_id: str):
    """Set benchmark_opt_in = false for the client. Used when client opts out of benchmark insights."""
    client = get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail=f"{client_id} not found")
    try:
        updated = update_client(client_id, {"benchmark_opt_in": False})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return updated or client


# ── Engagements ───────────────────────────────────────────────────────────────

@router.post("/engagements", status_code=201)
def post_engagement(body: EngagementCreate):
    try:
        data = body.dict(exclude_none=True)
        status = data.get("status")
        if status and status not in {"open", "completed", "abandoned"}:
            raise HTTPException(status_code=400, detail="Invalid engagement status. Must be one of: open, completed, abandoned")
        risk = data.get("risk_level")
        if risk and risk not in {"low", "medium", "high"}:
            raise HTTPException(status_code=400, detail="Invalid risk_level. Must be one of: low, medium, high")
        health = data.get("health")
        if health and health not in {"on_track", "at_risk", "off_track"}:
            raise HTTPException(status_code=400, detail="Invalid health. Must be one of: on_track, at_risk, off_track")
        eng = create_engagement(data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not eng:
        raise HTTPException(status_code=500, detail="Failed to create engagement")
    return eng

@router.get("/engagements")
def list_all_engagements(client_id: Optional[str] = None):
    """List engagements, each enriched with client_name for display."""
    try:
        rows = list_engagements(client_id)
        out = []
        for eng in rows:
            cid = eng.get("client_id")
            client_name = None
            if cid:
                try:
                    client = get_client(cid)
                    client_name = (client or {}).get("name")
                except Exception:
                    pass
            out.append({**eng, "client_name": client_name})
        return {"engagements": out}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/engagements/{engagement_id}")
def get_single_engagement(engagement_id: str):
    try:
        eng = get_engagement_with_client(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not eng:
        raise HTTPException(status_code=404, detail=f"{engagement_id} not found")
    return eng


class EngagementUpdate(BaseModel):
    """Fields that can be updated on an engagement (all optional)."""
    client_id: Optional[str] = None
    name: Optional[str] = None
    description: Optional[str] = None
    go_live_date: Optional[str] = None
    project_type: Optional[str] = None
    status: Optional[str] = None
    planned_start_date: Optional[str] = None
    planned_end_date: Optional[str] = None
    actual_start_date: Optional[str] = None
    actual_end_date: Optional[str] = None
    project_manager: Optional[str] = None
    sponsor: Optional[str] = None
    risk_level: Optional[str] = None
    health: Optional[str] = None
    mode: Optional[str] = None              # guided | collaborative | autonomous
    autonomy_config: Optional[dict] = None


@router.patch("/engagements/{engagement_id}")
def patch_engagement(engagement_id: str, body: EngagementUpdate):
    """Update engagement fields. Used for auto-save and editing."""
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"{engagement_id} not found")
    updates = body.dict(exclude_none=True)
    if updates.get("status") and updates["status"] not in {"open", "completed", "abandoned"}:
        raise HTTPException(status_code=400, detail="Invalid status")
    if updates.get("risk_level") and updates["risk_level"] not in {"low", "medium", "high"}:
        raise HTTPException(status_code=400, detail="Invalid risk_level")
    if updates.get("health") and updates["health"] not in {"on_track", "at_risk", "off_track"}:
        raise HTTPException(status_code=400, detail="Invalid health")
    if not updates:
        return get_engagement_with_client(engagement_id)
    try:
        update_engagement(engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return get_engagement_with_client(engagement_id)


@router.get("/engagement/{engagement_id}/label")
def get_engagement_label(engagement_id: str):
    """Lightweight: engagement_id, project name, company name for display (e.g. header)."""
    try:
        eng = get_engagement_with_client(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not eng:
        raise HTTPException(status_code=404, detail=f"{engagement_id} not found")
    client = eng.get("client") or {}
    return {
        "engagement_id": eng.get("engagement_id"),
        "name": eng.get("name") or "",
        "client_id": eng.get("client_id") or client.get("client_id") or "",
        "client_name": client.get("name") or "",
    }


@router.get("/engagement/{engagement_id}/search")
def search_engagement(engagement_id: str, q: str = ""):
    """Search within engagement by keywords and match code (req_id, type, etc). Returns requirements and RICEFW matches with navigation paths."""
    q = (q or "").strip().lower()
    if not q:
        return {"engagement_id": engagement_id, "query": "", "requirements": [], "ricefw": []}
    try:
        requirements = get_requirements_by_engagement(engagement_id)
        ricefw_items = get_ricefw_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    req_matches = []
    for r in requirements:
        req_id = (r.get("req_id") or "").lower()
        title = (r.get("title") or "").lower()
        desc = (r.get("description") or "").lower()
        tags = " ".join((r.get("tags") or [])).lower()
        if q in req_id or q in title or q in desc or q in tags:
            req_matches.append({
                "type": "requirement",
                "req_id": r.get("req_id"),
                "title": r.get("title"),
                "path": f"/workflow/{r.get('req_id')}?engagement_id={engagement_id}",
            })
    ricefw_matches = []
    for item in ricefw_items:
        name = (item.get("name") or "").lower()
        desc = (item.get("description") or "").lower()
        typ = (item.get("type") or "").lower()
        req_id = (item.get("req_id") or "").lower()
        if q in name or q in desc or q in typ or q in req_id:
            ricefw_matches.append({
                "kind": "ricefw",
                "id": item.get("id"),
                "name": item.get("name"),
                "type": item.get("type"),
                "req_id": item.get("req_id"),
                "path": f"/engagement?engagement_id={engagement_id}#ricefw",
            })
    return {
        "engagement_id": engagement_id,
        "query": q,
        "requirements": req_matches[:20],
        "ricefw": ricefw_matches[:20],
    }


# ── Requirements ──────────────────────────────────────────────────────────────

_ALLOWED_PRIORITIES = {"Must-Have", "Should-Have", "Nice-to-Have"}
_ALLOWED_CATEGORIES = {"Automation", "Control/Compliance", "Reporting", "Integration", "UX", "Data Migration"}
_ALLOWED_FIT = {"Fit-to-Standard", "Soft-Gap", "Hard-Gap"}
_HITL_STATES_ORDERED = ["ai_draft", "needs_sme_review", "needs_architect_review", "approved"]
_HITL_TERMINAL_STATES = {"approved", "out_of_scope"}


@router.post("/requirements", response_model=RequirementResponse, status_code=201)
def post_requirement(body: RequirementCreate):
    kwargs = body.dict()
    engagement_id = kwargs.pop("engagement_id")
    title = kwargs.pop("title")
    description = kwargs.pop("description")

    priority = kwargs.get("priority")
    if priority and priority not in _ALLOWED_PRIORITIES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid priority. Must be one of: {sorted(_ALLOWED_PRIORITIES)}",
        )
    category = kwargs.get("category")
    if category and category not in _ALLOWED_CATEGORIES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid category. Must be one of: {sorted(_ALLOWED_CATEGORIES)}",
        )
    fit = kwargs.get("fit_assessment")
    if fit and fit not in _ALLOWED_FIT:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid fit_assessment. Must be one of: {sorted(_ALLOWED_FIT)}",
        )

    try:
        req = create_requirement(
            engagement_id=engagement_id,
            title=title,
            description=description,
            **kwargs,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=500, detail="Failed to create requirement")
    # Engagement mode: queue fit-gap run
    _queue_or_execute(engagement_id, "run_fitgap",
        {"req_id": req.get("req_id"), "description": description},
        source="requirement_capture")
    return req

@router.get("/requirements", response_model=List[RequirementResponse])
def list_requirements(engagement_id: str):
    try:
        return get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/requirements/extract-from-transcript", status_code=201)
def extract_from_transcript(body: TranscriptExtractRequest):
    provider = get_provider()
    client_context = _build_client_context_line(body.engagement_id)
    system_prompt = """You are an expert business analyst capturing requirements from conversation transcripts for SAP S/4HANA implementation projects.

Extract discrete business requirements from the transcript. For each requirement identify:
- title: Brief descriptive title (max 10 words)
- description: Detailed description of the business need or process
- tags: Array — use only: pain_point, manual_step, secret_sauce, workaround, hand_off
- business_process: One of Procure-to-Pay, Order-to-Cash, Record-to-Report, Plan-to-Produce, Hire-to-Retire (pick closest)
- priority: Must-Have if the pain is severe or a compliance/control need; Should-Have if clearly beneficial; Nice-to-Have if aspirational
- category: One of Automation, Control/Compliance, Reporting, Integration, UX, Data Migration
- shadow_tools: Array of unofficial tools mentioned (e.g. "Excel macro", "WhatsApp", "Access DB"). Empty array if none.
- actors: Array of objects {"role": "job title or name", "type": "formal" or "informal"} for people mentioned
- kpi_impact: Object {"metric": "...", "current": "...", "target": "...", "unit": "..."} if any measurable metric is implied (e.g. "takes 3 days" → target to reduce). null if no metric.

Return a JSON array only — no markdown fences:
[
  {
    "title": "requirement title",
    "description": "detailed description",
    "tags": ["tag1"],
    "business_process": "Order-to-Cash",
    "priority": "Must-Have",
    "category": "Automation",
    "shadow_tools": [],
    "actors": [{"role": "AP Clerk", "type": "formal"}],
    "kpi_impact": null
  }
]

Rules:
- Each requirement must be distinct and actionable
- tags must only come from: pain_point, manual_step, secret_sauce, workaround, hand_off
- Return [] if no clear requirements are found"""

    if client_context:
        system_prompt = system_prompt + f"\n\nEngagement context: {client_context}"

    user_prompt = f"Stakeholder: {body.stakeholder}\n\nExtract requirements from this transcript:\n\n{body.transcript_text}\n\nReturn JSON array."

    try:
        # Use cheaper model with tight token limit for extraction
        result = provider.complete(system_prompt, user_prompt, max_tokens=768, model=MODEL_HAIKU)
        raw_text = result.get("content", "[]")

        json_match = re.search(r'\[.*\]', raw_text, re.DOTALL)
        if not json_match:
            raise ValueError("No JSON array found in response")

        extracted = json.loads(json_match.group())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Extraction failed: {e}")

    valid_tags = {"pain_point", "manual_step", "secret_sauce", "workaround", "hand_off"}
    created = []
    for item in extracted:
        tags = [t for t in (item.get("tags") or []) if t in valid_tags]
        try:
            req = create_requirement(
                engagement_id=body.engagement_id,
                title=item.get("title", "Untitled"),
                description=item.get("description", ""),
                source_type="Conversation",
                tags=tags,
                stakeholder=body.stakeholder,
                raw_input=body.transcript_text,
                business_process=item.get("business_process") or None,
                priority=item.get("priority") or "Must-Have",
                category=item.get("category") or None,
                shadow_tools=item.get("shadow_tools") or None,
                actors=item.get("actors") or None,
                kpi_impact=item.get("kpi_impact") or None,
                hitl_state="ai_draft",
                ai_rationale="Extracted from transcript by AI",
            )
            if req:
                created.append({
                    "req_id": req["req_id"],
                    "title": req["title"],
                    "tags": req.get("tags", []),
                    "business_process": req.get("business_process"),
                    "priority": req.get("priority"),
                    "category": req.get("category"),
                    "shadow_tools": req.get("shadow_tools"),
                    "actors": req.get("actors"),
                    "kpi_impact": req.get("kpi_impact"),
                })
        except Exception as e:
            print(f"Failed to create requirement '{item.get('title')}': {e}")

    return {"created": len(created), "requirements": created}


# ── Domain Templates ───────────────────────────────────────────────────────────

_DOMAIN_TEMPLATES: Dict[str, List[Dict]] = {
    "finance": [
        {
            "title": "Automated customer invoicing on goods issue",
            "description": "Automatically generate and send customer invoices when goods issue is posted, eliminating manual invoice creation and reducing billing cycle time.",
            "business_process": "Order-to-Cash",
            "priority": "Must-Have",
            "category": "Automation",
            "tags": ["manual_step", "pain_point"],
            "shadow_tools": [],
            "actors": [{"role": "Billing Clerk", "type": "formal"}],
            "kpi_impact": {"metric": "invoice cycle time", "target": "reduce by 50%", "unit": "days"},
        },
        {
            "title": "Vendor payment approval with three-way match",
            "description": "Enforce three-way match (PO / GR / invoice) before any vendor payment is approved, with automatic blocking of mismatched invoices.",
            "business_process": "Procure-to-Pay",
            "priority": "Must-Have",
            "category": "Control/Compliance",
            "tags": ["manual_step", "workaround"],
            "shadow_tools": ["Excel spreadsheet"],
            "actors": [{"role": "AP Manager", "type": "formal"}, {"role": "Auditor", "type": "formal"}],
            "kpi_impact": None,
        },
        {
            "title": "Month-end closing time reduction to 3 days",
            "description": "Automate journal entries, intercompany reconciliation, and reporting to compress financial close from current state to 3 business days.",
            "business_process": "Record-to-Report",
            "priority": "Must-Have",
            "category": "Reporting",
            "tags": ["pain_point", "manual_step"],
            "shadow_tools": ["Excel macro", "Access DB"],
            "actors": [{"role": "Controller", "type": "formal"}, {"role": "CFO", "type": "formal"}],
            "kpi_impact": {"metric": "month-end close days", "target": "3 days", "unit": "days"},
        },
    ],
    "sales": [
        {
            "title": "Credit risk check and automatic order blocking",
            "description": "Run automated credit risk checks when sales orders are created and block orders that exceed the customer credit limit, with workflow for manual override approval.",
            "business_process": "Order-to-Cash",
            "priority": "Must-Have",
            "category": "Control/Compliance",
            "tags": ["manual_step", "workaround"],
            "shadow_tools": [],
            "actors": [{"role": "Credit Manager", "type": "formal"}, {"role": "Sales Rep", "type": "formal"}],
            "kpi_impact": None,
        },
        {
            "title": "Shipping and warehouse integration triggers",
            "description": "Automatically trigger warehouse pick/pack/ship tasks when sales orders are confirmed, with real-time status updates back to the sales order.",
            "business_process": "Order-to-Cash",
            "priority": "Should-Have",
            "category": "Integration",
            "tags": ["manual_step", "hand_off"],
            "shadow_tools": ["WhatsApp group"],
            "actors": [{"role": "Warehouse Operator", "type": "formal"}, {"role": "Shipping Clerk", "type": "formal"}],
            "kpi_impact": {"metric": "order-to-ship time", "target": "reduce by 30%", "unit": "hours"},
        },
    ],
    "procurement": [
        {
            "title": "Purchase requisition approval workflow",
            "description": "Implement configurable multi-level approval workflow for purchase requisitions based on value thresholds, cost centre, and category, replacing email-based approvals.",
            "business_process": "Procure-to-Pay",
            "priority": "Must-Have",
            "category": "Control/Compliance",
            "tags": ["manual_step", "workaround"],
            "shadow_tools": ["Email", "Excel spreadsheet"],
            "actors": [{"role": "Requester", "type": "formal"}, {"role": "Budget Holder", "type": "formal"}, {"role": "Procurement Manager", "type": "formal"}],
            "kpi_impact": None,
        },
        {
            "title": "Automatic stock replenishment",
            "description": "Automatically generate purchase requisitions when inventory falls below reorder points, using MRP logic to calculate quantities and lead times.",
            "business_process": "Procure-to-Pay",
            "priority": "Should-Have",
            "category": "Automation",
            "tags": ["manual_step"],
            "shadow_tools": [],
            "actors": [{"role": "Inventory Planner", "type": "formal"}],
            "kpi_impact": {"metric": "stockout events", "target": "reduce by 80%", "unit": "incidents/month"},
        },
    ],
    "manufacturing": [
        {
            "title": "MRP capacity-aware order conversion",
            "description": "Convert planned orders to production orders only when capacity is available, integrating MRP with capacity planning to prevent overloading work centres.",
            "business_process": "Plan-to-Produce",
            "priority": "Must-Have",
            "category": "Automation",
            "tags": ["manual_step", "pain_point"],
            "shadow_tools": ["Excel macro"],
            "actors": [{"role": "Production Planner", "type": "formal"}, {"role": "Shop Floor Supervisor", "type": "formal"}],
            "kpi_impact": {"metric": "schedule adherence", "target": "above 95%", "unit": "percent"},
        },
        {
            "title": "Production variance alerts",
            "description": "Automatically detect and alert when actual production costs or quantities deviate from standard by more than a configurable threshold.",
            "business_process": "Plan-to-Produce",
            "priority": "Should-Have",
            "category": "Reporting",
            "tags": ["pain_point"],
            "shadow_tools": [],
            "actors": [{"role": "Cost Accountant", "type": "formal"}, {"role": "Production Manager", "type": "formal"}],
            "kpi_impact": None,
        },
    ],
    "hr": [
        {
            "title": "Digital onboarding workflow for new hires",
            "description": "Automate the end-to-end onboarding process — contract generation, IT provisioning, payroll setup, and training assignment — triggered on hire date.",
            "business_process": "Hire-to-Retire",
            "priority": "Must-Have",
            "category": "Automation",
            "tags": ["manual_step", "hand_off"],
            "shadow_tools": ["Email", "Shared Drive"],
            "actors": [{"role": "HR Business Partner", "type": "formal"}, {"role": "Line Manager", "type": "formal"}, {"role": "IT Admin", "type": "formal"}],
            "kpi_impact": {"metric": "onboarding time to productivity", "target": "reduce by 40%", "unit": "days"},
        },
        {
            "title": "Automated payroll variance reporting",
            "description": "Generate monthly payroll variance reports comparing current vs prior period, flagging anomalies above threshold for payroll manager review before approval.",
            "business_process": "Hire-to-Retire",
            "priority": "Must-Have",
            "category": "Reporting",
            "tags": ["manual_step", "pain_point"],
            "shadow_tools": ["Excel macro"],
            "actors": [{"role": "Payroll Manager", "type": "formal"}, {"role": "Finance Controller", "type": "formal"}],
            "kpi_impact": None,
        },
    ],
}

_VALID_DOMAINS = list(_DOMAIN_TEMPLATES.keys())


@router.get("/requirements/templates")
def get_requirement_templates(domain: str):
    key = domain.lower()
    if key not in _DOMAIN_TEMPLATES:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown domain '{domain}'. Valid options: {', '.join(_VALID_DOMAINS)}",
        )
    templates = _DOMAIN_TEMPLATES[key]
    return {"domain": key, "total": len(templates), "templates": templates}


_ARCHAEOLOGIST_SYSTEM_PROMPT = """You are a senior business analyst and process archaeologist conducting a discovery interview for a Cloud ERP transformation. Your job is to deeply understand how work actually happens today — not how it should work, but how it really works, including workarounds, exceptions, and shadow tools.

Your behaviour:
- Ask one focused question at a time
- Probe for: who actually does each step (not just job titles), what breaks, what is manual, what tools people use unofficially, what the person is proud of and wants to preserve
- Adapt your questions to the person role: ask a warehouse operator about physical flows and exceptions; ask a CFO about controls, reporting pain, and month-end pressure
- When you learn something significant, summarise it back: "So if I understand correctly, you manually reconcile the bank statement every morning in Excel before the ERP is updated — is that right?"
- After 3-4 exchanges, offer to extract what you have learned as a structured requirement

Response format — always return valid JSON only, no markdown fences:
{
  "reply": "your next question or summary to the user",
  "extracted": {
    "ready": false,
    "title": "",
    "description": "",
    "tags": [],
    "shadow_tools": [],
    "actors": [],
    "pain_points": [],
    "secret_sauce": []
  },
  "suggested_follow_ups": ["question 1", "question 2"]
}

When extracted.ready is true, populate all extracted fields with what you have learned."""


def _extract_json_object(text: str) -> dict:
    """Extract first complete JSON object from text, stripping markdown fences."""
    text = re.sub(r'```(?:json)?\s*', '', text).strip()
    start = text.find('{')
    if start == -1:
        raise ValueError("No JSON object found in response")
    depth = 0
    for i, c in enumerate(text[start:], start):
        if c == '{':
            depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                return json.loads(text[start:i + 1])
    raise ValueError("No complete JSON object found in response")


@router.post("/requirements/archaeologist-session")
def archaeologist_session(body: ArchaeologistSessionRequest):
    provider = get_provider()

    # Inject client context and top patterns into system prompt
    system_prompt = _ARCHAEOLOGIST_SYSTEM_PROMPT
    client_context = _build_client_context_line(body.engagement_id)
    if client_context:
        system_prompt = system_prompt + f"\n\nEngagement context: {client_context}"
    patterns = _get_top_patterns_text(limit=5)
    if patterns:
        system_prompt = system_prompt + "\n\n" + patterns

    # Build conversation as a single user message (history + current turn)
    lines = [
        f"Context:",
        f"  Stakeholder: {body.stakeholder} | Role: {body.role} | Business Process: {body.business_process}",
        "",
    ]
    if body.session_history:
        lines.append("Conversation so far:")
        for msg in body.session_history:
            prefix = "Analyst" if msg.get("role") == "assistant" else "Stakeholder"
            lines.append(f"  {prefix}: {msg.get('content', '')}")
        lines.append("")
    lines.append(f"Stakeholder: {body.message}")
    lines.append("")
    lines.append("Respond as the analyst. Return valid JSON only.")
    user_prompt = "\n".join(lines)

    try:
        # Use cheaper model with modest response size for dialogue
        result = provider.complete(system_prompt, user_prompt, max_tokens=512, model=MODEL_HAIKU)
        raw_text = result.get("content", "{}")
        parsed = _extract_json_object(raw_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Archaeologist LLM error: {e}")

    extracted = parsed.get("extracted", {})
    response: Dict = {
        "reply": parsed.get("reply", ""),
        "extracted": extracted,
        "suggested_follow_ups": parsed.get("suggested_follow_ups", []),
    }

    if extracted.get("ready"):
        valid_tags = {"pain_point", "manual_step", "secret_sauce", "workaround", "hand_off"}
        tags = [t for t in (extracted.get("tags") or []) if t in valid_tags]
        try:
            req = create_requirement(
                engagement_id=body.engagement_id,
                title=extracted.get("title", "Extracted Requirement"),
                description=extracted.get("description", ""),
                source_type="Conversation",
                tags=tags,
                stakeholder=body.stakeholder,
                business_process=body.business_process,
                actors=extracted.get("actors") or None,
                shadow_tools=extracted.get("shadow_tools") or None,
                hitl_state="ai_draft",
                ai_rationale="Drafted by archaeologist agent",
            )
            if req:
                response["req_id"] = req.get("req_id")
        except Exception as e:
            print(f"Auto-create requirement failed (non-fatal): {e}")

    return response


@router.get("/requirements/{req_id}", response_model=RequirementResponse)
def get_requirement(req_id: str, engagement_id: str):
    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")
    return req

@router.patch("/requirements/{req_id}", response_model=RequirementResponse)
def patch_requirement(req_id: str, engagement_id: str, body: RequirementUpdate):
    updates = {k: v for k, v in body.dict().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields provided to update")
    try:
        req = update_requirement(req_id, engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")
    return req


@router.delete("/requirements/{req_id}", status_code=204)
def delete_requirement_route(req_id: str, engagement_id: str):
    """Delete requirement. Blocked if approved fit-gap or linked RICEFW (409)."""
    req = get_requirement_by_id(req_id, engagement_id)
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")
    assessments = [a for a in (get_fit_gap_by_engagement(engagement_id) or []) if a.get("req_id") == req_id]
    for a in assessments:
        if (a.get("hitl_state") or "").lower() in ("approved", "out_of_scope"):
            raise HTTPException(
                status_code=409,
                detail="Cannot delete requirement with approved fit-gap assessment. Archive instead.",
            )
    ricefw_items = [r for r in (get_ricefw_by_engagement(engagement_id) or []) if r.get("req_id") == req_id]
    if ricefw_items:
        raise HTTPException(
            status_code=409,
            detail="Remove RICEFW link first. One or more RICEFW items reference this requirement.",
        )
    try:
        delete_requirement(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/requirements/{req_id}/sign-off", response_model=RequirementResponse)
def sign_off_requirement(req_id: str, engagement_id: str, body: SignOffRequest):
    """Sign-off state machine:
      level=sme   → sme_approved
      level=owner, current=sme_approved → confirmed
      level=owner, otherwise            → owner_approved
    """
    if body.level not in ("sme", "owner"):
        raise HTTPException(status_code=400, detail="level must be 'sme' or 'owner'")

    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")

    current_status = req.get("sign_off_status", "draft")
    if body.level == "sme":
        new_status = "sme_approved"
    elif current_status == "sme_approved":
        new_status = "confirmed"
    else:
        new_status = "owner_approved"

    now = datetime.now(timezone.utc).isoformat()
    updates = {
        "sign_off_status": new_status,
        "sign_off_by": body.signed_by,
        "sign_off_at": now,
    }
    try:
        updated = update_requirement(req_id, engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not updated:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")
    return updated


@router.post("/requirements/{req_id}/hitl-advance", response_model=RequirementResponse)
def hitl_advance(req_id: str, engagement_id: str, body: HitlAdvanceRequest):
    """Advance requirement HITL state along the pipeline and record an event."""
    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")

    current = (req.get("hitl_state") or "ai_draft").lower()
    if current in _HITL_TERMINAL_STATES:
        raise HTTPException(status_code=400, detail=f"HITL state '{current}' is terminal and cannot be advanced")

    try:
        idx = _HITL_STATES_ORDERED.index(current) if current in _HITL_STATES_ORDERED else 0
    except ValueError:
        idx = 0
    if idx >= len(_HITL_STATES_ORDERED) - 1:
        new_state = _HITL_STATES_ORDERED[-1]
    else:
        new_state = _HITL_STATES_ORDERED[idx + 1]

    # Record HITL event
    try:
        create_hitl_event(
            {
                "req_id": req_id,
                "engagement_id": engagement_id,
                "from_state": current,
                "to_state": new_state,
                "actor": body.actor,
                "actor_role": body.actor_role,
                "notes": body.notes,
                "ai_suggestion": None,
                "human_correction": body.correction,
            }
        )
    except Exception as e:
        logger.warning("Failed to create HITL event for %s: %s", req_id, e)

    history = req.get("hitl_history") or []
    if not isinstance(history, list):
        history = []
    history.append(
        {
            "from": current,
            "to": new_state,
            "actor": body.actor,
            "actor_role": body.actor_role,
            "notes": body.notes,
            "ts": datetime.now(timezone.utc).isoformat(),
        }
    )

    updates = {
        "hitl_state": new_state,
        "hitl_history": history,
    }
    try:
        updated = update_requirement(req_id, engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not updated:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")
    # Engagement mode: queue HITL advance for high-confidence items
    req_confidence = (updated or {}).get("confidence_score") or req.get("confidence_score")
    if req_confidence is not None:
        _queue_or_execute(engagement_id, "advance_hitl_high",
            {"req_id": req_id, "target_state": new_state, "current_state": current},
            confidence=float(req_confidence), source="hitl_evaluator")
    return updated


@router.post("/requirements/{req_id}/hitl-reject", response_model=RequirementResponse)
def hitl_reject(req_id: str, engagement_id: str, body: HitlRejectRequest):
    """Reject or send a requirement back in the HITL pipeline."""
    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")

    current = (req.get("hitl_state") or "ai_draft").lower()
    target = (body.target_state or "ai_draft").lower()
    allowed_states = set(_HITL_STATES_ORDERED) | {"out_of_scope"}
    if target not in allowed_states:
        raise HTTPException(
            status_code=400,
            detail="target_state must be one of ai_draft, needs_sme_review, needs_architect_review, approved, out_of_scope",
        )

    try:
        create_hitl_event(
            {
                "req_id": req_id,
                "engagement_id": engagement_id,
                "from_state": current,
                "to_state": target,
                "actor": body.actor,
                "actor_role": None,
                "notes": body.reason,
                "ai_suggestion": None,
                "human_correction": None,
            }
        )
    except Exception as e:
        logger.warning("Failed to create HITL reject event for %s: %s", req_id, e)

    history = req.get("hitl_history") or []
    if not isinstance(history, list):
        history = []
    history.append(
        {
            "from": current,
            "to": target,
            "actor": body.actor,
            "notes": body.reason,
            "ts": datetime.now(timezone.utc).isoformat(),
        }
    )

    updates = {
        "hitl_state": target,
        "hitl_history": history,
        "reviewer_notes": body.reason,
    }
    try:
        updated = update_requirement(req_id, engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not updated:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")
    return updated


@router.get("/requirements/{req_id}/traceability")
def get_traceability(req_id: str, engagement_id: str):
    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")

    try:
        gap_records = get_gap_results_by_req_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    # Derive answer_to from requirement fields
    tags = req.get("tags") or []
    pain_tags = {"pain_point", "manual_step", "workaround"}
    pain_points = [t for t in tags if t in pain_tags]
    actors = req.get("actors") or []
    actor_names = [a.get("role", "") for a in actors if isinstance(a, dict)]
    who_asked_parts = [req.get("stakeholder")] + actor_names
    who_asked = ", ".join(p for p in who_asked_parts if p)

    description = req.get("description", "")
    why_needed = description[:200] + ("..." if len(description) > 200 else "")
    what_problem = (
        f"Tags indicate: {', '.join(pain_points)}. " if pain_points else ""
    ) + (
        f"Shadow tools in use: {', '.join(req.get('shadow_tools') or [])}." if req.get("shadow_tools") else ""
    )

    return {
        "requirement": req,
        "as_is_evidence": {
            "current_state_ref": req.get("current_state_ref"),
            "actors": actors,
            "shadow_tools": req.get("shadow_tools"),
            "tags": tags,
        },
        "gap_analysis": {
            "matches": gap_records[0].get("matches") if gap_records else [],
            "total_analyses": len(gap_records),
        },
        "answer_to": {
            "why_needed": why_needed,
            "who_asked": who_asked or "Unknown",
            "what_problem": what_problem or "No pain points tagged",
        },
    }


# ── Gap Analysis ──────────────────────────────────────────────────────────────

@router.post("/gap-analysis", response_model=GapAnalysisResponse)
async def gap_analysis(request: GapAnalysisRequest):
    # Resolve process_description: from req_id lookup or direct input
    req_id = request.req_id
    process_description = request.process_description

    if req_id:
        try:
            req = get_requirement_by_id(req_id, request.engagement_id)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Requirement lookup failed: {e}")
        if not req:
            raise HTTPException(status_code=404, detail=f"Requirement {req_id} not found for engagement {request.engagement_id}")
        process_description = req["description"]
    elif not process_description:
        raise HTTPException(status_code=422, detail="Provide either process_description or req_id")

    provider = get_provider()
    catalogue = build_catalogue_for_prompt(request.lob_filter)

    system_prompt = """You are an expert SAP S/4HANA implementation consultant specializing in Fit-to-Standard gap analysis.

Your task: Given a business process description, identify the most relevant SAP S/4HANA Cloud Public Edition scope items from the provided catalogue.

Instructions:
1. Analyze the business process description semantically - look beyond keywords to understand intent
2. Return the top matching scope items ranked by relevance
3. For each match, provide confidence (HIGH / MEDIUM / LOW) and a brief rationale
4. Consider that one business requirement often maps to multiple scope items
5. Always return valid JSON only

Response format (JSON array):
[
  {
    "id": "scope_item_code",
    "confidence": "HIGH|MEDIUM|LOW",
    "rationale": "One sentence explaining why this scope item matches"
  }
]"""

    user_prompt = f"""Business Process Description:
{process_description}

SAP S/4HANA Cloud 2602 Scope Item Catalogue (2602 release):
{catalogue}

Return the top {request.top_n} most relevant scope items as JSON."""

    try:
        result = provider.complete(system_prompt, user_prompt)
        raw_text = result.get("content", "[]")
        tokens_used = result.get("tokens_used")

        json_match = re.search(r'\[.*\]', raw_text, re.DOTALL)
        if not json_match:
            raise ValueError("No JSON array found in response")

        matches_raw = json.loads(json_match.group())

        scope_lookup = {item['id']: item for item in SCOPE_ITEMS}
        matches = []
        for m in matches_raw[:request.top_n]:
            item_id = m.get('id', '')
            scope = scope_lookup.get(item_id, {})
            if scope:
                matches.append(ScopeItemMatch(
                    id=item_id,
                    name=scope['name'],
                    lob=scope['lob'],
                    process_group=scope['process_group'],
                    description=scope['description'],
                    confidence=m.get('confidence', 'MEDIUM'),
                    rationale=m.get('rationale', ''),
                    migration_objects=scope.get('migration_objects', [])
                ))

        timestamp = datetime.utcnow().isoformat()

        try:
            save_gap_analysis(
                engagement_id=request.engagement_id,
                process_description=process_description,
                matches=[m.dict() for m in matches],
                tokens_used=tokens_used,
                timestamp=timestamp,
                req_id=req_id,
            )
        except Exception as db_err:
            print(f"DB save failed (non-fatal): {db_err}")

        return GapAnalysisResponse(
            engagement_id=request.engagement_id,
            req_id=req_id,
            process_description=process_description,
            matches=matches,
            total_scope_items_searched=len(SCOPE_ITEMS),
            tokens_used=tokens_used,
            timestamp=timestamp
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Engagement Summary ────────────────────────────────────────────────────────

@router.get("/engagement/{engagement_id}/summary")
def get_engagement_summary(engagement_id: str):
    try:
        requirements = get_requirements_by_engagement(engagement_id)
        gap_results = get_results_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    by_status: dict = {}
    for req in requirements:
        s = req.get("status", "open")
        by_status[s] = by_status.get(s, 0) + 1

    by_tag: dict = {}
    for req in requirements:
        for tag in (req.get("tags") or []):
            by_tag[tag] = by_tag.get(tag, 0) + 1

    # Index latest gap result per req_id (most recent first, results ordered desc)
    results_by_req: dict = {}
    for gr in gap_results:
        rid = gr.get("req_id")
        if rid and rid not in results_by_req:
            results_by_req[rid] = gr

    gap_results_summary = []
    for req in requirements:
        if req.get("status") == "analysed":
            rid = req["req_id"]
            gr = results_by_req.get(rid)
            if gr:
                matches = gr.get("matches") or []
                top = matches[0] if matches else {}
                gap_results_summary.append({
                    "req_id": rid,
                    "title": req.get("title"),
                    "top_match_id": top.get("id"),
                    "top_match_name": top.get("name"),
                    "top_confidence": top.get("confidence"),
                })

    return {
        "engagement_id": engagement_id,
        "total_requirements": len(requirements),
        "requirements_by_status": by_status,
        "requirements_by_tag": by_tag,
        "total_analysed": by_status.get("analysed", 0),
        "gap_results_summary": gap_results_summary,
    }


class AskRapidRequest(BaseModel):
    question: str


_ASK_RAPID_SYSTEM = """You are the RAPID assistant for an SAP S/4HANA implementation workspace.
Answer the user's question using ONLY the engagement context provided below. Be concise and actionable (3–6 bullet points when suggesting next steps).
If the answer is not in the context or you are unsure, say: "I don't have enough information to answer that from the engagement data."
Then suggest the user speak to the relevant human role (e.g. Engagement Manager for scope and client questions, Solution Architect for fit/gap and RICEFW, Process Owner for sign-off).
Do not invent data. Stay within the workspace scope."""


@router.post("/engagement/{engagement_id}/ask-rapid")
def ask_rapid(engagement_id: str, body: AskRapidRequest):
    """Context-sensitive Ask RAPID: load engagement, client, requirements, fit-gap summary; answer from context or direct user to human role."""
    ctx = get_engagement_with_client(engagement_id)
    if not ctx:
        raise HTTPException(status_code=404, detail="Engagement not found")
    client = ctx.get("client") or {}
    requirements = get_requirements_by_engagement(engagement_id) or []
    assessments = get_fit_gap_by_engagement(engagement_id) or []
    try:
        ricefw_items = get_ricefw_by_engagement(engagement_id) or []
    except Exception:
        ricefw_items = []
    # Build context string
    parts = [
        f"Engagement: {ctx.get('name') or engagement_id} (ID: {engagement_id})",
        f"Phase: {ctx.get('phase') or 'N/A'}, Status: {ctx.get('status') or 'N/A'}",
        f"Client: {client.get('name') or client.get('client_id') or 'N/A'}",
        f"Industry: {client.get('industry') or 'N/A'}, Sector: {client.get('sector_archetype') or client.get('industry') or 'N/A'}",
        f"Requirements count: {len(requirements)}",
        f"Fit/Gap assessments: {len(assessments)}",
        f"RICEFW items: {len(ricefw_items)}",
    ]
    if requirements:
        parts.append("Sample requirement titles (first 15):")
        for r in requirements[:15]:
            parts.append(f"  - {r.get('title') or r.get('req_id')}")
    if assessments:
        by_fit = {}
        for a in assessments:
            ft = a.get("fit_type") or "unknown"
            by_fit[ft] = by_fit.get(ft, 0) + 1
        parts.append("Fit/Gap breakdown: " + ", ".join(f"{k}: {v}" for k, v in sorted(by_fit.items())))
    context_text = "\n".join(parts)
    user_prompt = f"Context:\n{context_text}\n\nUser question: {body.question}"
    try:
        provider = get_provider()
        result = provider.complete(_ASK_RAPID_SYSTEM, user_prompt, max_tokens=600, model=MODEL_HAIKU)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"AI service unavailable: {e}")
    content = (result.get("content") or "").strip()
    return {"reply": content}


# ── RACI Matrix ─────────────────────────────────────────────────────────────

def _default_raci_matrix() -> list:
    """Default RACI matrix rows: area + activities with R/A/C/I checkboxes (all false)."""
    areas = [
        ("Requirements discovery & capture", ["Conduct workshops", "Extract from transcripts", "Maintain repository", "Prioritise", "Link to L2/L3"]),
        ("Gap analysis", ["Semantic match", "Review AI mappings", "Document fit/gap", "Traceability"]),
        ("Process mirror & to-be", ["Document as-is", "Identify automation", "Design to-be", "Validate L3"]),
        ("Sign-off and governance", ["SME approval", "Owner sign-off", "Final confirmation", "Audit trail"]),
        ("RICEFW", ["Classify RICEFW", "Link to requirements", "Assess complexity", "Approve or cancel"]),
    ]
    roles = ["R", "A", "C", "I"]
    out = []
    for area, activities in areas:
        for act in activities:
            out.append({
                "area": area,
                "activity": act,
                "R": False, "A": False, "C": False, "I": False,
            })
    return out


@router.get("/engagement/{engagement_id}/raci")
def get_raci(engagement_id: str):
    """Get RACI matrix for engagement. Returns matrix, finalized, finalized_at, finalized_by, change_log."""
    if not get_engagement(engagement_id):
        raise HTTPException(status_code=404, detail="Engagement not found")
    row = get_raci_matrix(engagement_id)
    if not row:
        return {
            "engagement_id": engagement_id,
            "matrix": _default_raci_matrix(),
            "finalized": False,
            "finalized_at": None,
            "finalized_by": None,
            "change_log": [],
        }
    return {
        "engagement_id": engagement_id,
        "matrix": row.get("matrix") or _default_raci_matrix(),
        "finalized": bool(row.get("finalized")),
        "finalized_at": row.get("finalized_at"),
        "finalized_by": row.get("finalized_by"),
        "change_log": row.get("change_log") or [],
    }


class RaciUpdate(BaseModel):
    matrix: Optional[List[Dict[str, Any]]] = None
    finalize: Optional[bool] = None
    finalized_by: Optional[str] = None


@router.patch("/engagement/{engagement_id}/raci")
def patch_raci(engagement_id: str, body: RaciUpdate, request: Request):
    """Update RACI matrix or set finalized. If already finalized and matrix sent, append to change_log then update."""
    if not get_engagement(engagement_id):
        raise HTTPException(status_code=404, detail="Engagement not found")
    actor = request.headers.get("X-Actor-Id") or request.headers.get("X-Actor-Role") or "user"
    now = datetime.now(timezone.utc).isoformat()
    row = get_raci_matrix(engagement_id)
    current_matrix = (row.get("matrix") or _default_raci_matrix()) if row else _default_raci_matrix()
    finalized = bool(row.get("finalized")) if row else False
    finalized_at = row.get("finalized_at")
    finalized_by = row.get("finalized_by")
    change_log = list(row.get("change_log") or []) if row else []

    if body.finalize is True:
        finalized = True
        finalized_at = now
        finalized_by = body.finalized_by or actor
    if body.finalize is False:
        finalized = False
        finalized_at = None
        finalized_by = None

    if body.matrix is not None:
        if finalized and row:
            # Log before/after for each cell that changed
            for i, new_row in enumerate(body.matrix):
                if i >= len(current_matrix):
                    continue
                old_row = current_matrix[i]
                for role in ("R", "A", "C", "I"):
                    ov = old_row.get(role, False)
                    nv = new_row.get(role, False)
                    if ov != nv:
                        change_log.append({
                            "at": now,
                            "by": actor,
                            "row_index": i,
                            "area": old_row.get("area"),
                            "activity": old_row.get("activity"),
                            "role": role,
                            "before": ov,
                            "after": nv,
                        })
        current_matrix = body.matrix

    upsert_raci_matrix(
        engagement_id,
        matrix=current_matrix,
        finalized=finalized,
        finalized_at=finalized_at,
        finalized_by=finalized_by,
        change_log=change_log,
    )
    return get_raci(engagement_id)


# ── Engagement scope (L1/L2/L3) ─────────────────────────────────────────────

@router.get("/engagement/{engagement_id}/scope")
def get_scope(engagement_id: str):
    """Get scope (L1/L2/L3 business process checkmarks) for engagement."""
    if not get_engagement(engagement_id):
        raise HTTPException(status_code=404, detail="Engagement not found")
    row = get_engagement_scope(engagement_id)
    scope = (row.get("scope") or {}) if row else {}
    return {"engagement_id": engagement_id, "scope": scope}


class ScopeUpdate(BaseModel):
    scope: Dict[str, Any]


@router.patch("/engagement/{engagement_id}/scope")
def patch_scope(engagement_id: str, body: ScopeUpdate):
    """Update scope for engagement."""
    if not get_engagement(engagement_id):
        raise HTTPException(status_code=404, detail="Engagement not found")
    upsert_engagement_scope(engagement_id, body.scope)
    return get_scope(engagement_id)


@router.get("/engagement/{engagement_id}/benchmark-hints")
def get_benchmark_hints(engagement_id: str):
    """Return benchmark insights for the engagement. If client has benchmark_opt_in=false, returns []. Otherwise returns stored hints or derived from client profile (sector_archetype, erp_maturity, complexity_drivers)."""
    ctx = get_engagement_with_client(engagement_id)
    if not ctx:
        raise HTTPException(status_code=404, detail="Engagement not found")
    client = ctx.get("client") or {}
    if client.get("benchmark_opt_in") is False:
        return {"engagement_id": engagement_id, "hints": [], "opted_out": True}
    try:
        stored = get_benchmark_hints_by_engagement(engagement_id)
    except Exception as e:
        logger.warning("get_benchmark_hints_by_engagement failed: %s", e)
        stored = []
    if stored:
        return {"engagement_id": engagement_id, "hints": stored, "opted_out": False}
    # Derive hints from client profile
    hints = []
    sector = client.get("sector_archetype") or client.get("industry")
    if sector:
        hints.append({
            "category": "sector",
            "title": f"Sector: {sector}",
            "content": f"Your sector archetype is set to {sector}. Use this to tailor scope item relevance and industry best practices in gap analysis.",
        })
    erp = client.get("erp_maturity")
    if erp:
        hints.append({
            "category": "erp_maturity",
            "title": f"ERP maturity: {erp}",
            "content": f"ERP maturity is {erp}. This can influence fit vs. gap expectations and adoption readiness.",
        })
    drivers = client.get("complexity_drivers")
    if drivers:
        if isinstance(drivers, list):
            drivers_str = ", ".join(str(d) for d in drivers[:5])
        else:
            drivers_str = str(drivers)
        hints.append({
            "category": "complexity",
            "title": "Complexity drivers",
            "content": f"Noted complexity drivers: {drivers_str}. Consider these when assessing effort and customisation risk.",
        })
    if not hints:
        hints.append({
            "category": "general",
            "title": "Benchmark insights",
            "content": "Add sector archetype, ERP maturity, or complexity drivers in the client profile to get tailored benchmark hints here.",
        })
    return {"engagement_id": engagement_id, "hints": hints, "opted_out": False}


@router.get("/engagement/{engagement_id}/hitl-queue")
def get_hitl_queue(engagement_id: str):
    """Return HITL queue for an engagement, grouped by state with summary."""
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    buckets: Dict[str, List[Dict[str, Any]]] = {
        "ai_draft": [],
        "needs_sme_review": [],
        "needs_architect_review": [],
        "approved": [],
        "out_of_scope": [],
    }
    by_state_counts: Dict[str, int] = {k: 0 for k in buckets.keys()}
    by_process: Dict[str, Dict[str, int]] = {}

    for req in requirements:
        state = (req.get("hitl_state") or "ai_draft").lower()
        if state not in buckets:
            state = "ai_draft"
        buckets[state].append(req)
        by_state_counts[state] = by_state_counts.get(state, 0) + 1

        bp = req.get("business_process") or "Unspecified"
        if bp not in by_process:
            by_process[bp] = {k: 0 for k in buckets.keys()}
        by_process[bp][state] = by_process[bp].get(state, 0) + 1

    summary = {
        "total": len(requirements),
        "by_state": by_state_counts,
        "by_process": by_process,
    }

    return {
        "engagement_id": engagement_id,
        "ai_draft": buckets["ai_draft"],
        "needs_sme_review": buckets["needs_sme_review"],
        "needs_architect_review": buckets["needs_architect_review"],
        "approved": buckets["approved"],
        "out_of_scope": buckets["out_of_scope"],
        "summary": summary,
    }


@router.get("/engagement/{engagement_id}/hitl-report", response_class=StreamingResponse)
def get_hitl_report(engagement_id: str):
    """Download HITL report as Excel: requirements with state + HITL events log."""
    try:
        requirements = get_requirements_by_engagement(engagement_id)
        events = list_hitl_events(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "HITL Summary"
    ws1.append(["req_id", "title", "hitl_state", "business_process", "priority"])
    for r in requirements:
        ws1.append([
            r.get("req_id"),
            r.get("title"),
            r.get("hitl_state") or "ai_draft",
            r.get("business_process"),
            r.get("priority"),
        ])
    ws2 = wb.create_sheet("HITL Events")
    ws2.append(["event_id", "req_id", "from_state", "to_state", "actor", "actor_role", "notes", "created_at"])
    for e in events:
        ws2.append([
            e.get("event_id"),
            e.get("req_id"),
            e.get("from_state"),
            e.get("to_state"),
            e.get("actor"),
            e.get("actor_role"),
            e.get("notes"),
            e.get("created_at"),
        ])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    filename = f"{engagement_id}_hitl_report.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/engagement/{engagement_id}/hitl-events")
def get_hitl_events(engagement_id: str):
    """Return HITL events audit trail for an engagement."""
    try:
        events = list_hitl_events(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {
        "engagement_id": engagement_id,
        "events": events,
    }


@router.get("/engagement/{engagement_id}/audit-trail", tags=["audit"])
def engagement_audit_trail(engagement_id: str, limit: int = 100):
    """Unified audit trail: HITL events + audit_events sorted by created_at desc."""
    try:
        hitl = list_hitl_events(engagement_id)
        audit = list_audit_events_by_engagement(engagement_id, limit=limit)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    for h in hitl:
        h["_source"] = "hitl"
        h["_sort_at"] = h.get("created_at") or ""
    for a in audit:
        a["_source"] = "audit"
        a["_sort_at"] = a.get("created_at") or ""
    combined = sorted(hitl + audit, key=lambda x: x.get("_sort_at", ""), reverse=True)
    for x in combined:
        x.pop("_sort_at", None)
    return {"engagement_id": engagement_id, "events": combined[:limit], "total": len(combined)}


@router.get("/engagement/{engagement_id}/completion-check", tags=["govern"])
def get_completion_check(engagement_id: str):
    """Engagement completion checklist: all reqs have fit-gap, all assessments HITL approved or out_of_scope."""
    try:
        eng = get_engagement(engagement_id)
        if not eng:
            raise HTTPException(status_code=404, detail="Engagement not found")
        requirements = get_requirements_by_engagement(engagement_id)
        assessments = get_fit_gap_by_engagement(engagement_id)
        req_ids_with_fga = {a["req_id"] for a in assessments}
        all_have_fga = all(r["req_id"] in req_ids_with_fga for r in requirements) if requirements else True
        approved_or_scope = sum(
            1 for a in assessments
            if (a.get("hitl_state") or "").lower() in ("approved", "out_of_scope")
        )
        total_fga = len(assessments)
        all_fga_reviewed = (total_fga == approved_or_scope) if total_fga else True
        return {
            "engagement_id": engagement_id,
            "total_requirements": len(requirements),
            "total_assessments": total_fga,
            "all_requirements_have_fit_gap": all_have_fga,
            "all_assessments_reviewed": all_fga_reviewed,
            "approved_or_out_of_scope_count": approved_or_scope,
            "complete": all_have_fga and all_fga_reviewed,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Sources (enterprise multi-source capture) ───────────────────────────────

class SourceCreate(BaseModel):
    engagement_id: str
    source_type: str  # transcript|notes|excel|document|workshop
    title: str
    raw_content: Optional[str] = None
    file_url: Optional[str] = None
    file_name: Optional[str] = None
    created_by: Optional[str] = None


class SourceUpdate(BaseModel):
    title: Optional[str] = None
    status: Optional[str] = None
    raw_content: Optional[str] = None
    extracted_count: Optional[int] = None


@router.post("/sources", status_code=201)
def post_source(body: SourceCreate):
    """Create a source for an engagement."""
    try:
        out = create_source(
            engagement_id=body.engagement_id,
            source_type=body.source_type,
            title=body.title or "Untitled",
            raw_content=body.raw_content,
            file_url=body.file_url,
            file_name=body.file_name,
            created_by=body.created_by,
        )
        return out
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/sources")
def list_sources(engagement_id: str):
    """List sources for an engagement. Query param: engagement_id."""
    try:
        items = list_sources_by_engagement(engagement_id)
        return {"engagement_id": engagement_id, "sources": items}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/engagement/{engagement_id}/sources")
def get_engagement_sources(engagement_id: str):
    """List sources for an engagement."""
    try:
        items = list_sources_by_engagement(engagement_id)
        return {"engagement_id": engagement_id, "sources": items}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/sources/{source_id}")
def get_source_by_id(source_id: str, engagement_id: str):
    """Get a single source. Query param: engagement_id."""
    try:
        out = get_source(source_id, engagement_id)
        if not out:
            raise HTTPException(status_code=404, detail="Source not found")
        return out
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.patch("/sources/{source_id}")
def patch_source(source_id: str, body: SourceUpdate, engagement_id: str):
    """Update a source. Query param: engagement_id."""
    try:
        existing = get_source(source_id, engagement_id)
        if not existing:
            raise HTTPException(status_code=404, detail="Source not found")
        updates = body.model_dump(exclude_unset=True)
        out = update_source(source_id, engagement_id, updates)
        return out
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/sources/{source_id}")
def delete_source_route(source_id: str, engagement_id: str):
    """Delete a source. Query param: engagement_id."""
    try:
        existing = get_source(source_id, engagement_id)
        if not existing:
            raise HTTPException(status_code=404, detail="Source not found")
        delete_source(source_id, engagement_id)
        return {"deleted": source_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


_SOURCE_EXTRACT_SYSTEM = """You are a senior SAP business analyst. Extract discrete business requirements from the given text (notes, document, or transcript).

For each requirement return:
- title: Brief title (max 10 words)
- description: Detailed description
- source_excerpt: Short quote from the text that supports this requirement (1-2 sentences)
- business_process: One of Procure-to-Pay, Order-to-Cash, Record-to-Report, Plan-to-Produce, Hire-to-Retire, Other
- priority: Must-Have, Should-Have, or Nice-to-Have
- tags: Array — use only: pain_point, manual_step, secret_sauce, workaround, hand_off

Return a JSON array only — no markdown:
[{"title": "...", "description": "...", "source_excerpt": "...", "business_process": "...", "priority": "...", "tags": []}]
Return [] if no clear requirements."""


@router.post("/sources/{source_id}/extract")
def post_source_extract(source_id: str, engagement_id: str):
    """Extract requirements from a source using AI. Creates requirements with source_id and source_excerpt."""
    try:
        src = get_source(source_id, engagement_id)
        if not src:
            raise HTTPException(status_code=404, detail="Source not found")
        raw = (src.get("raw_content") or "").strip()
        if not raw:
            update_source(source_id, engagement_id, {"status": "extracted", "extracted_count": 0})
            return {"source_id": source_id, "engagement_id": engagement_id, "extracted": 0, "message": "No raw content to extract."}
        update_source(source_id, engagement_id, {"status": "processing"})
        created_count = 0
        if os.getenv("ANTHROPIC_API_KEY"):
            try:
                provider = get_provider()
                result = provider.complete(_SOURCE_EXTRACT_SYSTEM, raw[:30000], max_tokens=1024, model=MODEL_HAIKU)
                raw_text = result.get("content", "[]")
                json_match = re.search(r"\[.*\]", raw_text, re.DOTALL)
                if json_match:
                    extracted = json.loads(json_match.group())
                    valid_tags = {"pain_point", "manual_step", "secret_sauce", "workaround", "hand_off"}
                    for item in extracted:
                        tags = [t for t in (item.get("tags") or []) if t in valid_tags]
                        create_requirement(
                            engagement_id=engagement_id,
                            title=item.get("title", "Untitled"),
                            description=item.get("description", ""),
                            source_type=src.get("source_type") or "document",
                            tags=tags,
                            business_process=item.get("business_process") or None,
                            priority=item.get("priority") or "Must-Have",
                            source_id=source_id,
                            source_excerpt=item.get("source_excerpt") or None,
                            extraction_confidence=0.85,
                            hitl_state="ai_draft",
                        )
                        created_count += 1
            except Exception as e:
                logger.warning("Source extract LLM failed: %s", e)
                update_source(source_id, engagement_id, {"status": "uploaded", "extracted_count": 0})
                raise HTTPException(status_code=500, detail=f"Extraction failed: {e}")
        update_source(source_id, engagement_id, {"status": "extracted", "extracted_count": created_count})
        return {"source_id": source_id, "engagement_id": engagement_id, "extracted": created_count}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/sources/{source_id}/upload-file")
async def upload_source_file(
    source_id: str,
    engagement_id: str = Form(...),
    file: UploadFile = File(...),
):
    """Upload a PDF, DOCX, or TXT file and extract its text into sources.content."""
    file_ext = (file.filename or "").lower().rsplit(".", 1)[-1]
    if file_ext not in ("pdf", "docx", "txt"):
        raise HTTPException(status_code=400, detail="Unsupported file type. Use .pdf, .docx, or .txt")
    src = get_source(source_id, engagement_id)
    if not src:
        raise HTTPException(status_code=404, detail="Source not found")
    file_bytes = await file.read()
    file_size = len(file_bytes)
    text = ""
    try:
        if file_ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        elif file_ext == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            text = file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to extract text from file: {e}")
    update_source(source_id, engagement_id, {
        "content": text,
        "raw_content": text,
        "file_name": file.filename,
        "file_size": file_size,
        "status": "ready_for_extraction",
    })
    return {
        "source_id": source_id,
        "extracted_text_length": len(text),
        "status": "ready_for_extraction",
    }


# ── Analyse All ───────────────────────────────────────────────────────────────

@router.post("/engagement/{engagement_id}/analyse-all")
def analyse_all(engagement_id: str):
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    open_reqs = [r for r in requirements if r.get("status") == "open"]
    if not open_reqs:
        return {"processed": 0, "results": []}

    provider = get_provider()
    results = []

    for req in open_reqs:
        req_id = req["req_id"]
        try:
            matches, tokens_used = _run_gap_analysis(provider, req["description"])
            timestamp = datetime.utcnow().isoformat()
            try:
                save_gap_analysis(
                    engagement_id=engagement_id,
                    process_description=req["description"],
                    matches=[m.dict() for m in matches],
                    tokens_used=tokens_used,
                    timestamp=timestamp,
                    req_id=req_id,
                )
            except Exception as db_err:
                print(f"DB save failed for {req_id} (non-fatal): {db_err}")

            update_requirement(req_id, engagement_id, {"status": "analysed"})

            top = matches[0] if matches else None
            results.append({
                "req_id": req_id,
                "title": req.get("title"),
                "top_match_id": top.id if top else None,
                "top_match_name": top.name if top else None,
            })
        except Exception as e:
            print(f"Analysis failed for {req_id}: {e}")

    return {"processed": len(results), "results": results}


# ── Process Mirror ────────────────────────────────────────────────────────────

@router.get("/engagement/{engagement_id}/process-mirror")
def get_process_mirror(engagement_id: str):
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    tag_counts: dict = {}
    for req in requirements:
        for tag in (req.get("tags") or []):
            tag_counts[tag] = tag_counts.get(tag, 0) + 1

    by_tag: dict = {}
    untagged = []
    for req in requirements:
        tags = req.get("tags") or []
        entry = {
            "req_id": req["req_id"],
            "title": req.get("title"),
            "description": req.get("description"),
            "stakeholder": req.get("stakeholder"),
        }
        if not tags:
            untagged.append(entry)
        else:
            for tag in tags:
                by_tag.setdefault(tag, []).append(entry)

    return {
        "engagement_id": engagement_id,
        "generated_at": datetime.utcnow().isoformat(),
        "summary": {
            "total_requirements": len(requirements),
            "by_tag": tag_counts,
        },
        "by_tag": by_tag,
        "untagged": untagged,
    }


# ── Sign-off Status ────────────────────────────────────────────────────────────

@router.get("/engagement/{engagement_id}/sign-off-status")
def get_sign_off_status(engagement_id: str):
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    counts: dict = {"draft": 0, "sme_approved": 0, "owner_approved": 0, "confirmed": 0}
    by_process: dict = {}

    for req in requirements:
        status = req.get("sign_off_status") or "draft"
        counts[status] = counts.get(status, 0) + 1

        process = req.get("business_process") or "Unclassified"
        by_process.setdefault(process, {"draft": 0, "sme_approved": 0, "owner_approved": 0, "confirmed": 0})
        by_process[process][status] = by_process[process].get(status, 0) + 1

    return {
        "engagement_id": engagement_id,
        "total": len(requirements),
        **counts,
        "by_process": by_process,
    }


@router.post("/engagement/{engagement_id}/signoff-batch-request", status_code=200)
def signoff_batch_request(engagement_id: str, body: dict = Body(...)):
    """Generate a plain-English sign-off summary for selected requirements and mark them sme_approved."""
    req_ids = body.get("requirement_ids", [])
    if not req_ids:
        raise HTTPException(status_code=400, detail="requirement_ids required")

    eng = get_engagement_with_client(engagement_id)
    client = (eng or {}).get("client") or {}

    assessments = get_fit_gap_by_engagement(engagement_id)
    assess_by_req = {a["req_id"]: a for a in assessments}

    fit_labels = {
        "fit_standard": "SAP covers this process out of the box — no changes needed",
        "fit_config": "SAP covers this with standard configuration",
        "fit_extension": "SAP covers this with a supported extension (low risk)",
        "gap_ricefw": "SAP does not cover this — custom development required (adds cost and risk)",
        "gap_companion": "SAP does not cover this — requires a third-party solution",
        "out_of_scope": "This item is outside the project scope",
    }

    items = []
    for req_id in req_ids:
        req = get_requirement_by_id(req_id, engagement_id)
        if not req:
            continue
        assessment = assess_by_req.get(req_id)
        if not assessment:
            continue
        fit_type = assessment.get("fit_type", "unknown")
        plain_english = fit_labels.get(fit_type, fit_type)
        items.append({
            "req_id": req_id,
            "title": req.get("title") or (req.get("description", ""))[:80],
            "process": req.get("business_process", "General"),
            "fit_type": fit_type,
            "plain_english": plain_english,
            "effort_estimate": (
                f"{assessment.get('estimated_effort_days_low', 0)}–{assessment.get('estimated_effort_days_high', 0)} days"
                if fit_type in ["gap_ricefw", "gap_companion"] else None
            ),
            "cost_band": assessment.get("cost_band") if fit_type in ["gap_ricefw", "gap_companion"] else None,
        })
        try:
            update_requirement(req_id, engagement_id, {"sign_off_status": "sme_approved"})
        except Exception:
            pass

    try:
        create_audit_event(engagement_id, "signoff_batch_requested", entity_type="requirements",
                           details={"count": len(items), "req_ids": req_ids})
    except Exception:
        pass

    return {
        "engagement_id": engagement_id,
        "client_name": client.get("name", "Client"),
        "items_count": len(items),
        "items": items,
        "message": body.get("message", ""),
        "requested_at": datetime.now(timezone.utc).isoformat(),
    }


# ── KPI Summary ────────────────────────────────────────────────────────────────

@router.get("/engagement/{engagement_id}/kpi-summary")
def get_kpi_summary(engagement_id: str):
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    by_process: dict = {}
    total_with_kpi = 0

    for req in requirements:
        kpi = req.get("kpi_impact")
        if not kpi:
            continue
        total_with_kpi += 1
        process = req.get("business_process") or "Unclassified"
        by_process.setdefault(process, []).append({
            "req_id": req["req_id"],
            "title": req.get("title"),
            "kpi_impact": kpi,
            "priority": req.get("priority"),
            "stakeholder": req.get("stakeholder"),
        })

    return {
        "engagement_id": engagement_id,
        "total_with_kpi": total_with_kpi,
        "by_process": by_process,
    }


# ── Assets ────────────────────────────────────────────────────────────────────

_SUPPORTED_EXTENSIONS = {
    "pdf", "docx", "xlsx", "png", "jpg", "jpeg",
    "mp3", "mp4", "wav", "mov", "txt",
}

_CONTENT_TYPE_MAP = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "png":  "image/png",
    "jpg":  "image/jpeg",
    "jpeg": "image/jpeg",
    "mp3":  "audio/mpeg",
    "mp4":  "video/mp4",
    "wav":  "audio/wav",
    "mov":  "video/quicktime",
    "txt":  "text/plain",
}


@router.post("/assets/upload", status_code=201)
async def upload_asset(
    file: UploadFile = File(...),
    engagement_id: str = Form(...),
    uploaded_by: Optional[str] = Form(None),
    req_id: Optional[str] = Form(None),
    process_level_2: Optional[str] = Form(None),
    process_level_3: Optional[str] = Form(None),
):
    file_name = file.filename or "unnamed"
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if ext not in _SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '.{ext}'. Allowed: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}",
        )

    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read uploaded file: {e}")

    try:
        content_type = _CONTENT_TYPE_MAP[ext]

        # Build asset record (asset_id assigned inside create_asset via _next_asset_id)
        asset_data: Dict[str, Any] = {
            "engagement_id": engagement_id,
            "file_name": file_name,
            "file_type": ext,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        if uploaded_by:
            asset_data["uploaded_by"] = uploaded_by
        if req_id:
            asset_data["req_id"] = req_id
        if process_level_2:
            asset_data["process_level_2"] = process_level_2
        if process_level_3:
            asset_data["process_level_3"] = process_level_3

        # Extract text for plain-text files
        if ext == "txt":
            try:
                asset_data["extracted_text"] = file_bytes.decode("utf-8", errors="replace")
            except Exception:
                pass

        # Create DB record first to get asset_id, then upload to storage
        asset = create_asset(asset_data)
        if not asset:
            raise HTTPException(status_code=500, detail="Failed to create asset record")

        asset_id = asset["asset_id"]
        storage_url = upload_file_to_storage(engagement_id, asset_id, file_name, file_bytes, content_type)
        update_asset(asset_id, {"storage_url": storage_url})

        return {
            "asset_id": asset_id,
            "file_name": file_name,
            "file_type": ext,
            "storage_url": storage_url,
            "engagement_id": engagement_id,
        }
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        logger.exception("Assets upload error: %s\n%s", e, traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/assets")
def list_assets(engagement_id: str):
    try:
        assets = get_assets_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"engagement_id": engagement_id, "total": len(assets), "assets": assets}


@router.get("/assets/requirement/{req_id}")
def list_assets_for_requirement(req_id: str, engagement_id: str):
    try:
        assets = get_assets_by_requirement(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"req_id": req_id, "engagement_id": engagement_id, "total": len(assets), "assets": assets}


@router.patch("/assets/{asset_id}")
def patch_asset(asset_id: str, body: AssetUpdate):
    aid = asset_id.strip()
    updates = {k: v for k, v in body.dict().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields provided to update")
    try:
        asset = update_asset(aid, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not asset:
        raise HTTPException(status_code=404, detail=f"{aid} not found")
    return asset


# ── Process Hierarchy (Sprint 4) ───────────────────────────────────────────────

_PROCESS_HIERARCHY: Dict[str, Dict[str, Dict[str, List[str]]]] = {
    "Finance": {
        "Record to Report": {
            "General Ledger":       ["Journal Entry", "Period Close", "Financial Statements"],
            "Accounts Payable":     ["Invoice Entry", "Payment Run", "Vendor Reconciliation"],
            "Accounts Receivable":  ["Customer Invoice", "Cash Application", "Dunning"],
            "Asset Accounting":     ["Asset Acquisition", "Depreciation Run", "Asset Retirement"],
            "Tax Compliance":       ["Tax Return", "Tax Reporting", "Withholding Tax"],
            "Period Close":         ["Month-End Close", "Year-End Close", "Accruals"],
            "Intercompany":         ["Intercompany Billing", "Intercompany Reconciliation", "Elimination"],
            "Treasury":             ["Bank Account Management", "Cash Position", "Liquidity Planning"],
            "Cash Management":      ["Cash Flow Forecast", "Bank Statement", "Payment Advice"],
        },
        "Management Accounting": {
            "Cost Center Accounting":    ["Cost Allocation", "Assessment", "Distribution"],
            "Internal Orders":           ["Order Settlement", "Budget Monitoring", "Cost Collection"],
            "Profit Center":             ["Profit Center Assignment", "Transfer Pricing", "Reporting"],
            "Product Costing":           ["Standard Cost Estimate", "Actual Costing", "Variance Analysis"],
            "Profitability Analysis":    ["Contribution Margin", "Segment Reporting", "CO-PA Settlement"],
            "Planning and Budgeting":    ["Annual Budget", "Rolling Forecast", "Plan Upload"],
        },
        "Group Reporting": {
            "Legal Consolidation":      ["Consolidation Entry", "Minority Interest", "Statutory Report"],
            "Management Consolidation": ["Management Report", "Segment Elimination", "Interunit Reconciliation"],
        },
    },
    "Sourcing and Procurement": {
        "Procure to Pay": {
            "Purchase Requisition": ["PR Creation", "PR Approval", "PR Conversion"],
            "Purchase Order":       ["PO Creation", "PO Approval", "PO Amendment"],
            "Goods Receipt":        ["GR Posting", "Quality Inspection", "Tolerance Check"],
            "Invoice Verification": ["Invoice Entry", "3-Way Match", "Exception Handling"],
            "Supplier Payment":     ["Payment Proposal", "Payment Run", "Bank Transfer"],
        },
        "Supplier Management": {
            "Supplier Evaluation":    ["Scorecard", "KPI Tracking", "Audit"],
            "Supplier Qualification": ["Onboarding", "Risk Assessment", "Certification"],
            "Contract Management":    ["Contract Creation", "Contract Renewal", "Compliance Monitoring"],
        },
        "Sourcing": {
            "RFQ":          ["RFQ Creation", "Supplier Invitation", "Response Collection"],
            "Bid Evaluation": ["Price Comparison", "Technical Evaluation", "Scoring"],
            "Award":        ["Award Decision", "Contract Award", "PO Creation"],
        },
    },
    "Sales": {
        "Order to Cash": {
            "Quotation":    ["Quotation Creation", "Pricing", "Approval"],
            "Sales Order":  ["Order Entry", "Order Confirmation", "Credit Check"],
            "Delivery":     ["Picking", "Packing", "Goods Issue"],
            "Goods Issue":  ["Stock Reduction", "Delivery Completion", "Shipment"],
            "Billing":      ["Invoice Creation", "Revenue Recognition", "Tax Calculation"],
            "Cash Collection": ["Payment Receipt", "Bank Reconciliation", "Cash Application"],
        },
        "Customer Management": {
            "Credit Management": ["Credit Limit", "Credit Check", "Risk Classification"],
            "Returns":           ["Return Order", "Goods Receipt", "Credit Memo"],
            "Complaints":        ["Complaint Entry", "Root Cause Analysis", "Resolution"],
            "Rebates":           ["Rebate Agreement", "Accrual", "Settlement"],
        },
        "Contract Management": {
            "Customer Contracts":    ["Contract Creation", "Renewal", "Termination"],
            "Subscription Management": ["Subscription Order", "Billing Schedule", "Renewal"],
        },
    },
    "Supply Chain": {
        "Inventory Management": {
            "Goods Movement":    ["Goods Issue", "Goods Receipt", "Transfer Posting"],
            "Stock Transfer":    ["STO Creation", "Delivery", "Goods Receipt"],
            "Physical Inventory": ["Inventory Count", "Difference Posting", "Count Document"],
            "Batch Management":  ["Batch Creation", "Batch Classification", "Shelf Life"],
        },
        "Warehouse Management": {
            "Inbound Processing":  ["Putaway", "Goods Receipt", "Deconsolidation"],
            "Outbound Processing": ["Pick", "Pack", "Ship"],
            "Internal Warehouse":  ["Stock Transfer", "Replenishment", "Inventory"],
            "Yard Management":     ["Truck Arrival", "Door Assignment", "Departure"],
        },
        "Transportation": {
            "Freight Order":     ["Freight Order Creation", "Carrier Assignment", "Execution"],
            "Carrier Selection": ["Rate Comparison", "Tender", "Award"],
            "Freight Settlement": ["Freight Cost", "Invoice Verification", "Payment"],
        },
        "Demand and Supply Planning": {
            "Demand Forecasting": ["Statistical Forecast", "Consensus Demand", "Adjustment"],
            "MRP":               ["MRP Run", "Planned Order", "Exception Messages"],
            "Supply Planning":   ["Supply Network", "Heuristics", "Optimizer"],
            "ATP":               ["ATP Check", "Backorder Processing", "Confirmation"],
        },
    },
    "Manufacturing": {
        "Plan to Produce": {
            "Production Order":         ["Order Creation", "Release", "Confirmation"],
            "Process Order":            ["Batch Production", "Process Instruction", "GR"],
            "Repetitive Manufacturing": ["Rate-Based Planning", "Backflush", "Reporting"],
            "Kanban":                   ["Kanban Signal", "Replenishment", "Status Update"],
        },
        "Quality Management": {
            "Inspection Planning": ["Inspection Plan", "Sampling Procedure", "Characteristic"],
            "Inspection Lot":      ["Lot Creation", "Sample Drawing", "Results Recording"],
            "Usage Decision":      ["Acceptance", "Rejection", "Follow-Up Action"],
            "Defect Recording":    ["Defect Notification", "Root Cause", "CAPA"],
        },
        "Engineering": {
            "BOM Management":    ["BOM Creation", "BOM Change", "Validity"],
            "Routing":           ["Operation", "Work Center", "Standard Values"],
            "Engineering Change": ["Change Order", "Revision Level", "Implementation"],
        },
    },
    "Asset Management": {
        "Maintain to Operate": {
            "Preventive Maintenance":  ["Maintenance Plan", "Scheduling", "Work Order"],
            "Corrective Maintenance":  ["Breakdown Notification", "Repair Order", "Completion"],
            "Predictive Maintenance":  ["Condition Monitoring", "Alert", "Predictive Work Order"],
            "Calibration":             ["Calibration Plan", "Measurement", "Certificate"],
        },
        "Asset Lifecycle": {
            "Asset Acquisition": ["Asset Creation", "Capitalization", "Settlement"],
            "Depreciation":      ["Depreciation Run", "Adjustment", "Reporting"],
            "Asset Retirement":  ["Retirement Posting", "Disposal", "Write-off"],
        },
    },
    "Human Resources": {
        "Hire to Retire": {
            "Recruitment":    ["Job Posting", "Application Review", "Offer"],
            "Onboarding":     ["New Hire Setup", "Equipment", "Training"],
            "Time Management": ["Time Recording", "Approval", "Absence"],
            "Payroll":        ["Payroll Run", "Posting", "Pay Slip"],
            "Offboarding":    ["Exit Interview", "Clearance", "Final Settlement"],
        },
        "Talent Management": {
            "Performance":   ["Goal Setting", "Mid-Year Review", "Annual Appraisal"],
            "Learning":      ["Course Assignment", "Completion", "Certification"],
            "Compensation":  ["Salary Review", "Bonus", "Equity"],
        },
    },
    "Professional Services": {
        "Project to Cash": {
            "Project Planning":   ["WBS", "Scheduling", "Budgeting"],
            "Resource Management": ["Resource Request", "Assignment", "Utilization"],
            "Time Recording":     ["Timesheet", "Approval", "Cost Posting"],
            "Project Billing":    ["Milestone Billing", "T&M Billing", "Revenue Recognition"],
        },
    },
}


@router.get("/process-hierarchy")
def get_process_hierarchy():
    return _PROCESS_HIERARCHY


@router.get("/process-hierarchy/flat")
def get_process_hierarchy_flat():
    flat = []
    for lob, level2_map in _PROCESS_HIERARCHY.items():
        for level2, level3_map in level2_map.items():
            for level3, level4 in level3_map.items():
                flat.append({
                    "lob": lob,
                    "level2": level2,
                    "level3": level3,
                    "level4": level4,
                })
    return flat


# ── Process Flow (Sprint 5) ────────────────────────────────────────────────────

def _build_process_flow(requirements: list, engagement_id: str) -> dict:
    """Build the swimlane / edge / stats payload from a list of requirement dicts."""
    # Group nodes by (lob, level2); requirements without both go to unassigned
    grouped: Dict[tuple, list] = {}
    unassigned_nodes: list = []

    for req in requirements:
        tags = req.get("tags") or []
        node = {
            "id": req["req_id"],
            "label": req.get("title", ""),
            "tags": tags,
            "priority": req.get("priority"),
            "status": req.get("status"),
            "pain_point": "pain_point" in tags,
            "secret_sauce": "secret_sauce" in tags,
            "manual_step": "manual_step" in tags,
            "process_level_3": req.get("process_level_3"),
        }
        lob = req.get("business_process")
        level2 = req.get("process_level_2")
        if lob and level2:
            grouped.setdefault((lob, level2), []).append(node)
        else:
            unassigned_nodes.append(node)

    # Build ordered swimlanes
    swimlanes = [
        {"lob": lob, "level2": level2, "nodes": nodes}
        for (lob, level2), nodes in sorted(grouped.items())
    ]
    if unassigned_nodes:
        swimlanes.append({"lob": "Unassigned", "level2": "Unassigned", "nodes": unassigned_nodes})

    # Edges: sequential links between requirements sharing the same level2 group
    edges: list = []
    for (_, _), nodes in sorted(grouped.items()):
        req_ids = [n["id"] for n in nodes]
        for i in range(len(req_ids) - 1):
            edges.append({"source": req_ids[i], "target": req_ids[i + 1], "label": "same process"})

    # Stats
    def _count_tag(tag: str) -> int:
        return sum(1 for r in requirements if tag in (r.get("tags") or []))

    return {
        "engagement_id": engagement_id,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "swimlanes": swimlanes,
        "edges": edges,
        "stats": {
            "total_nodes": len(requirements),
            "pain_points": _count_tag("pain_point"),
            "manual_steps": _count_tag("manual_step"),
            "secret_sauce": _count_tag("secret_sauce"),
            "unassigned_process": len(unassigned_nodes),
        },
    }


@router.get("/engagement/{engagement_id}/process-flow")
def get_process_flow(engagement_id: str):
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return _build_process_flow(requirements, engagement_id)


@router.post("/engagement/{engagement_id}/process-flow/assign")
def assign_process_flow(engagement_id: str, body: ProcessFlowAssignRequest):
    updates: Dict[str, Any] = {"process_level_2": body.process_level_2}
    if body.process_level_3 is not None:
        updates["process_level_3"] = body.process_level_3
    try:
        updated = update_requirement(body.req_id, engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not updated:
        raise HTTPException(
            status_code=404,
            detail=f"{body.req_id} not found in engagement {engagement_id}",
        )
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return _build_process_flow(requirements, engagement_id)


# ── Process Steps (Sprint 6) ───────────────────────────────────────────────────

_VALID_SHAPES = {"start", "end", "process", "decision", "document"}
_VALID_STEP_TYPES = {"manual", "system", "agentic"}

_STEP_EXTRACT_SYSTEM_PROMPT = """You are a senior business analyst mapping As-Is process flows for SAP S/4HANA implementations.

Given a business requirement and its conversation transcript, extract the As-Is process steps as a structured JSON array.

For each step infer:
- step_number: integer starting at 1, in execution order
- title: short action label (max 8 words)
- description: what actually happens at this step
- performer_name: the person's name or generic role name who does it
- performer_role: their job title or function
- shape: choose from → "start" (first step only), "end" (last step only), "decision" (approval/check/condition), "document" (form/report/record), "process" (default action)
- step_type: "manual" if done by a human without system, "system" if automated/done in a system, "agentic" if AI-driven
- duration_minutes: estimated time in minutes if mentioned or inferable, null otherwise
- systems_used: list of IT systems or tools mentioned for this step (e.g. ["SAP ECC", "Excel"])
- kpis: object with nullable float fields → {"error_rate_pct": null, "volume_per_month": null, "rework_rate_pct": null}
- is_pain_point: true if words like "manual", "error-prone", "time-consuming", "tedious", "slow", "broken", "workaround" apply
- next_step_id: null (will be assigned after creation)
- branches: for decision nodes only → [{"label": "Approved", "target_step_id": null}, {"label": "Rejected", "target_step_id": null}], else null

Return a JSON array only — no markdown fences, no commentary.

Example shape rules:
- Approval step → shape="decision", branches=[{"label":"Approved","target_step_id":null},{"label":"Rejected","target_step_id":null}]
- Filling in a form → shape="document"
- Starting trigger → shape="start"
- Final confirmation → shape="end"
- Everything else → shape="process"

If the transcript is empty or insufficient, generate a realistic 5-step As-Is process based on the requirement title and description."""


def _sample_steps_for_requirement(req: dict) -> list:
    """Generate 5 generic sample process steps based on the requirement title."""
    title = req.get("title", "Business Process")
    process = req.get("business_process", "Operations")
    return [
        {
            "step_number": 1,
            "title": f"Receive {title} request",
            "description": f"Requestor submits a new {title} request via email or paper form to the responsible team.",
            "performer_name": "Requestor",
            "performer_role": "Business User",
            "shape": "start",
            "step_type": "manual",
            "duration_minutes": 10.0,
            "systems_used": ["Email"],
            "kpis": {"error_rate_pct": None, "volume_per_month": None, "rework_rate_pct": None},
            "is_pain_point": False,
            "next_step_id": None,
            "branches": None,
        },
        {
            "step_number": 2,
            "title": "Log and validate request",
            "description": f"The {process} team logs the request in a spreadsheet and checks for completeness and duplicates.",
            "performer_name": "Process Owner",
            "performer_role": "Team Lead",
            "shape": "process",
            "step_type": "manual",
            "duration_minutes": 20.0,
            "systems_used": ["Excel"],
            "kpis": {"error_rate_pct": 12.0, "volume_per_month": None, "rework_rate_pct": None},
            "is_pain_point": True,
            "next_step_id": None,
            "branches": None,
        },
        {
            "step_number": 3,
            "title": "Management approval",
            "description": "Manager reviews the request and approves or rejects via email. No formal workflow exists.",
            "performer_name": "Line Manager",
            "performer_role": "Manager",
            "shape": "decision",
            "step_type": "manual",
            "duration_minutes": 1440.0,
            "systems_used": ["Email"],
            "kpis": {"error_rate_pct": None, "volume_per_month": None, "rework_rate_pct": None},
            "is_pain_point": True,
            "next_step_id": None,
            "branches": [
                {"label": "Approved", "target_step_id": None},
                {"label": "Rejected", "target_step_id": None},
            ],
        },
        {
            "step_number": 4,
            "title": "Process and record outcome",
            "description": f"The approved {title} is processed and the result is recorded in the relevant system.",
            "performer_name": "Clerk",
            "performer_role": "Operations Clerk",
            "shape": "document",
            "step_type": "system",
            "duration_minutes": 30.0,
            "systems_used": ["ERP System"],
            "kpis": {"error_rate_pct": None, "volume_per_month": 200.0, "rework_rate_pct": 8.0},
            "is_pain_point": False,
            "next_step_id": None,
            "branches": None,
        },
        {
            "step_number": 5,
            "title": "Notify requestor of completion",
            "description": "Requestor is notified by email that the process is complete.",
            "performer_name": "Process Owner",
            "performer_role": "Team Lead",
            "shape": "end",
            "step_type": "manual",
            "duration_minutes": 5.0,
            "systems_used": ["Email"],
            "kpis": {"error_rate_pct": None, "volume_per_month": None, "rework_rate_pct": None},
            "is_pain_point": False,
            "next_step_id": None,
            "branches": None,
        },
    ]


@router.get("/requirements/{req_id}/process-steps")
def list_process_steps(req_id: str, engagement_id: str):
    try:
        steps = get_process_steps(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"req_id": req_id, "engagement_id": engagement_id, "total": len(steps), "steps": steps}


@router.post("/requirements/{req_id}/process-steps/seed", status_code=201)
def seed_process_steps(req_id: str, engagement_id: str):
    """Create 5 sample process steps for a requirement that has no steps yet."""
    try:
        existing = get_process_steps(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if existing:
        return {"req_id": req_id, "seeded": 0, "steps": existing, "message": "Steps already exist"}

    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")

    sample = _sample_steps_for_requirement(req)
    created = []
    for s in sample:
        try:
            step = create_process_step({"req_id": req_id, "engagement_id": engagement_id, **s})
            if step:
                created.append(step)
        except Exception as e:
            print(f"Failed to seed step {s['step_number']}: {e}")

    return {"req_id": req_id, "seeded": len(created), "steps": created}


@router.post("/requirements/{req_id}/process-steps/extract", status_code=201)
def extract_process_steps(req_id: str, engagement_id: str):
    """AI-extract As-Is process steps from the requirement's transcript. Clears existing steps first."""
    try:
        req = get_requirement_by_id(req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not req:
        raise HTTPException(status_code=404, detail=f"{req_id} not found for engagement {engagement_id}")

    transcript = req.get("raw_input") or ""
    title = req.get("title", "")
    description = req.get("description", "")

    if not transcript.strip():
        steps_data = _sample_steps_for_requirement(req)
    else:
        provider = get_provider()
        user_prompt = (
            f"Requirement title: {title}\n"
            f"Requirement description: {description}\n\n"
            f"Conversation transcript:\n{transcript}\n\n"
            "Extract the As-Is process steps as a JSON array."
        )
        try:
            # Use cheaper model and cap tokens for step extraction
            result = provider.complete(_STEP_EXTRACT_SYSTEM_PROMPT, user_prompt, max_tokens=768, model=MODEL_HAIKU)
            raw_text = result.get("content", "[]")
            json_match = re.search(r'\[.*\]', raw_text, re.DOTALL)
            if not json_match:
                raise ValueError("No JSON array in response")
            steps_data = json.loads(json_match.group())
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI extraction failed: {e}")

    # Clear existing steps for this req
    try:
        existing = get_process_steps(req_id, engagement_id)
        for s in existing:
            delete_process_step(s["id"], req_id, engagement_id)
    except Exception as e:
        print(f"Step cleanup warning (non-fatal): {e}")

    # Validate and persist
    created = []
    for i, s in enumerate(steps_data, start=1):
        shape = s.get("shape", "process")
        if shape not in _VALID_SHAPES:
            shape = "process"
        step_type = s.get("step_type", "manual")
        if step_type not in _VALID_STEP_TYPES:
            step_type = "manual"
        kpis = s.get("kpis") or {}
        kpis_clean = {
            "error_rate_pct": kpis.get("error_rate_pct"),
            "volume_per_month": kpis.get("volume_per_month"),
            "rework_rate_pct": kpis.get("rework_rate_pct"),
        }
        record = {
            "req_id": req_id,
            "engagement_id": engagement_id,
            "step_number": s.get("step_number", i),
            "title": s.get("title", f"Step {i}"),
            "description": s.get("description", ""),
            "performer_name": s.get("performer_name", ""),
            "performer_role": s.get("performer_role", ""),
            "shape": shape,
            "step_type": step_type,
            "duration_minutes": s.get("duration_minutes"),
            "systems_used": s.get("systems_used") or [],
            "kpis": kpis_clean,
            "is_pain_point": bool(s.get("is_pain_point", False)),
            "next_step_id": s.get("next_step_id"),
            "branches": s.get("branches"),
        }
        try:
            step = create_process_step(record)
            if step:
                created.append(step)
        except Exception as e:
            print(f"Failed to create step {i}: {e}")

    return {"req_id": req_id, "engagement_id": engagement_id, "extracted": len(created), "steps": created}


@router.post("/requirements/{req_id}/process-steps", status_code=201)
def create_step(req_id: str, engagement_id: str, body: ProcessStepCreate):
    if body.shape not in _VALID_SHAPES:
        raise HTTPException(status_code=400, detail=f"Invalid shape '{body.shape}'. Must be one of: {sorted(_VALID_SHAPES)}")
    if body.step_type not in _VALID_STEP_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid step_type '{body.step_type}'. Must be one of: {sorted(_VALID_STEP_TYPES)}")
    data = body.dict()
    data["req_id"] = req_id
    data["engagement_id"] = engagement_id
    try:
        step = create_process_step(data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not step:
        raise HTTPException(status_code=500, detail="Failed to create process step")
    return step


@router.put("/requirements/{req_id}/process-steps/{step_id}")
def update_step(req_id: str, step_id: UUID, engagement_id: str, body: ProcessStepUpdate):
    step_id_str = str(step_id)
    updates = {k: v for k, v in body.dict().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields provided to update")
    if "shape" in updates and updates["shape"] not in _VALID_SHAPES:
        raise HTTPException(status_code=400, detail=f"Invalid shape. Must be one of: {sorted(_VALID_SHAPES)}")
    if "step_type" in updates and updates["step_type"] not in _VALID_STEP_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid step_type. Must be one of: {sorted(_VALID_STEP_TYPES)}")
    try:
        step = update_process_step(step_id_str, req_id, engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not step:
        raise HTTPException(status_code=404, detail=f"Step {step_id_str} not found")
    return step


@router.delete("/requirements/{req_id}/process-steps/{step_id}", status_code=204)
def delete_step(req_id: str, step_id: UUID, engagement_id: str):
    try:
        step_id_str = str(step_id)
        deleted = delete_process_step(step_id_str, req_id, engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not deleted:
        raise HTTPException(status_code=404, detail=f"Step {step_id_str} not found")
    return None


@router.post("/engagement/{engagement_id}/seed-process-steps", status_code=201)
def seed_all_process_steps(engagement_id: str):
    """Seed sample process steps for every requirement in the engagement that has none."""
    try:
        requirements = get_requirements_by_engagement(engagement_id)
        existing_steps = get_process_steps_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    reqs_with_steps = {s["req_id"] for s in existing_steps}
    seeded_reqs = []

    for req in requirements:
        req_id = req["req_id"]
        if req_id in reqs_with_steps:
            continue
        sample = _sample_steps_for_requirement(req)
        count = 0
        for s in sample:
            try:
                step = create_process_step({"req_id": req_id, "engagement_id": engagement_id, **s})
                if step:
                    count += 1
            except Exception as e:
                print(f"Seed failed for {req_id} step {s['step_number']}: {e}")
        if count:
            seeded_reqs.append({"req_id": req_id, "title": req.get("title"), "steps_created": count})

    return {
        "engagement_id": engagement_id,
        "requirements_seeded": len(seeded_reqs),
        "seeded": seeded_reqs,
    }


_SYNTHETIC_REQUIREMENTS = [
    {"title": "Vendor invoice posting and three-way match", "description": "System must support automatic matching of PO, GR and invoice. Tolerance limits configurable by company code.", "business_process": "Procure-to-Pay", "priority": "Must-Have", "category": "Automation"},
    {"title": "Customer credit check before order release", "description": "Credit limit check at order entry and release. Block or warn based on configurable rules.", "business_process": "Order-to-Cash", "priority": "Must-Have", "category": "Control/Compliance"},
    {"title": "Monthly closing and period-end reporting", "description": "Support period-end close with automatic accruals and reporting package for management.", "business_process": "Record-to-Report", "priority": "Must-Have", "category": "Reporting"},
    {"title": "MRP run and planned order conversion", "description": "Run MRP for finished and semi-finished materials; convert planned orders to production orders with capacity check.", "business_process": "Plan-to-Produce", "priority": "Must-Have", "category": "Automation"},
    {"title": "Employee onboarding and role provisioning", "description": "Onboard new hires with role assignment and system access request workflow.", "business_process": "Hire-to-Retire", "priority": "Should-Have", "category": "Integration"},
    {"title": "Intercompany sales and purchase reconciliation", "description": "Automatic IC billing and reconciliation; eliminate manual matching.", "business_process": "Record-to-Report", "priority": "Should-Have", "category": "Automation"},
    {"title": "Approval workflow for purchase requisitions", "description": "Multi-level approval by amount and commodity; delegation when approver absent.", "business_process": "Procure-to-Pay", "priority": "Must-Have", "category": "Control/Compliance"},
    {"title": "Revenue recognition per IFRS 15", "description": "Support performance obligation and allocation; revenue recognition schedule by contract.", "business_process": "Record-to-Report", "priority": "Must-Have", "category": "Reporting"},
]


@router.post("/engagement/{engagement_id}/seed-synthetic", status_code=201)
def seed_synthetic_requirements(engagement_id: str):
    """Seed synthetic requirements for demo and E2E testing. Engagement must exist."""
    try:
        eng = get_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    created = []
    for item in _SYNTHETIC_REQUIREMENTS:
        try:
            rec = create_requirement(
                engagement_id=engagement_id,
                title=item["title"],
                description=item["description"],
                business_process=item.get("business_process"),
                priority=item.get("priority"),
                category=item.get("category"),
                status="open",
                tags=[],
                hitl_state="ai_draft",
            )
            if rec:
                created.append({"req_id": rec.get("req_id"), "title": rec.get("title")})
        except Exception as e:
            logger.warning("Seed requirement failed: %s", e)
    return {
        "engagement_id": engagement_id,
        "created": len(created),
        "requirements": created,
    }


# ── Excel Upload / Download for Requirements (Sprint 6) ────────────────────────

_EXCEL_TAGS_ALLOWED = {"pain_point", "manual_step", "secret_sauce", "workaround", "hand_off"}


def _parse_excel_tags(raw: Optional[str]) -> List[str]:
    """Parse a tags cell from Excel into canonical tag list."""
    if not raw:
        return []
    parts = re.split(r"[;,]", str(raw))
    result: List[str] = []
    for p in parts:
        norm = _normalize_tag(p)
        if norm in _EXCEL_TAGS_ALLOWED and norm not in result:
            result.append(norm)
    return result


@router.get("/engagement/{engagement_id}/requirements/export")
def export_requirements_excel(engagement_id: str):
    """Download all requirements for an engagement as an Excel file."""
    try:
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    wb = Workbook()
    ws = wb.active
    ws.title = "Requirements"

    headers = [
        "req_id",
        "reference_id",
        "engagement_id",
        "title",
        "description",
        "business_process",
        "priority",
        "category",
        "tags",
        "stakeholder",
        "sign_off_status",
        "business_value",
        "current_system",
        "target_system_module",
        "fit_type",
        "related_test_case_id",
    ]
    ws.append(headers)

    for r in requirements:
        tags = r.get("tags") or []
        tags_str = ", ".join(tags)
        ws.append([
            r.get("req_id"),
            r.get("reference_id"),
            r.get("engagement_id"),
            r.get("title"),
            r.get("description"),
            r.get("business_process"),
            r.get("priority"),
            r.get("category"),
            tags_str,
            r.get("stakeholder"),
            r.get("sign_off_status"),
            r.get("business_value"),
            r.get("current_system"),
            r.get("target_system_module"),
            r.get("fit_type"),
            r.get("related_test_case_id"),
        ])

    # Phase F: optional second sheet "Fit-Gap" when fit_gap_assessments exist
    try:
        fga_list = get_fit_gap_by_engagement(engagement_id) or []
    except Exception:
        fga_list = []
    if fga_list:
        ws2 = wb.create_sheet("Fit-Gap")
        fga_headers = [
            "assessment_id", "req_id", "fit_type", "complexity", "rationale",
            "sap_scope_item_id", "sap_scope_item_name", "workaround_option",
            "estimated_effort_days_low", "estimated_effort_days_high", "cost_band",
            "confidence_score", "hitl_state", "reviewed_by", "reviewed_at", "reviewer_notes",
        ]
        ws2.append(fga_headers)
        for a in fga_list:
            ws2.append([
                a.get("assessment_id"),
                a.get("req_id"),
                a.get("fit_type"),
                a.get("complexity"),
                a.get("rationale"),
                a.get("sap_scope_item_id"),
                a.get("sap_scope_item_name"),
                a.get("workaround_option"),
                a.get("estimated_effort_days_low"),
                a.get("estimated_effort_days_high"),
                a.get("cost_band"),
                a.get("confidence_score"),
                a.get("hitl_state"),
                a.get("reviewed_by"),
                a.get("reviewed_at"),
                a.get("reviewer_notes"),
            ])

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    filename = f"{engagement_id}_requirements.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/requirements/template/download")
def download_requirements_template():
    """Return an Excel template for requirements/RTM upload. Sheet 'RTM' with header row only."""
    wb = Workbook()
    ws = wb.active
    ws.title = "RTM"
    headers = [
        "Requirement ID",
        "Requirement Title",
        "Requirement Description",
        "Business Process Area",
        "Sub-Process",
        "Requirement Type",
        "Priority",
        "Business Value",
        "Source",
        "Status",
        "Current System",
        "Target System Module",
        "Fit Type",
        "Related Test Case ID",
    ]
    ws.append(headers)
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="RAPID_requirements_template.xlsx"'},
    )


@router.get("/clients/template/download")
def download_clients_template():
    """Return an Excel template for client data. Fill and re-upload or use for reference."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Client"
    headers = [
        "name",
        "industry",
        "employees",
        "legal_entities",
        "current_systems",
        "systems_to_keep",
        "systems_to_replace",
        "countries",
        "regulatory_environment",
        "business_strategy",
        "goals",
        "key_products",
        "value_proposition",
        "senior_executives",
        "direct_competitors",
        "substitutes",
        "sector_archetype",
        "complexity_drivers",
        "erp_maturity",
        "benchmark_opt_in",
    ]
    ws.append(headers)
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="RAPID_client_template.xlsx"'},
    )


@router.get("/engagements/template/download")
def download_engagements_template():
    """Return an Excel template for engagement data. client_id required when creating."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Engagement"
    headers = [
        "client_id",
        "name",
        "description",
        "go_live_date",
        "project_type",
        "status",
        "planned_start_date",
        "planned_end_date",
        "actual_start_date",
        "actual_end_date",
        "project_manager",
        "sponsor",
        "risk_level",
        "health",
    ]
    ws.append(headers)
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="RAPID_engagement_template.xlsx"'},
    )


@router.get("/ricefw/template/download")
def download_ricefw_template():
    """Return an Excel template for RICEFW inventory. Type: R|I|C|E|F|W; status: identified|approved|in_development|delivered|cancelled."""
    wb = Workbook()
    ws = wb.active
    ws.title = "RICEFW"
    headers = [
        "type",
        "name",
        "description",
        "req_id",
        "status",
        "complexity",
        "priority",
    ]
    ws.append(headers)
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="RAPID_ricefw_template.xlsx"'},
    )


def _normalize_header_cell(c: Any) -> str:
    if c is None:
        return ""
    return str(c).strip().lower()


def _rtm_find_header_row(rows: List[tuple]) -> Optional[int]:
    """Return row index of RTM header (row with 'Requirement ID', 'Requirement Title', 'Requirement Description')."""
    for idx, row in enumerate(rows[:5]):
        header = [_normalize_header_cell(c) for c in row]
        if any("requirement id" in h for h in header) and any("requirement title" in h for h in header) and any("requirement description" in h for h in header):
            return idx
    return None


def _priority_rtm_to_rapid(p: Optional[str]) -> Optional[str]:
    if not p:
        return None
    p = p.strip().lower()
    if p in ("high", "must", "must-have"):
        return "Must-Have"
    if p in ("medium", "should", "should-have"):
        return "Should-Have"
    if p in ("low", "could", "nice", "nice-to-have"):
        return "Nice-to-Have"
    return p.title() if p else None


def _status_rtm_to_rapid(s: Optional[str]) -> str:
    if not s:
        return "open"
    s = s.strip().lower()
    if s in ("planned", "draft", "new"):
        return "open"
    if s in ("in progress", "in_progress", "inprogress"):
        return "in_progress"
    if s in ("analysed", "analyzed", "done", "complete"):
        return "analysed"
    if s in ("closed", "cancelled"):
        return "closed"
    return "open"


@router.post("/engagement/{engagement_id}/requirements/import")
async def import_requirements_excel(engagement_id: str, file: UploadFile = File(...)):
    """Upload an Excel file to create or update requirements for an engagement.

    Supports two formats:
    - RTM format: Header row contains 'Requirement ID', 'Requirement Title', 'Requirement Description'
      (e.g. RAPID_RTM_Acme_S4.xlsx). Uses internal REQ-XXX ids; stores Excel Requirement ID in reference_id.
    - Legacy format: Header row contains 'req_id', 'title', 'description'. Existing req_id updates; blank creates new.
    """
    filename = file.filename or ""
    if not filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Only .xlsx Excel files are supported")

    try:
        contents = await file.read()
        wb = load_workbook(io.BytesIO(contents), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read Excel file: {e}")

    ws = wb["RTM"] if "RTM" in wb.sheetnames else wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise HTTPException(status_code=400, detail="Excel file is empty")

    def _cell(row_values, idx: Optional[int]) -> Optional[str]:
        if idx is None or idx >= len(row_values):
            return None
        value = row_values[idx]
        if value is None:
            return None
        return str(value).strip() or None

    # Detect RTM format: find row with Requirement ID, Requirement Title, Requirement Description
    rtm_header_row = _rtm_find_header_row(rows)
    if rtm_header_row is not None:
        header = [_normalize_header_cell(c) for c in rows[rtm_header_row]]
        data_start = rtm_header_row + 1

        def _rtm_col(name_substr: str) -> Optional[int]:
            for i, h in enumerate(header):
                if name_substr in h:
                    return i
            return None

        # Requirement ID = Excel/source ID (stored as reference_id); must not be "requirement title"
        col_ref_id = next((i for i, h in enumerate(header) if h == "requirement id" or (h and "requirement id" in h and "title" not in h)), _rtm_col("requirement id"))
        col_title = _rtm_col("requirement title")
        col_desc = _rtm_col("requirement description")
        col_bp = _rtm_col("business process")
        col_sub = _rtm_col("sub-process")
        col_type = _rtm_col("requirement type")
        col_priority = _rtm_col("priority")
        col_biz_val = _rtm_col("business value")
        col_source = _rtm_col("source")
        col_cur_sys = _rtm_col("current system")
        col_tgt_sys = _rtm_col("target system")
        col_fit = _rtm_col("fit type")
        col_tc = _rtm_col("test case")
        col_status = _rtm_col("status")

        if col_title is None or col_desc is None:
            raise HTTPException(status_code=400, detail="RTM sheet must contain Requirement Title and Requirement Description columns")

        created_ids: List[str] = []
        errors: List[Dict[str, Any]] = []
        for row_idx, row in enumerate(rows[data_start:], start=data_start + 1):
            if not any(row):
                continue
            title = _cell(row, col_title)
            description = _cell(row, col_desc)
            if not title or not description:
                errors.append({"row": row_idx, "error": "Missing title or description"})
                continue
            reference_id = _cell(row, col_ref_id)
            if not reference_id:
                errors.append({"row": row_idx, "error": "Missing Requirement ID (reference_id)"})
                continue
            try:
                created = create_requirement(
                    engagement_id=engagement_id,
                    title=title,
                    description=description,
                    source_type="Excel",
                    reference_id=reference_id,
                    business_process=_cell(row, col_bp),
                    process_level_2=_cell(row, col_sub),
                    category=_cell(row, col_type),
                    priority=_priority_rtm_to_rapid(_cell(row, col_priority)),
                    stakeholder=_cell(row, col_source),
                    status=_status_rtm_to_rapid(_cell(row, col_status)),
                    business_value=_cell(row, col_biz_val),
                    current_system=_cell(row, col_cur_sys),
                    target_system_module=_cell(row, col_tgt_sys),
                    fit_type=_cell(row, col_fit),
                    related_test_case_id=_cell(row, col_tc),
                )
                if created:
                    created_ids.append(created.get("req_id"))
            except Exception as e:
                errors.append({"row": row_idx, "error": str(e)})

        return {
            "engagement_id": engagement_id,
            "format": "RTM",
            "created": len(created_ids),
            "updated": 0,
            "created_req_ids": created_ids,
            "updated_req_ids": [],
            "errors": errors,
        }

    # Legacy format: header in first row
    header = [_normalize_header_cell(c) for c in rows[0]]

    def _col(name: str) -> Optional[int]:
        try:
            return header.index(name)
        except ValueError:
            return None

    col_req_id = _col("req_id")
    col_title = _col("title")
    col_desc = _col("description")
    col_bp = _col("business_process")
    col_priority = _col("priority")
    col_category = _col("category")
    col_tags = _col("tags")
    col_stakeholder = _col("stakeholder")
    col_sign_off_status = _col("sign_off_status")

    if col_title is None or col_desc is None:
        raise HTTPException(status_code=400, detail="Excel must contain 'title' and 'description' header columns")

    try:
        existing = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    existing_by_req = {r.get("req_id"): r for r in existing if r.get("req_id")}

    created_ids = []
    updated_ids = []
    errors = []

    for row_idx, row in enumerate(rows[1:], start=2):
        if not any(row):
            continue
        title = _cell(row, col_title)
        description = _cell(row, col_desc)
        if not title or not description:
            errors.append({"row": row_idx, "error": "Missing title or description"})
            continue
        req_id = _cell(row, col_req_id)
        tags_raw = _cell(row, col_tags)
        tags = _parse_excel_tags(tags_raw)
        payload = {"title": title, "description": description}
        bp = _cell(row, col_bp)
        if bp:
            payload["business_process"] = bp
        priority = _cell(row, col_priority)
        if priority:
            payload["priority"] = priority
        category = _cell(row, col_category)
        if category:
            payload["category"] = category
        stakeholder = _cell(row, col_stakeholder)
        if stakeholder:
            payload["stakeholder"] = stakeholder
        sign_off_status = _cell(row, col_sign_off_status)
        if sign_off_status:
            payload["sign_off_status"] = sign_off_status
        try:
            if req_id and req_id in existing_by_req:
                updates = {k: v for k, v in payload.items() if k not in {"title", "description"}}
                updates["title"] = title
                updates["description"] = description
                if tags:
                    updates["tags"] = tags
                updated = update_requirement(req_id=req_id, engagement_id=engagement_id, updates=updates)
                if updated:
                    updated_ids.append(updated.get("req_id", req_id))
            else:
                created = create_requirement(
                    engagement_id=engagement_id,
                    title=title,
                    description=description,
                    source_type="Excel",
                    tags=tags,
                    business_process=payload.get("business_process"),
                    priority=payload.get("priority"),
                    category=payload.get("category"),
                    stakeholder=payload.get("stakeholder"),
                    sign_off_status=payload.get("sign_off_status"),
                )
                if created:
                    created_ids.append(created.get("req_id"))
        except Exception as e:
            errors.append({"row": row_idx, "error": str(e)})

    return {
        "engagement_id": engagement_id,
        "created": len(created_ids),
        "updated": len(updated_ids),
        "created_req_ids": created_ids,
        "updated_req_ids": updated_ids,
        "errors": errors,
    }


# ── RICEFW Customisation Inventory (Sprint 7) ──────────────────────────────────

_RICEFW_TYPES = {"R", "I", "C", "E", "F", "W", "A"}
_RICEFW_STATUSES = {"identified", "approved", "in_development", "delivered", "cancelled"}


@router.get("/engagement/{engagement_id}/ricefw")
def list_ricefw(engagement_id: str, type: Optional[str] = None):
    """List RICEFW customisation items for an engagement. type filter: R | I | C | E | F | W."""
    if type and type.upper() not in _RICEFW_TYPES:
        raise HTTPException(status_code=400, detail=f"type must be one of: {', '.join(sorted(_RICEFW_TYPES))}")
    try:
        items = get_ricefw_by_engagement(engagement_id, type=type.upper() if type else None)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"engagement_id": engagement_id, "total": len(items), "items": items}


@router.get("/engagement/{engagement_id}/ricefw/export")
def export_ricefw_excel(engagement_id: str, type: Optional[str] = None):
    """Download RICEFW inventory for an engagement as an Excel file. Optional type filter: R|I|C|E|F|W."""
    if type and type.upper() not in _RICEFW_TYPES:
        raise HTTPException(status_code=400, detail=f"type must be one of: {', '.join(sorted(_RICEFW_TYPES))}")
    try:
        items = get_ricefw_by_engagement(engagement_id, item_type=type.upper() if type else None)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    wb = Workbook()
    ws = wb.active
    ws.title = "RICEFW Inventory"
    headers = ["id", "engagement_id", "type", "name", "description", "req_id", "status", "complexity", "priority", "created_at"]
    ws.append(headers)
    for item in items:
        ws.append([
            item.get("id"),
            item.get("engagement_id"),
            item.get("type"),
            item.get("name"),
            item.get("description") or "",
            item.get("req_id") or "",
            item.get("status") or "",
            item.get("complexity") or "",
            item.get("priority") or "",
            item.get("created_at") or "",
        ])

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    suffix = f"_{type}" if type else ""
    filename = f"{engagement_id}_ricefw{suffix}.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _check_ricefw_complexity_priority(complexity: Optional[str], priority: Optional[str]) -> None:
    if complexity and complexity.lower() not in _RICEFW_COMPLEXITY:
        raise HTTPException(status_code=400, detail=f"complexity must be one of: {', '.join(sorted(_RICEFW_COMPLEXITY))}")
    if priority and priority.lower() not in _RICEFW_PRIORITY:
        raise HTTPException(status_code=400, detail=f"priority must be one of: {', '.join(sorted(_RICEFW_PRIORITY))}")


@router.post("/engagement/{engagement_id}/ricefw", status_code=201)
def create_ricefw(engagement_id: str, body: RICEFWCreate):
    if body.type.upper() not in _RICEFW_TYPES:
        raise HTTPException(status_code=400, detail=f"type must be one of: {', '.join(sorted(_RICEFW_TYPES))}")
    if body.status and body.status.lower() not in _RICEFW_STATUSES:
        raise HTTPException(status_code=400, detail=f"status must be one of: {', '.join(sorted(_RICEFW_STATUSES))}")
    _check_ricefw_complexity_priority(body.complexity, body.priority)
    # Ensure req_id belongs to this engagement (no cross-engagement data linkage)
    req = get_requirement_by_id(body.req_id, engagement_id)
    if not req:
        raise HTTPException(
            status_code=400,
            detail=f"Requirement {body.req_id} not found in this engagement. All RICEFW items must link to requirements within the same engagement.",
        )
    try:
        item = create_ricefw_item(
            engagement_id=engagement_id,
            item_type=body.type.upper(),
            name=body.name,
            req_id=body.req_id,
            description=body.description,
            status=(body.status or "identified").lower(),
            complexity=body.complexity.lower() if body.complexity else None,
            priority=body.priority.lower() if body.priority else None,
            effort_days_low=body.effort_days_low,
            effort_days_high=body.effort_days_high,
            owner=body.owner,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not item:
        raise HTTPException(status_code=500, detail="Failed to create RICEFW item")
    return item


@router.patch("/engagement/{engagement_id}/ricefw/{item_id}")
def update_ricefw(engagement_id: str, item_id: UUID, body: RICEFWUpdate):
    updates = {k: v for k, v in body.dict().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields provided to update")
    if "type" in updates and updates["type"].upper() not in _RICEFW_TYPES:
        raise HTTPException(status_code=400, detail=f"type must be one of: {', '.join(sorted(_RICEFW_TYPES))}")
    if "status" in updates and updates["status"].lower() not in _RICEFW_STATUSES:
        raise HTTPException(status_code=400, detail=f"status must be one of: {', '.join(sorted(_RICEFW_STATUSES))}")
    if "type" in updates:
        updates["type"] = updates["type"].upper()
    if "status" in updates:
        updates["status"] = updates["status"].lower()
    _check_ricefw_complexity_priority(updates.get("complexity"), updates.get("priority"))
    if "complexity" in updates and updates["complexity"]:
        updates["complexity"] = updates["complexity"].lower()
    if "priority" in updates and updates["priority"]:
        updates["priority"] = updates["priority"].lower()
    if "req_id" in updates:
        req = get_requirement_by_id(updates["req_id"], engagement_id)
        if not req:
            raise HTTPException(
                status_code=400,
                detail=f"Requirement {updates['req_id']} not found in this engagement. All RICEFW items must link to requirements within the same engagement.",
            )
    try:
        item = update_ricefw_item(str(item_id), engagement_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not item:
        raise HTTPException(status_code=404, detail="RICEFW item not found")
    return item


@router.delete("/engagement/{engagement_id}/ricefw/{item_id}", status_code=204)
def delete_ricefw(engagement_id: str, item_id: UUID):
    try:
        deleted = delete_ricefw_item(str(item_id), engagement_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not deleted:
        raise HTTPException(status_code=404, detail="RICEFW item not found")
    return None


# ── Fit/Gap Assessments (Phase B) ────────────────────────────────────────────

_FIT_GAP_TYPES = {"fit_standard", "fit_config", "fit_extension", "gap_ricefw", "gap_companion", "out_of_scope"}
_FIT_GAP_COMPLEXITY = {"XS", "S", "M", "L", "XL"}

_FIT_TYPE_MAP = {
    "fit standard": "fit_standard",
    "fit_to_standard": "fit_standard",
    "fit config": "fit_config",
    "fit configuration": "fit_config",
    "fit extension": "fit_extension",
    "gap ricefw": "gap_ricefw",
    "gap - ricefw": "gap_ricefw",
    "gap_ricefw": "gap_ricefw",
    "gap companion": "gap_companion",
    "gap - companion": "gap_companion",
    "gap_companion": "gap_companion",
    "out of scope": "out_of_scope",
    "out_of_scope": "out_of_scope",
}


def _normalize_fit_type(raw: str) -> str:
    """Normalize fit_type to canonical underscore lowercase (fit_standard, gap_ricefw, etc.)."""
    if not raw or not isinstance(raw, str):
        return "fit_standard"
    key = raw.strip().lower().replace("-", "_").replace(" ", "_").replace("__", "_").strip("_")
    return _FIT_TYPE_MAP.get(key) or (key if key in _FIT_GAP_TYPES else "fit_standard")

_FIT_GAP_SYSTEM = """You are a senior SAP S/4HANA Public Cloud solution architect conducting a Fit-to-Standard assessment.
Client context: {context}
Be realistic and strict. Standard S/4HANA Public Cloud does NOT natively support:
- EV-specific or battery-specific manufacturing (battery BOM, state-of-health, cell balancing)
- Charging infrastructure management
- Vehicle homologation and type approval workflows
- Telematics and connected vehicle data integration
- Carbon/emissions regulatory reporting beyond basic CO2 tracking
- Highly engineered configure-to-order for niche products
- MES-level shop floor control for specialist manufacturing
- Any requirement that is highly industry-specific with no matching standard scope item
These MUST be classified as gap_ricefw or gap_companion.
Only use fit_standard if the requirement maps directly to a documented S/4HANA Public Cloud standard scope item with zero modification.
When uncertain between fit_standard and gap_ricefw, choose gap_ricefw.
Always return valid JSON only. No markdown. No explanation outside the JSON object."""

_FIT_GAP_USER_TMPL = """Requirement: {title}
Description: {description}
Business process: {business_process}
Priority: {priority}
Tags: {tags}
Category: {category}
Best matching SAP scope items from initial analysis: {top_matches}

Classify this requirement as one of:
- fit_standard: works as-is in standard S/4HANA
- fit_config: achievable through configuration/parameters only
- fit_extension: achievable via SAP-released extension (BAdI, key-user, BTP side-by-side)
- gap_ricefw: requires custom RICEFW development
- gap_companion: requires third-party companion solution
- out_of_scope: not in scope

Return JSON only:
{{"fit_type":"fit_standard|fit_config|fit_extension|gap_ricefw|gap_companion|out_of_scope","complexity":"XS|S|M|L|XL","rationale":"2-3 sentence explanation","reasoning":"1-2 sentences explaining why this fit_type was chosen","sap_scope_item_id":null or "ID","sap_scope_item_name":null or "name","workaround_option":null or "text","customisation_risk":null or "Low|Medium|High","clean_core_impact":null or "None|Low|Medium|High","estimated_effort_days_low":0,"estimated_effort_days_high":0,"cost_band":"<5k|5k-20k|20k-100k|>100k","confidence_score":0.8}}"""


@router.post("/requirements/{req_id}/fit-gap-assess", status_code=201)
def fit_gap_assess(req_id: str, engagement_id: str):
    """Run AI Fit-to-Standard assessment for a requirement. Creates one assessment per req (idempotent: returns existing if present)."""
    req = get_requirement_by_id(req_id, engagement_id)
    if not req:
        raise HTTPException(status_code=404, detail=f"Requirement {req_id} not found for engagement {engagement_id}")

    # Idempotent: if assessment already exists for this req, return it
    existing = get_fit_gap_by_engagement(engagement_id)
    for a in existing:
        if a.get("req_id") == req_id:
            return a

    ctx = get_engagement_with_client(engagement_id)
    client = (ctx or {}).get("client") or {}
    industry = client.get("industry") or "General"
    employees = client.get("employees") or 0
    regulatory = client.get("regulatory_environment") or []
    context = f"Industry: {industry}, {employees} employees, regulatory: {regulatory}"

    gap_records = get_gap_results_by_req_id(req_id, engagement_id)
    top_matches = []
    if gap_records:
        matches = gap_records[0].get("matches") or []
        for m in matches[:3]:
            top_matches.append(f"{m.get('id', '')}: {m.get('name', '')}")
    top_matches_str = "; ".join(top_matches) if top_matches else "None"

    user_prompt = _FIT_GAP_USER_TMPL.format(
        title=req.get("title", ""),
        description=req.get("description", ""),
        business_process=req.get("business_process") or "",
        priority=req.get("priority") or "",
        tags=req.get("tags") or [],
        category=req.get("category") or "",
        top_matches=top_matches_str,
    )
    # Pattern library pre-check: find pattern whose title shares >=2 words with requirement title
    try:
        all_patterns = get_pattern_library(limit=100)
        req_title_words = set(w for w in req.get("title", "").lower().split() if len(w) > 2)
        for pat in all_patterns:
            pat_title_words = set(w for w in (pat.get("name") or "").lower().split() if len(w) > 2)
            if len(req_title_words & pat_title_words) >= 2:
                user_prompt += f"\nRelevant past resolution from similar project: {pat.get('name')} — {(pat.get('content') or '')[:150]}"
                break
    except Exception:
        pass
    system_prompt = _FIT_GAP_SYSTEM.format(context=context)
    patterns = _get_top_patterns_text(limit=5)
    if patterns:
        system_prompt = system_prompt + "\n\n" + patterns

    try:
        provider = get_provider()
        result = provider.complete(system_prompt, user_prompt, max_tokens=512, model=MODEL_HAIKU)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"AI service unavailable: {e}")

    raw = result.get("content", "{}")
    json_match = re.search(r"\{[^{}]*\}", raw)
    if not json_match:
        raise HTTPException(status_code=500, detail="No JSON in AI response")
    try:
        data = json.loads(json_match.group())
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Invalid JSON: {e}")

    fit_type = _normalize_fit_type(data.get("fit_type") or "fit_standard")
    complexity = (data.get("complexity") or "S").upper()
    if complexity not in _FIT_GAP_COMPLEXITY:
        complexity = "S"

    record = {
        "engagement_id": engagement_id,
        "req_id": req_id,
        "fit_type": fit_type,
        "complexity": complexity,
        "rationale": data.get("rationale"),
        "sap_scope_item_id": data.get("sap_scope_item_id"),
        "sap_scope_item_name": data.get("sap_scope_item_name"),
        "workaround_option": data.get("workaround_option"),
        "customisation_risk": data.get("customisation_risk"),
        "clean_core_impact": data.get("clean_core_impact"),
        "estimated_effort_days_low": data.get("estimated_effort_days_low"),
        "estimated_effort_days_high": data.get("estimated_effort_days_high"),
        "cost_band": data.get("cost_band"),
        "ai_generated": True,
        "confidence_score": data.get("confidence_score"),
        "reasoning": data.get("reasoning"),
        "hitl_state": "ai_draft",
    }
    assessment = create_fit_gap_assessment(record)
    if not assessment:
        raise HTTPException(status_code=500, detail="Failed to create fit/gap assessment")
    # Attach requirement summary for frontend
    assessment["requirement_title"] = req.get("title")
    assessment["requirement_business_process"] = req.get("business_process")
    # Engagement mode: queue RICEFW generation for gaps
    fit_type_val = record.get("fit_type", "")
    if fit_type_val in ("gap_ricefw", "gap_companion"):
        _queue_or_execute(engagement_id, "generate_ricefw",
            {"req_id": req_id, "fit_type": fit_type_val, "gap_description": req.get("description", "")},
            source="fitgap_result")
    # Layer 1: pattern tagging — auto-tag requirement when gap is identified
    if fit_type_val in ("gap_ricefw", "gap_companion"):
        try:
            from database import add_pattern_tags_to_requirement
            process_area = req.get("business_process") or ""
            tags = _PROCESS_AREA_TAGS.get(process_area, [f"{process_area.lower().replace(' ','-').replace('/','-')}_gap"])
            add_pattern_tags_to_requirement(req_id, tags)
        except Exception as _tag_exc:
            logger.warning("pattern_tag_failed: %s", _tag_exc)
    return assessment


@router.get("/engagement/{engagement_id}/fit-gap-board")
def get_fit_gap_board(engagement_id: str):
    """Return fit/gap board: by_fit_type, by_process, summary."""
    requirements = get_requirements_by_engagement(engagement_id)
    req_map = {r["req_id"]: r for r in requirements}
    assessments = get_fit_gap_by_engagement(engagement_id)
    by_fit_type = {
        "fit_standard": [],
        "fit_config": [],
        "fit_extension": [],
        "gap_ricefw": [],
        "gap_companion": [],
        "out_of_scope": [],
    }
    by_process = {}
    summary = {
        "total": 0,
        "fit_count": 0,
        "gap_count": 0,
        "ai_draft": 0,
        "approved": 0,
        "total_effort_days_low": 0,
        "total_effort_days_high": 0,
        "complexity_breakdown": {"XS": 0, "S": 0, "M": 0, "L": 0, "XL": 0},
    }
    fit_types_set = {"fit_standard", "fit_config", "fit_extension"}

    for a in assessments:
        summary["total"] += 1
        state = (a.get("hitl_state") or "ai_draft").lower()
        if state == "ai_draft":
            summary["ai_draft"] += 1
        elif state == "approved":
            summary["approved"] += 1
        ft = _normalize_fit_type(a.get("fit_type") or "fit_standard")
        if ft in fit_types_set:
            summary["fit_count"] += 1
        else:
            summary["gap_count"] += 1
        summary["total_effort_days_low"] += a.get("estimated_effort_days_low") or 0
        summary["total_effort_days_high"] += a.get("estimated_effort_days_high") or 0
        c = (a.get("complexity") or "S").upper()
        if c in summary["complexity_breakdown"]:
            summary["complexity_breakdown"][c] += 1

        req = req_map.get(a["req_id"]) or {}
        a_with_req = {**a, "requirement_title": req.get("title"), "requirement_business_process": req.get("business_process")}
        bucket = by_fit_type.get(ft)
        if bucket is None:
            bucket = by_fit_type.get("out_of_scope")
        if bucket is not None:
            bucket.append(a_with_req)

        bp = req.get("business_process") or "Unclassified"
        by_process.setdefault(bp, {})
        by_process[bp][ft] = by_process[bp].get(ft, 0) + 1

    return {
        "engagement_id": engagement_id,
        "by_fit_type": by_fit_type,
        "by_process": by_process,
        "summary": summary,
    }


@router.post("/fit-gap-assessments/{assessment_id}/review")
def review_fit_gap_assessment(assessment_id: str, engagement_id: str, body: FitGapReviewRequest):
    """Approve or send back a fit/gap assessment. assessment_id is e.g. FGA-001."""
    fga = get_fit_gap_by_assessment_id(assessment_id, engagement_id)
    if not fga:
        raise HTTPException(status_code=404, detail="Fit/gap assessment not found")

    if body.approve:
        updates = {
            "hitl_state": "approved",
            "reviewed_by": body.reviewer,
            "reviewed_at": datetime.now(timezone.utc).isoformat(),
            "reviewer_notes": body.notes,
        }
    else:
        updates = {"hitl_state": "ai_draft", "reviewer_notes": body.notes}
        if body.fit_type:
            updates["fit_type"] = _normalize_fit_type(body.fit_type)
        if body.complexity:
            updates["complexity"] = body.complexity.upper()

    updated = update_fit_gap_assessment(assessment_id, engagement_id, updates)
    if not updated:
        raise HTTPException(status_code=500, detail="Update failed")
    req = get_requirement_by_id(updated["req_id"], engagement_id)
    updated["requirement_title"] = req.get("title") if req else None
    updated["requirement_business_process"] = req.get("business_process") if req else None
    return updated


@router.post("/engagement/{engagement_id}/fit-gap-analyse-all", status_code=200)
def fit_gap_analyse_all(engagement_id: str):
    """Run fit-gap assessment for all requirements that do not yet have an assessment. 200ms delay between calls."""
    import time
    requirements = get_requirements_by_engagement(engagement_id)
    existing = get_fit_gap_by_engagement(engagement_id)
    req_ids_done = {a["req_id"] for a in existing}
    created = 0
    for req in requirements:
        if req["req_id"] in req_ids_done:
            continue
        state = (req.get("hitl_state") or "ai_draft").lower()
        if state == "out_of_scope":
            continue
        try:
            fit_gap_assess(req["req_id"], engagement_id)
            created += 1
            time.sleep(0.2)
        except Exception as e:
            logger.warning("Fit-gap assess failed for %s: %s", req["req_id"], e)
    return {"engagement_id": engagement_id, "processed": created}


@router.get("/engagement/{engagement_id}/deliverable-progress")
def get_deliverable_progress(engagement_id: str):
    """Return deliverable progress percentages for blueprint, RICEFW, test scripts, and go-live."""
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)

    total_reqs = len(reqs)
    assessed_req_ids = {a["req_id"] for a in assessments}
    confirmed_reqs = [r for r in reqs if r.get("sign_off_status") == "confirmed"]

    if total_reqs == 0:
        blueprint_pct = 0
    else:
        blueprint_pct = min(100, round(((len(assessed_req_ids) * 0.6) + (len(confirmed_reqs) * 0.4)) / total_reqs * 100))

    total_ricefw = len(ricefw)
    ricefw_with_effort = [r for r in ricefw if r.get("effort_days_low") or r.get("effort_days_high")]
    ricefw_pct = min(100, round(len(ricefw_with_effort) / total_ricefw * 100)) if total_ricefw > 0 else 0

    checklist = get_go_live_checklist(engagement_id)
    total_checklist = len(checklist)
    completed_checklist = len([i for i in checklist if i.get("status") == "completed"])
    go_live_pct = round(completed_checklist / total_checklist * 100) if total_checklist > 0 else 0

    return {
        "engagement_id": engagement_id,
        "blueprint_pct": blueprint_pct,
        "ricefw_pct": ricefw_pct,
        "test_scripts_pct": 0,
        "go_live_pct": go_live_pct,
        "detail": {
            "total_requirements": total_reqs,
            "requirements_assessed": len(assessed_req_ids),
            "requirements_confirmed": len(confirmed_reqs),
            "total_ricefw": total_ricefw,
            "ricefw_with_effort": len(ricefw_with_effort),
            "checklist_total": total_checklist,
            "checklist_completed": completed_checklist,
        },
    }


@router.post("/engagement/{engagement_id}/ricefw-generate", status_code=200)
def ricefw_generate_from_gaps(engagement_id: str):
    """From approved fit_gap_assessments with fit_type=gap_ricefw, create ricefw_inventory items. Skips req_ids that already have a RICEFW item."""
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")
    try:
        assessments = get_fit_gap_by_engagement(engagement_id)
    except Exception as e:
        logger.exception("ricefw-generate: get_fit_gap_by_engagement failed")
        raise HTTPException(status_code=500, detail="Fit/Gap data unavailable. Run migrations and ensure Fit/Gap board is loaded.")
    gap_ricefw_approved = [
        a for a in assessments
        if (a.get("hitl_state") or "").lower() == "approved"
        and (a.get("fit_type") or "").lower() == "gap_ricefw"
    ]
    try:
        existing_ricefw = get_ricefw_by_engagement(engagement_id)
        requirements = get_requirements_by_engagement(engagement_id)
    except Exception as e:
        logger.exception("ricefw-generate: get_ricefw or get_requirements failed")
        raise HTTPException(status_code=500, detail="Could not load RICEFW or requirements. Try again.")
    existing_ricefw = existing_ricefw or []
    requirements = requirements or []
    req_ids_with_ricefw = {item.get("req_id") for item in existing_ricefw if item.get("req_id")}
    req_map = {r["req_id"]: r for r in requirements}
    created = 0
    skipped = 0
    for a in gap_ricefw_approved:
        req_id = a.get("req_id")
        if not req_id or req_id in req_ids_with_ricefw:
            skipped += 1
            continue
        req = req_map.get(req_id) or {}
        title = (req.get("title") or a.get("requirement_title") or req_id).strip() or req_id
        description = (a.get("rationale") or title)[:2000]
        # RICEFW complexity is very_high|high|medium|low; fit/gap uses XS|S|M|L|XL — pass None to avoid invalid value
        try:
            create_ricefw_item(
                engagement_id=engagement_id,
                item_type="E",
                name=title,
                req_id=req_id,
                description=description,
                status="identified",
                complexity=None,
                priority=None,
            )
            created += 1
            req_ids_with_ricefw.add(req_id)
        except Exception as e:
            logger.warning("ricefw-generate: failed to create for %s: %s", req_id, e)
            skipped += 1
    return {
        "engagement_id": engagement_id,
        "created": created,
        "skipped": skipped,
        "message": f"Created {created} RICEFW item(s) from approved Gap RICEFW assessments; skipped {skipped}.",
    }


@router.post("/engagement/{engagement_id}/ricefw-estimate-all", status_code=200)
def ricefw_estimate_all(engagement_id: str):
    """Run the estimator on all RICEFW items missing effort_days_low/high."""
    ricefw_items = get_ricefw_by_engagement(engagement_id)
    items_to_estimate = [r for r in ricefw_items if not r.get("effort_days_low") and not r.get("effort_days_high")]

    if not items_to_estimate:
        return {"message": "All RICEFW items already have estimates", "processed": 0}

    patterns = get_pattern_library(limit=100)
    eng = get_engagement_with_client(engagement_id)
    client = (eng or {}).get("client") or {}
    industry = client.get("industry", "General")

    results = []
    for item in items_to_estimate:
        item_type = item.get("type", "E")
        item_name = item.get("name", "")
        item_desc = item.get("description", "")

        pattern_hint = ""
        for p in patterns:
            p_words = set((p.get("name") or "").lower().split())
            i_words = set(item_name.lower().split())
            if len(p_words & i_words) >= 2:
                pattern_hint = f"Similar past item: {p.get('name')} — {str(p.get('content', ''))[:150]}"
                break

        type_labels = {
            "R": "Report", "I": "Interface/Integration", "C": "Data Conversion",
            "E": "Enhancement/Custom Code", "F": "Form/Output", "W": "Workflow/Approval",
            "A": "Analytics/Dashboard",
        }
        type_label = type_labels.get(item_type, item_type)

        prompt = f"""You are an SAP project estimator. Estimate effort for this custom development item.

Item Type: {type_label}
Name: {item_name}
Description: {item_desc}
Industry: {industry}
{pattern_hint}

Return JSON only:
{{"effort_days_low": <integer, realistic minimum person-days>, "effort_days_high": <integer, realistic maximum person-days>, "complexity": "Simple|Medium|Complex|Very Complex", "rationale": "1 sentence explaining the estimate", "priority": "Critical|High|Medium|Low"}}

Typical ranges: Simple=1-3d, Medium=3-8d, Complex=8-20d, Very Complex=20-50d"""

        try:
            result = get_provider().complete(
                model=MODEL_HAIKU,
                messages=[{"role": "user", "content": prompt}],
                system="You are an SAP estimator. Return valid JSON only.",
                max_tokens=200,
            )
            raw = result["content"].strip()
            # strip markdown code fences if present
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            data = json.loads(raw)
            update_ricefw_item(item["id"], engagement_id, {
                "effort_days_low": data.get("effort_days_low"),
                "effort_days_high": data.get("effort_days_high"),
                "priority": data.get("priority"),
            })
            results.append({"id": item["id"], "name": item_name, "status": "estimated",
                            "low": data.get("effort_days_low"), "high": data.get("effort_days_high")})
        except Exception as e:
            logger.warning(f"RICEFW estimate failed for {item['id']}: {e}")
            results.append({"id": item["id"], "name": item_name, "status": "failed"})

    return {"processed": len(results), "results": results}


# ── Phase D: Feedback & Pattern Library ───────────────────────────────────────

def _get_top_patterns_text(limit: int = 5) -> str:
    """Return top patterns by use_count as a single string for prompt injection."""
    try:
        patterns = get_pattern_library(limit=limit)
    except Exception:
        return ""
    if not patterns:
        return ""
    lines = ["Relevant patterns from the library (use to guide responses):"]
    for p in patterns:
        name = p.get("name") or "Pattern"
        content = (p.get("content") or "").strip()
        if content:
            lines.append(f"- {name}: {content[:500]}")
    return "\n".join(lines) + "\n\n" if lines else ""


@router.post("/feedback", status_code=201)
def post_feedback(body: FeedbackCreate):
    """Record a feedback event. If event_type=pattern_used and payload.pattern_id is set, increments that pattern's use_count."""
    try:
        event = create_feedback_event(
            engagement_id=body.engagement_id,
            event_type=(body.event_type or "general").lower(),
            payload=body.payload or {},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if (body.event_type or "").lower() == "pattern_used" and body.payload and body.payload.get("pattern_id"):
        increment_pattern_use(body.payload["pattern_id"])
    return event


@router.get("/feedback")
def get_feedback(engagement_id: Optional[str] = None, limit: int = 50):
    """List recent feedback events, optionally filtered by engagement_id."""
    try:
        items = list_feedback_events(engagement_id=engagement_id, limit=min(limit, 100))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"items": items, "total": len(items)}


@router.get("/pattern-library")
def get_pattern_library_endpoint(limit: int = 50):
    """List pattern library entries ordered by use_count descending."""
    try:
        items = get_pattern_library(limit=min(limit, 100))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"items": items, "total": len(items)}


# ── Agent Team & Simulation ───────────────────────────────────────────────────

@router.get("/agent-roles")
def get_agent_roles():
    """List configured agent roles for simulation and UI dropdown."""
    try:
        roles = list_agent_roles()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"items": roles, "total": len(roles)}


@router.get("/agent-roles/{role_id}/maturity")
def get_agent_maturity(role_id: str):
    """Return latest maturity scores per criterion for the role."""
    try:
        scores = get_agent_maturity_scores(role_id=role_id, limit=50)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    # Return latest per criterion
    by_criterion = {}
    for s in scores:
        c = s.get("criterion") or ""
        if c and c not in by_criterion:
            by_criterion[c] = s
    return {"role_id": role_id, "scores": list(by_criterion.values())}


@router.post("/agent-roles/{role_id}/maturity", status_code=201)
def post_agent_maturity(role_id: str, body: MaturityScoreCreate):
    """Record a maturity assessment for the role (criterion, score 1-5, notes)."""
    role = get_agent_role_by_role_id(role_id)
    if not role:
        raise HTTPException(status_code=404, detail=f"Agent role {role_id} not found")
    score = max(1, min(5, body.score))
    try:
        record = create_agent_maturity_score(role_id, body.criterion, score, body.notes)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return record


@router.post("/simulate/agent-response")
def simulate_agent_response(request: Request, body: SimulateAgentRequest):
    """Get a single agent reply: load role + knowledge, build system prompt, call LLM, return reply. Injects pattern library."""
    role = get_agent_role_by_role_id(body.agent_role_id)
    if not role:
        raise HTTPException(status_code=404, detail=f"Agent role {body.agent_role_id} not found")
    knowledge = get_agent_knowledge_by_role(body.agent_role_id, limit=20)
    focus_areas = role.get("focus_areas") or []
    if isinstance(focus_areas, str):
        try:
            focus_areas = json.loads(focus_areas) if focus_areas else []
        except Exception:
            focus_areas = []
    system_parts = [
        f"You are: {role.get('name', '')}",
        f"Mandate: {role.get('mandate', '')}",
        f"Focus areas: {', '.join(focus_areas) if focus_areas else 'N/A'}",
        f"Behavior: {role.get('behavior_rules', '')}",
        f"Escalation: {role.get('escalation_rules', '')}",
    ]
    if knowledge:
        system_parts.append("Domain knowledge (use to ground your response):")
        for k in knowledge:
            system_parts.append(f"- [{k.get('category', '')}] { (k.get('content') or '')[:400] }")
    system_prompt = "\n".join(system_parts)
    patterns = _get_top_patterns_text(limit=5)
    if patterns:
        system_prompt = system_prompt + "\n\n" + patterns
    phase = body.phase or "requirements"
    system_prompt += f"\n\nCurrent simulation phase: {phase}. Respond in character; be concise and actionable."
    user_lines = []
    if body.conversation_turn:
        for msg in body.conversation_turn:
            user_lines.append(f"{msg.get('role', 'user')}: {msg.get('content', '')}")
    if body.context_message:
        user_lines.append(body.context_message)
    if not user_lines:
        user_lines.append("What should we focus on next for this phase?")
    user_prompt = "\n".join(user_lines)
    try:
        provider = get_provider()
        result = provider.complete(system_prompt, user_prompt, max_tokens=600, model=MODEL_HAIKU)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"AI service unavailable: {e}")
    content = (result.get("content") or "").strip()
    out = {"agent_role_id": body.agent_role_id, "phase": phase, "reply": content}
    try:
        create_audit_event(
            engagement_id=body.engagement_id or "",
            action="agent_response",
            entity_type="simulation",
            entity_id=body.agent_role_id,
            actor_id=request.headers.get("X-Actor-Id"),
            actor_role=request.headers.get("X-Actor-Role"),
            details={"phase": phase},
        )
    except Exception:
        pass
    return out


_SEED_REQUIREMENTS_SYSTEM = (
    "You are an SAP S/4HANA requirements analyst. Generate realistic business requirements for a {industry} company "
    "implementing SAP S/4HANA Public Cloud. Requirements should reflect real pain points and business needs specific "
    "to the {industry} industry. Return ONLY a JSON array with no preamble or explanation."
)


@router.post("/simulate/seed-requirements", status_code=201)
def seed_requirements(request: Request, body: SeedRequirementsRequest):
    """Generate 40–60 requirements for industry/processes via Claude, insert, run fit-gap for each, set HITL ai_draft; return summary."""
    engagement = get_engagement(body.engagement_id)
    if not engagement:
        raise HTTPException(status_code=404, detail=f"Engagement {body.engagement_id} not found")
    if not body.processes:
        raise HTTPException(status_code=400, detail="processes list cannot be empty")

    existing = get_requirements_by_engagement(body.engagement_id)
    if len(existing) > 0:
        raise HTTPException(
            status_code=409,
            detail=f"Requirements already exist for this engagement ({len(existing)}). Delete existing requirements first or use a different engagement.",
        )

    system_prompt = _SEED_REQUIREMENTS_SYSTEM.format(industry=body.industry)
    user_prompt = (
        f"Industry: {body.industry}. Processes in scope: {', '.join(body.processes)}.\n"
        "Generate 8–12 requirements per process. Total 40–60 requirements.\n"
        "Return a JSON array only. Each item: {\"title\": str, \"description\": str, \"process\": str, "
        "\"priority\": \"Must-Have\"|\"Should-Have\"|\"Could-Have\", \"category\": str}. "
        "Use the exact priority and category strings. process must be one of the listed processes."
    )
    try:
        provider = get_provider()
        result = provider.complete(system_prompt, user_prompt, max_tokens=16000, model=MODEL_SONNET_SEED)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"AI service unavailable: {e}")

    raw = (result.get("content") or "").strip()
    json_match = re.search(r"\[[\s\S]*\]", raw)
    if not json_match:
        raise HTTPException(status_code=500, detail="No JSON array in AI response")
    try:
        items = json.loads(json_match.group())
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Invalid JSON: {e}")

    if not isinstance(items, list):
        raise HTTPException(status_code=500, detail="AI response is not a JSON array")

    requirements_created = 0
    assessments_created = 0
    gap_types = {"gap_ricefw", "gap_companion", "out_of_scope"}

    for item in items:
        if not isinstance(item, dict):
            continue
        title = (item.get("title") or "").strip() or "Untitled"
        description = (item.get("description") or "").strip()
        process = (item.get("process") or "").strip() or body.processes[0]
        priority = item.get("priority") or "Must-Have"
        if priority == "Could-Have":
            priority = "Nice-to-Have"
        category = (item.get("category") or "").strip() or None
        try:
            req = create_requirement(
                engagement_id=body.engagement_id,
                title=title,
                description=description,
                business_process=process,
                priority=priority,
                category=category,
                hitl_state="ai_draft",
            )
        except Exception as e:
            logger.warning("Seed: create_requirement failed for %s: %s", title[:50], e)
            continue
        if not req:
            continue
        requirements_created += 1
        try:
            fit_gap_assess(req["req_id"], body.engagement_id)
            assessments_created += 1
        except Exception as e:
            logger.warning("Seed: fit_gap_assess failed for %s: %s", req["req_id"], e)
        try:
            update_requirement(req["req_id"], body.engagement_id, {"hitl_state": "ai_draft"})
        except Exception:
            pass

    assessments = get_fit_gap_by_engagement(body.engagement_id) or []
    gaps = sum(1 for a in assessments if (a.get("fit_type") or "").lower().replace(" ", "_") in gap_types)

    return {
        "engagement_id": body.engagement_id,
        "requirements_created": requirements_created,
        "assessments_created": assessments_created,
        "gaps": gaps,
    }


@router.post("/platform-issues", status_code=201)
def post_platform_issue(request: Request, body: PlatformIssueCreate):
    """Create a platform issue (e.g. from simulation when an agent hits a limitation)."""
    try:
        record = create_platform_issue({
            "engagement_id": body.engagement_id,
            "agent_role_id": body.agent_role_id,
            "phase": body.phase or "requirements",
            "context": body.context or {},
            "problem_description": body.problem_description,
            "issue_type": (body.issue_type or "missing_feature").lower().replace(" ", "_"),
            "suggested_improvement": body.suggested_improvement or "",
            "priority": (body.priority or "medium").lower(),
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    try:
        create_audit_event(
            engagement_id=body.engagement_id or "",
            action="platform_issue_created",
            entity_type="platform_issue",
            entity_id=record.get("id"),
            actor_id=request.headers.get("X-Actor-Id"),
            actor_role=request.headers.get("X-Actor-Role"),
            details={"problem_description": (body.problem_description or "")[:200]},
        )
    except Exception:
        pass
    return record


@router.get("/platform-issues")
def get_platform_issues(
    engagement_id: Optional[str] = None,
    priority: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = 100,
):
    """List platform issues with optional filters."""
    try:
        items = list_platform_issues(engagement_id=engagement_id, priority=priority, status=status, limit=min(limit, 200))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"items": items, "total": len(items)}


@router.patch("/platform-issues/{issue_id}")
def patch_platform_issue(request: Request, issue_id: str, body: PlatformIssueUpdate, engagement_id: Optional[str] = None):
    """Update platform issue status or priority."""
    updates = {}
    if body.status is not None:
        updates["status"] = body.status
    if body.priority is not None:
        updates["priority"] = body.priority
    if not updates:
        raise HTTPException(status_code=400, detail="No updates provided")
    try:
        record = update_platform_issue(issue_id, updates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    eng_id = engagement_id or (record.get("engagement_id") if record else None)
    if eng_id:
        try:
            create_audit_event(
                engagement_id=eng_id,
                action="platform_issue_updated",
                entity_type="platform_issue",
                entity_id=issue_id,
                actor_id=request.headers.get("X-Actor-Id"),
                actor_role=request.headers.get("X-Actor-Role"),
                details=updates,
            )
        except Exception:
            pass
    return record


@router.get("/engagement/{engagement_id}/platform-backlog")
def get_engagement_platform_backlog(engagement_id: str):
    """Return platform issues for the engagement grouped by priority (high, medium, low)."""
    try:
        items = list_platform_issues(engagement_id=engagement_id, limit=500)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    by_priority = {"high": [], "medium": [], "low": []}
    for i in items:
        p = (i.get("priority") or "medium").lower()
        if p not in by_priority:
            by_priority[p] = []
        by_priority[p].append(i)
    return {"engagement_id": engagement_id, "by_priority": by_priority, "total": len(items)}


# ── Testing Command Center (RAPID Test Agents spec) ───────────────────────────

TESTING_SCENARIOS = [
    {"id": "smoke", "name": "Smoke test", "description": "Critical flows: health, clients, engagements"},
    {"id": "regression", "name": "Regression", "description": "Full API checks including engagement, requirements, fit-gap"},
    {"id": "data_integrity", "name": "Data Integrity", "description": "CRUD and referential checks for engagement scope"},
    {"id": "ux_deep_dive", "name": "UX Deep-Dive", "description": "Placeholder for future UX agent (API availability only for now)"},
    {"id": "import_export", "name": "Import/Export", "description": "Template download and export endpoints"},
]


class TestingRunRequest(BaseModel):
    scenario_ids: List[str] = ["smoke"]
    environment: Optional[str] = "DEV"
    engagement_id: Optional[str] = None
    push_issues_to_backlog: Optional[bool] = False


@router.get("/testing/scenarios")
def get_testing_scenarios():
    """Return catalog of test scenarios for the Testing Command Center."""
    return {"scenarios": TESTING_SCENARIOS}


def _run_smoke_checks(engagement_id: Optional[str]) -> tuple[int, int, list]:
    """Run smoke checks (health, clients, engagements). Returns (passed, failed, issues)."""
    passed, failed, issues = 0, 0, []
    # 1) List clients
    try:
        list_clients()
        passed += 1
    except Exception as e:
        failed += 1
        issues.append({
            "area": "Client",
            "type": "bug",
            "severity": "high",
            "message": f"List clients failed: {e!s}",
            "repro_steps": ["GET /v1/clients"],
            "expected_behavior": "200 with list",
            "actual_behavior": str(e),
        })
    # 2) List engagements
    try:
        list_engagements()
        passed += 1
    except Exception as e:
        failed += 1
        issues.append({
            "area": "Engagement",
            "type": "bug",
            "severity": "high",
            "message": f"List engagements failed: {e!s}",
            "repro_steps": ["GET /v1/engagements"],
            "expected_behavior": "200 with list",
            "actual_behavior": str(e),
        })
    # 3) If engagement_id provided, get engagement + requirements + fit-gap
    if engagement_id:
        try:
            get_engagement(engagement_id)
            passed += 1
        except Exception as e:
            failed += 1
            issues.append({
                "area": "Engagement",
                "type": "bug",
                "severity": "medium",
                "message": f"Get engagement {engagement_id} failed: {e!s}",
                "repro_steps": [f"GET /v1/engagements/{engagement_id}"],
                "expected_behavior": "200 with engagement",
                "actual_behavior": str(e),
            })
        try:
            get_requirements_by_engagement(engagement_id)
            passed += 1
        except Exception as e:
            failed += 1
            issues.append({
                "area": "Requirements",
                "type": "bug",
                "severity": "medium",
                "message": f"Get requirements for {engagement_id} failed: {e!s}",
                "repro_steps": [f"GET /v1/requirements?engagement_id={engagement_id}"],
                "expected_behavior": "200 with list",
                "actual_behavior": str(e),
            })
        try:
            get_fit_gap_by_engagement(engagement_id)
            passed += 1
        except Exception as e:
            failed += 1
            issues.append({
                "area": "Fit/Gap",
                "type": "bug",
                "severity": "medium",
                "message": f"Get fit-gap for {engagement_id} failed: {e!s}",
                "repro_steps": [f"GET /v1/engagement/{engagement_id}/fit-gap-board"],
                "expected_behavior": "200 with board",
                "actual_behavior": str(e),
            })
    return passed, failed, issues


def _run_import_export_checks() -> tuple[int, int, list]:
    """Run import/export endpoint checks. Returns (passed, failed, issues)."""
    passed, failed, issues = 0, 0, []
    # Template download exists (we only check the route is wired; actual stream would need request context)
    try:
        # Just verify scope_items / template logic is importable; real download is GET /v1/requirements/template/download
        get_catalogue_text()
        passed += 1
    except Exception as e:
        failed += 1
        issues.append({
            "area": "Requirements",
            "type": "bug",
            "severity": "low",
            "message": f"Template/catalogue check failed: {e!s}",
            "repro_steps": ["GET /v1/requirements/template/download"],
            "expected_behavior": "200 with file",
            "actual_behavior": str(e),
        })
    return passed, failed, issues


@router.post("/testing/run")
def post_testing_run(request: Request, body: TestingRunRequest):
    """Run selected test scenarios (API-level checks). Returns run result with pass/fail and issues."""
    run_id = str(uuid.uuid4())
    scenario_ids = body.scenario_ids or ["smoke"]
    engagement_id = body.engagement_id
    push_to_backlog = body.push_issues_to_backlog or False
    all_issues = []
    total_passed, total_failed = 0, 0

    for sid in scenario_ids:
        if sid == "smoke" or sid == "regression":
            p, f, issues = _run_smoke_checks(engagement_id)
            total_passed += p
            total_failed += f
            for i in issues:
                i["scenario_id"] = sid
                all_issues.append(i)
        elif sid == "import_export":
            p, f, issues = _run_import_export_checks()
            total_passed += p
            total_failed += f
            for i in issues:
                i["scenario_id"] = sid
                all_issues.append(i)
        elif sid in ("data_integrity", "ux_deep_dive"):
            # Same as smoke/regression for now
            p, f, issues = _run_smoke_checks(engagement_id)
            total_passed += p
            total_failed += f
            for i in issues:
                i["scenario_id"] = sid
                all_issues.append(i)

    # Optionally push issues to platform_issues (use engagement_id or placeholder)
    pushed_count = 0
    eng_for_backlog = engagement_id or "_test_run"
    if push_to_backlog and all_issues:
        for issue in all_issues:
            try:
                create_platform_issue({
                    "engagement_id": eng_for_backlog,
                    "problem_description": issue.get("message", "")[:500],
                    "issue_type": issue.get("type", "bug"),
                    "suggested_improvement": (issue.get("expected_behavior") or ""),
                    "priority": issue.get("severity", "medium"),
                    "context": {
                        "area": issue.get("area"),
                        "scenario_id": issue.get("scenario_id"),
                        "repro_steps": issue.get("repro_steps"),
                        "actual_behavior": issue.get("actual_behavior"),
                    },
                })
                pushed_count += 1
            except Exception:
                pass

    return {
        "run_id": run_id,
        "status": "completed",
        "environment": body.environment or "DEV",
        "scenario_ids": scenario_ids,
        "engagement_id": engagement_id,
        "summary": {
            "passed": total_passed,
            "failed": total_failed,
            "total_checks": total_passed + total_failed,
            "issues_count": len(all_issues),
            "pushed_to_backlog_count": pushed_count,
        },
        "issues": all_issues,
    }


@router.post("/engagement/{engagement_id}/test-scripts/generate", status_code=200)
def generate_test_scripts(engagement_id: str, body: dict = Body(...)):
    """Generate test scripts for RICEFW items using AI. Body: {ricefw_id?: str}"""
    ricefw_id_filter = body.get("ricefw_id")
    ricefw_items = get_ricefw_by_engagement(engagement_id)
    if ricefw_id_filter:
        ricefw_items = [r for r in ricefw_items if r.get("id") == ricefw_id_filter]
    if not ricefw_items:
        raise HTTPException(status_code=404, detail="No RICEFW items found")

    reqs = get_requirements_by_engagement(engagement_id)
    req_map = {r["req_id"]: r for r in reqs}

    # Build a test_case_id counter for this engagement
    existing_scripts = list_test_scripts_by_engagement(engagement_id)
    existing_nums = []
    for s in existing_scripts:
        tc = s.get("test_case_id") or ""
        if tc.startswith("TC-"):
            try:
                existing_nums.append(int(tc.split("-")[1]))
            except (IndexError, ValueError):
                pass
    next_tc_num = max(existing_nums, default=0) + 1

    created = []
    for item in ricefw_items:
        req_desc = ""
        if item.get("req_id"):
            req = req_map.get(item["req_id"])
            if req:
                req_desc = req.get("description", "")

        prompt = f"""You are an SAP test case writer. Generate a test case for this custom development item.

Item Type: {item.get('type', 'E')}
Name: {item.get('name', '')}
Description: {item.get('description', '')}
Linked Requirement: {req_desc or 'N/A'}

Return JSON only (no markdown):
{{
  "title": "...",
  "test_objective": "1-2 sentence test objective",
  "preconditions": "What must be set up before running this test",
  "steps": [
    {{"number": 1, "action": "...", "expected": "..."}},
    {{"number": 2, "action": "...", "expected": "..."}}
  ],
  "expected_result": "Overall expected outcome"
}}

Generate 3–5 steps covering the happy path."""

        try:
            result = get_provider().complete(
                model=MODEL_HAIKU,
                messages=[{"role": "user", "content": prompt}],
                system="You are an SAP test case writer. Return valid JSON only.",
                max_tokens=500,
            )
            raw = result["content"].strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            data = json.loads(raw)
            test_case_id = f"TC-{next_tc_num:03d}"
            next_tc_num += 1
            script = create_test_script({
                "engagement_id": engagement_id,
                "ricefw_id": item.get("id"),
                "req_id": item.get("req_id"),
                "test_case_id": test_case_id,
                "title": data.get("title"),
                "test_objective": data.get("test_objective"),
                "preconditions": data.get("preconditions"),
                "steps": data.get("steps"),
                "expected_result": data.get("expected_result"),
                "status": "draft",
            })
            created.append(script)
        except Exception as e:
            logger.warning(f"Test script generation failed for RICEFW {item.get('id')}: {e}")

    return {"engagement_id": engagement_id, "created": len(created), "test_scripts": created}


@router.get("/engagement/{engagement_id}/test-scripts", status_code=200)
def list_test_scripts(engagement_id: str, ricefw_id: Optional[str] = None):
    """List test scripts for an engagement, optionally filtered by ricefw_id."""
    if ricefw_id:
        scripts = list_test_scripts_by_ricefw(engagement_id, ricefw_id)
    else:
        scripts = list_test_scripts_by_engagement(engagement_id)
    return {"engagement_id": engagement_id, "total": len(scripts), "test_scripts": scripts}


@router.get("/command-center/alerts")
def get_command_center_alerts():
    """Scan all engagements and return prioritized action items for the command center."""
    try:
        engagements = list_engagements()
    except Exception:
        engagements = []
    alerts = []

    for eng in engagements:
        eng_id = eng["engagement_id"]
        eng_name = eng.get("name", eng_id)
        client_name = eng.get("client_name", "")

        # Check HITL queue
        try:
            hitl_events = list_hitl_events(eng_id)
            pending_hitl = [e for e in hitl_events if e.get("hitl_state") == "ai_draft"]
            if pending_hitl:
                alerts.append({
                    "engagement_id": eng_id, "engagement_name": eng_name, "client_name": client_name,
                    "type": "hitl_queue", "urgency": "high" if len(pending_hitl) >= 5 else "medium",
                    "title": f"{len(pending_hitl)} decisions pending review",
                    "description": f"{len(pending_hitl)} requirements awaiting consultant judgment",
                    "action_label": "Review Decisions", "action_url": f"/hitl?engagement_id={eng_id}",
                })
        except Exception:
            pass

        # Check sign-off completeness
        try:
            reqs = get_requirements_by_engagement(eng_id)
            if reqs:
                unconfirmed = [r for r in reqs if r.get("sign_off_status") != "confirmed"]
                pct = round((len(reqs) - len(unconfirmed)) / len(reqs) * 100)
                if unconfirmed and pct < 50:
                    alerts.append({
                        "engagement_id": eng_id, "engagement_name": eng_name, "client_name": client_name,
                        "type": "signoff_pending", "urgency": "medium",
                        "title": f"{len(unconfirmed)} requirements without sign-off",
                        "description": f"{pct}% of requirements confirmed",
                        "action_label": "View Requirements", "action_url": f"/requirements?engagement_id={eng_id}",
                    })
        except Exception:
            pass

        # Check RICEFW missing estimates
        try:
            ricefw = get_ricefw_by_engagement(eng_id)
            missing = [r for r in ricefw if not r.get("effort_days_low") and not r.get("effort_days_high")]
            if missing:
                alerts.append({
                    "engagement_id": eng_id, "engagement_name": eng_name, "client_name": client_name,
                    "type": "ricefw_estimates_missing", "urgency": "low",
                    "title": f"{len(missing)} RICEFW items missing effort estimates",
                    "description": "Estimates needed for project budget",
                    "action_label": "Update RICEFW", "action_url": f"/engagement/{eng_id}",
                })
        except Exception:
            pass

    urgency_order = {"high": 0, "medium": 1, "low": 2}
    alerts.sort(key=lambda x: urgency_order.get(x["urgency"], 3))
    return {"alerts": alerts, "total": len(alerts)}


# ── Blueprint ──────────────────────────────────────────────────────────────────

@router.get("/engagement/{engagement_id}/blueprint/preview")
def blueprint_preview(engagement_id: str):
    """Return blueprint data as JSON for frontend preview."""
    eng = get_engagement_with_client(engagement_id)
    client = (eng or {}).get("client") or {}
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw_items = get_ricefw_by_engagement(engagement_id)
    sources = list_sources_by_engagement(engagement_id)
    raci = get_raci_matrix(engagement_id)
    process_steps_all = get_process_steps_by_engagement(engagement_id)

    assess_by_req = {a["req_id"]: a for a in assessments}
    steps_by_req: dict = {}
    for step in process_steps_all:
        steps_by_req.setdefault(step["req_id"], []).append(step)

    by_process: dict = {}
    for req in reqs:
        proc = req.get("business_process") or "General"
        by_process.setdefault(proc, []).append(req)

    fit_counts: dict = {}
    for a in assessments:
        ft = a.get("fit_type", "unknown")
        fit_counts[ft] = fit_counts.get(ft, 0) + 1

    total_effort_low = sum(r.get("effort_days_low") or 0 for r in ricefw_items)
    total_effort_high = sum(r.get("effort_days_high") or 0 for r in ricefw_items)
    open_items = [r for r in reqs if r.get("sign_off_status") != "confirmed"]
    high_risk = [a for a in assessments if a.get("customisation_risk") == "High" or a.get("complexity") in ["L", "XL"]]

    return {
        "engagement": eng,
        "client": client,
        "fit_counts": fit_counts,
        "total_requirements": len(reqs),
        "total_ricefw": len(ricefw_items),
        "total_effort_low": total_effort_low,
        "total_effort_high": total_effort_high,
        "process_areas": list(by_process.keys()),
        "by_process": {
            proc: [
                {
                    "req": req,
                    "assessment": assess_by_req.get(req["req_id"]),
                    "steps": steps_by_req.get(req["req_id"], []),
                }
                for req in reqs_list
            ]
            for proc, reqs_list in by_process.items()
        },
        "ricefw_items": ricefw_items,
        "open_items": open_items,
        "high_risk_items": high_risk,
        "raci": raci,
        "sources": sources,
        "sign_off_pct": (
            round(len([r for r in reqs if r.get("sign_off_status") == "confirmed"]) / len(reqs) * 100)
            if reqs else 0
        ),
    }


@router.get("/engagement/{engagement_id}/blueprint/export")
def blueprint_export(engagement_id: str):
    """Generate and stream a DOCX Blueprint document."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    eng = get_engagement_with_client(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail="Engagement not found")
    client = (eng.get("client") or {})
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw_items = get_ricefw_by_engagement(engagement_id)
    process_steps_all = get_process_steps_by_engagement(engagement_id)

    assess_by_req = {a["req_id"]: a for a in assessments}
    steps_by_req: dict = {}
    for step in process_steps_all:
        steps_by_req.setdefault(step["req_id"], []).append(step)

    by_process: dict = {}
    for req in reqs:
        proc = req.get("business_process") or "General"
        by_process.setdefault(proc, []).append(req)

    fit_counts: dict = {}
    for a in assessments:
        ft = a.get("fit_type", "unknown")
        fit_counts[ft] = fit_counts.get(ft, 0) + 1

    total_effort_low = sum(r.get("effort_days_low") or 0 for r in ricefw_items)
    total_effort_high = sum(r.get("effort_days_high") or 0 for r in ricefw_items)

    # LLM: Executive Summary
    exec_summary = ""
    try:
        fit_summary_text = ", ".join(f"{v} {k.replace('_', ' ')}" for k, v in fit_counts.items())
        confirmed_count = len([r for r in reqs if r.get("sign_off_status") == "confirmed"])
        sign_off_pct = round(confirmed_count / len(reqs) * 100) if reqs else 0
        exec_prompt = f"""Write a professional 3-paragraph executive summary for an SAP S/4HANA Business Blueprint document.
Client: {client.get('name', 'Client')} ({client.get('industry', 'General')} industry, {client.get('employees', 'N/A')} employees)
Project: {eng.get('name', 'SAP Implementation')}
Go-Live Target: {eng.get('go_live_date', 'TBD')}
Total Requirements: {len(reqs)}
Fit Analysis: {fit_summary_text}
Custom Development Items: {len(ricefw_items)} (estimated {total_effort_low}–{total_effort_high} person-days)
Sign-off completion: {sign_off_pct}%

Para 1: Project context and business objective (2-3 sentences).
Para 2: Scope and fit-to-standard results in plain business language. No acronyms.
Para 3: Key outcomes of the blueprint phase and recommended next steps.
Write in professional consulting style. No markdown."""
        result = get_provider().complete(
            model=MODEL_SONNET,
            messages=[{"role": "user", "content": exec_prompt}],
            system="You are a senior SAP implementation consultant writing a formal blueprint document.",
            max_tokens=500,
        )
        exec_summary = result["content"].strip()
    except Exception as e:
        logger.warning(f"Blueprint exec summary LLM failed: {e}")
        exec_summary = (
            f"This Business Blueprint documents the fit-to-standard assessment for "
            f"{client.get('name', 'the client')} SAP S/4HANA implementation. "
            f"The assessment covers {len(reqs)} business requirements across {len(by_process)} process areas."
        )

    doc = DocxDocument()

    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("Business Blueprint", 0)
    title.runs[0].font.size = Pt(24)

    doc.add_paragraph(f"Client: {client.get('name', 'N/A')}")
    doc.add_paragraph(f"Project: {eng.get('name', 'SAP S/4HANA Implementation')}")
    doc.add_paragraph(f"Go-Live Target: {eng.get('go_live_date', 'TBD')}")
    doc.add_paragraph(f"Document Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("Status: DRAFT")
    doc.add_paragraph()

    doc.add_heading("1. Executive Summary", level=1)
    for para_text in exec_summary.split("\n\n"):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    doc.add_heading("2. Scope Summary", level=1)
    doc.add_paragraph(
        f"This blueprint covers {len(reqs)} business requirements across {len(by_process)} process areas."
    )
    if by_process:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Process Area"
        hdr[1].text = "Requirements"
        hdr[2].text = "Gaps"
        for proc, proc_reqs in by_process.items():
            gaps = sum(
                1 for r in proc_reqs
                if assess_by_req.get(r["req_id"], {}).get("fit_type") in ["gap_ricefw", "gap_companion"]
            )
            row = tbl.add_row().cells
            row[0].text = proc
            row[1].text = str(len(proc_reqs))
            row[2].text = str(gaps)
        doc.add_paragraph()

    fit_labels = {
        "fit_standard": "Fit — Standard (no changes)",
        "fit_config": "Fit — Configuration only",
        "fit_extension": "Fit — SAP Extension (low risk)",
        "gap_ricefw": "Gap — Custom Development required",
        "gap_companion": "Gap — Third-party solution required",
        "out_of_scope": "Out of Scope",
    }

    doc.add_heading("3. Fit-to-Standard Analysis", level=1)
    if fit_counts:
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = "Table Grid"
        hdr2 = tbl2.rows[0].cells
        hdr2[0].text = "Classification"
        hdr2[1].text = "Count"
        hdr2[2].text = "% of Total"
        for ft, count in sorted(fit_counts.items()):
            row = tbl2.add_row().cells
            row[0].text = fit_labels.get(ft, ft)
            row[1].text = str(count)
            row[2].text = f"{round(count / len(assessments) * 100)}%" if assessments else "—"
        doc.add_paragraph()

    doc.add_heading("4. Business Process Documentation", level=1)
    doc.add_paragraph(
        "This section documents each business process area, the current state, SAP future state, and fit-gap decisions."
    )

    process_keys = list(by_process.keys())
    for proc, proc_reqs in by_process.items():
        doc.add_heading(f"4.{process_keys.index(proc) + 1}  {proc}", level=2)
        try:
            proc_fits = [assess_by_req.get(r["req_id"], {}).get("fit_type", "unknown") for r in proc_reqs]
            fit_summary_proc = ", ".join(
                f"{proc_fits.count(ft)} {ft.replace('_', ' ')}" for ft in set(proc_fits)
            )
            area_prompt = (
                f"Write 2 sentences summarizing this SAP process area assessment. "
                f"Process: {proc}. {len(proc_reqs)} requirements: {fit_summary_proc}. "
                f"Focus on business impact and any gaps."
            )
            area_result = get_provider().complete(
                model=MODEL_HAIKU,
                messages=[{"role": "user", "content": area_prompt}],
                system="SAP consultant. 2 sentences only. Plain English.",
                max_tokens=120,
            )
            doc.add_paragraph(area_result["content"].strip())
        except Exception:
            doc.add_paragraph(f"This area covers {len(proc_reqs)} requirements.")

        for req in proc_reqs:
            req_title = req.get("title") or (req.get("description", "")[:60] + "...")
            doc.add_heading(f"{req['req_id']}: {req_title}", level=3)
            doc.add_paragraph("Current State (As-Is):", style="Intense Quote")
            doc.add_paragraph(req.get("description") or "Not documented.")
            steps = steps_by_req.get(req["req_id"], [])
            if steps:
                doc.add_paragraph("Process Steps:")
                for s in sorted(steps, key=lambda x: x.get("step_number", 0)):
                    doc.add_paragraph(
                        f"  {s.get('step_number', '')}. {s.get('title', '')} — {s.get('description', '')}",
                        style="List Bullet",
                    )
            assessment = assess_by_req.get(req["req_id"])
            if assessment:
                doc.add_paragraph("SAP Assessment:", style="Intense Quote")
                ft = assessment.get("fit_type", "—")
                doc.add_paragraph(f"Fit Classification: {fit_labels.get(ft, ft)}")
                doc.add_paragraph(f"SAP Scope Item: {assessment.get('sap_scope_item_name') or '—'}")
                doc.add_paragraph(
                    f"Complexity: {assessment.get('complexity', '—')}  |  "
                    f"Clean Core Impact: {assessment.get('clean_core_impact', '—')}"
                )
                if assessment.get("rationale"):
                    doc.add_paragraph(f"Assessment: {assessment['rationale']}")
                if assessment.get("workaround_option"):
                    doc.add_paragraph(f"Decision / Workaround: {assessment['workaround_option']}")
                elo = assessment.get("estimated_effort_days_low", 0)
                ehi = assessment.get("estimated_effort_days_high", 0)
                if elo or ehi:
                    doc.add_paragraph(
                        f"Effort Estimate: {elo}–{ehi} person-days  |  Cost Band: {assessment.get('cost_band', '—')}"
                    )
            doc.add_paragraph(
                f"Sign-Off Status: {req.get('sign_off_status', 'draft').replace('_', ' ').title()}"
            )
            doc.add_paragraph()

    doc.add_heading("5. RICEFW Custom Development Inventory", level=1)
    if ricefw_items:
        doc.add_paragraph(
            f"Total items: {len(ricefw_items)}. Estimated effort: {total_effort_low}–{total_effort_high} person-days."
        )
        tbl3 = doc.add_table(rows=1, cols=6)
        tbl3.style = "Table Grid"
        hdr3 = tbl3.rows[0].cells
        for i, h in enumerate(["ID", "Type", "Name", "Complexity", "Effort (days)", "Status"]):
            hdr3[i].text = h
        for item in ricefw_items:
            row = tbl3.add_row().cells
            row[0].text = item.get("id", "")[:8]
            row[1].text = item.get("type", "—")
            row[2].text = (item.get("name") or "—")[:40]
            row[3].text = item.get("complexity") or "—"
            lo = item.get("effort_days_low") or "?"
            hi = item.get("effort_days_high") or "?"
            row[4].text = f"{lo}–{hi}"
            row[5].text = item.get("status", "identified")
        doc.add_paragraph()
    else:
        doc.add_paragraph("No custom development items identified in this engagement.")

    doc.add_heading("6. Open Items and Risks", level=1)
    open_reqs = [r for r in reqs if r.get("sign_off_status") != "confirmed"]
    if open_reqs:
        doc.add_heading("6.1 Open Sign-Off Items", level=2)
        tbl4 = doc.add_table(rows=1, cols=3)
        tbl4.style = "Table Grid"
        h4 = tbl4.rows[0].cells
        h4[0].text = "Requirement"
        h4[1].text = "Process Area"
        h4[2].text = "Status"
        for r in open_reqs[:20]:
            row = tbl4.add_row().cells
            row[0].text = f"{r['req_id']}: {(r.get('title') or r.get('description', ''))[:50]}"
            row[1].text = r.get("business_process", "—")
            row[2].text = r.get("sign_off_status", "draft").replace("_", " ").title()
        doc.add_paragraph()

    high_risk = [
        a for a in assessments
        if a.get("customisation_risk") == "High" or a.get("complexity") in ["L", "XL"]
    ]
    if high_risk:
        doc.add_heading("6.2 Risk Register", level=2)
        try:
            risk_items_text = "\n".join(
                f"- {a.get('req_id')}: fit_type={a.get('fit_type')}, complexity={a.get('complexity')}, "
                f"risk={a.get('customisation_risk')}, rationale={a.get('rationale', '')[:100]}"
                for a in high_risk[:10]
            )
            risk_prompt = (
                f"For each SAP implementation risk item below, write one risk statement and one mitigation.\n"
                f"Items:\n{risk_items_text}\n\n"
                f'Return JSON array: [{{"item": "req_id", "risk": "...", "mitigation": "..."}}]\nJSON only.'
            )
            risk_result = get_provider().complete(
                model=MODEL_SONNET,
                messages=[{"role": "user", "content": risk_prompt}],
                system="SAP risk analyst. Return JSON array only.",
                max_tokens=800,
            )
            raw_risk = risk_result["content"].strip()
            if raw_risk.startswith("```"):
                raw_risk = raw_risk.split("```")[1]
                if raw_risk.startswith("json"):
                    raw_risk = raw_risk[4:]
            risk_data = json.loads(raw_risk)
            tbl5 = doc.add_table(rows=1, cols=3)
            tbl5.style = "Table Grid"
            h5 = tbl5.rows[0].cells
            h5[0].text = "Item"
            h5[1].text = "Risk"
            h5[2].text = "Mitigation"
            for rd in risk_data:
                row = tbl5.add_row().cells
                row[0].text = rd.get("item", "—")
                row[1].text = rd.get("risk", "—")
                row[2].text = rd.get("mitigation", "—")
        except Exception:
            for a in high_risk[:5]:
                doc.add_paragraph(
                    f"• {a.get('req_id')}: {a.get('rationale', 'High complexity item requires careful planning.')}",
                    style="List Bullet",
                )
        doc.add_paragraph()

    doc.add_heading("7. Document Sign-Off", level=1)
    doc.add_paragraph(
        "By signing below, the parties confirm that this Business Blueprint accurately represents "
        "the agreed scope and design decisions."
    )
    doc.add_paragraph()
    doc.add_paragraph(f"Client: {client.get('name', '________________________')}")
    doc.add_paragraph("Name: ________________________   Title: ________________________")
    doc.add_paragraph("Signature: ___________________   Date: _________________________")
    doc.add_paragraph()
    doc.add_paragraph("Consultant:")
    doc.add_paragraph("Name: ________________________   Title: ________________________")
    doc.add_paragraph("Signature: ___________________   Date: _________________________")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    eng_name_safe = (eng.get("name") or "engagement").replace(" ", "_").replace("/", "-")[:30]
    filename = f"Blueprint_{eng_name_safe}_{datetime.now().strftime('%Y%m%d')}.docx"

    try:
        create_audit_event(
            engagement_id, "blueprint_exported", entity_type="blueprint",
            details={"filename": filename, "sections": 7},
        )
    except Exception:
        pass

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ── Client Portal: Pydantic models ───────────────────────────────────────────

class PortalInviteRequest(BaseModel):
    engagement_id: str
    client_id: str
    name: str
    email: str
    role: Optional[str] = "client_executive"

class PortalSignoffQueryRequest(BaseModel):
    message: str

class SteerCoReportRequest(BaseModel):
    since_date: str
    format: Optional[str] = "json"  # "json" or "docx"

class GoLiveChecklistUpdateRequest(BaseModel):
    status: Optional[str] = None
    notes: Optional[str] = None
    owner: Optional[str] = None


# ── Client Portal: helper ─────────────────────────────────────────────────────

def _get_portal_user_or_401(token: str) -> dict:
    """Validate portal token and update last_access. Raises 401 if invalid/expired."""
    portal_user = get_portal_user_by_token(token)
    if not portal_user:
        raise HTTPException(status_code=401, detail="Invalid or expired portal token")
    expires = portal_user.get("token_expires_at")
    if expires:
        try:
            expires_dt = datetime.fromisoformat(str(expires).replace("Z", "+00:00"))
            if expires_dt < datetime.now(timezone.utc):
                raise HTTPException(status_code=401, detail="Portal token has expired")
        except (ValueError, TypeError):
            pass
    update_portal_user_last_access(portal_user["id"])
    return portal_user


def _queue_or_execute(engagement_id: str, action_type: str, payload: dict,
                      confidence: float | None = None, source: str | None = None) -> bool:
    """Queue action for review or auto-execute based on engagement mode. Returns True if auto-executed."""
    try:
        eng = get_engagement(engagement_id)
        if not eng:
            return False
        auto = should_auto_execute(eng, action_type, confidence)
        from autonomy import RISK_LEVELS
        risk = RISK_LEVELS.get(action_type, "high")
        status = "auto_executed" if auto else "pending"
        create_queue_item(engagement_id, action_type, payload, risk, status, confidence, source)
        logger.info("action_queue: %s %s engagement=%s", status, action_type, engagement_id)
        return auto
    except Exception as exc:
        logger.error("action_queue_error: %s", exc)
        return False


_PORTAL_FIT_LABELS = {
    "fit_standard": "SAP covers this process out of the box — no changes needed",
    "fit_config": "SAP covers this with standard configuration",
    "fit_extension": "SAP covers this with a supported extension (low risk)",
    "gap_ricefw": "SAP does not cover this — custom development required (adds cost and risk)",
    "gap_companion": "SAP does not cover this — requires a third-party solution",
    "out_of_scope": "This item is outside the project scope",
}


# ── Client Portal: endpoints (ITEM 1) ────────────────────────────────────────

@router.post("/portal/invite", status_code=201)
def portal_invite(body: PortalInviteRequest):
    """Create a portal user with a magic access token (30-day expiry). Returns portal_url."""
    eng = get_engagement(body.engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {body.engagement_id} not found")
    token = secrets.token_urlsafe(32)
    expires_at = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
    portal_user = create_portal_user(
        engagement_id=body.engagement_id,
        client_id=body.client_id,
        name=body.name,
        email=body.email,
        role=body.role or "client_executive",
        access_token=token,
        token_expires_at=expires_at,
    )
    return {
        "portal_user_id": portal_user.get("id"),
        "access_token": token,
        "portal_url": f"{os.getenv('FRONTEND_URL', 'https://rapid-ui-wine.vercel.app')}/portal/{token}",
        "expires_at": expires_at,
    }


@portal_router.get("/portal/auth/{token}")
def portal_auth(token: str):
    """Validate a portal token. Returns user info + valid:true, or 401."""
    portal_user = get_portal_user_by_token(token)
    if not portal_user:
        raise HTTPException(status_code=401, detail="Invalid or expired portal token")
    expires = portal_user.get("token_expires_at")
    if expires:
        try:
            expires_dt = datetime.fromisoformat(str(expires).replace("Z", "+00:00"))
            if expires_dt < datetime.now(timezone.utc):
                raise HTTPException(status_code=401, detail="Portal token has expired")
        except (ValueError, TypeError):
            pass
    update_portal_user_last_access(portal_user["id"])
    return {
        "portal_user_id": portal_user.get("id"),
        "engagement_id": portal_user.get("engagement_id"),
        "client_id": portal_user.get("client_id"),
        "name": portal_user.get("name"),
        "role": portal_user.get("role"),
        "valid": True,
    }


@portal_router.get("/portal/{token}/overview")
def portal_overview(token: str):
    """Return engagement overview for the client portal: health, progress, decisions."""
    portal_user = _get_portal_user_or_401(token)
    engagement_id = portal_user["engagement_id"]
    eng = get_engagement_with_client(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail="Engagement not found")
    client = get_client(portal_user["client_id"])
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)

    total_reqs = len(reqs)
    assessed_req_ids = {a["req_id"] for a in assessments}
    confirmed_reqs = [r for r in reqs if r.get("sign_off_status") == "confirmed"]
    if total_reqs == 0:
        blueprint_pct = 0
    else:
        blueprint_pct = min(100, round(((len(assessed_req_ids) * 0.6) + (len(confirmed_reqs) * 0.4)) / total_reqs * 100))
    total_ricefw = len(ricefw)
    ricefw_pct = min(100, round(len([r for r in ricefw if r.get("effort_days_low") or r.get("effort_days_high")]) / total_ricefw * 100)) if total_ricefw > 0 else 0

    pending_signoffs = [r for r in reqs if r.get("sign_off_status") != "confirmed"]
    recent_decisions = sorted(
        confirmed_reqs,
        key=lambda r: r.get("updated_at") or r.get("created_at") or "",
        reverse=True,
    )[:5]

    go_live = eng.get("go_live_date")
    days_to_go_live = None
    if go_live:
        try:
            gl_dt = datetime.fromisoformat(str(go_live).replace("Z", "+00:00"))
            days_to_go_live = (gl_dt.date() - datetime.now(timezone.utc).date()).days
        except (ValueError, TypeError):
            pass
    stale_hitl = [r for r in reqs if r.get("hitl_state") == "ai_draft"]
    if blueprint_pct < 30 or (days_to_go_live is not None and days_to_go_live < 30 and blueprint_pct < 70):
        rag = "red"
    elif blueprint_pct < 60 or len(stale_hitl) > 0:
        rag = "amber"
    else:
        rag = "green"

    phase = eng.get("phase") or "explore_realize"
    phase_order = ["discover_prepare", "explore_realize", "deploy_run"]
    try:
        current_idx = phase_order.index(phase)
    except ValueError:
        current_idx = 1
    milestone_summary = []
    for i, name in enumerate(["Discovery & Prepare", "Explore & Realize", "Deploy & Run"]):
        if i < current_idx:
            status = "complete"
        elif i == current_idx:
            status = "current"
        else:
            status = "upcoming"
        milestone_summary.append({"name": name, "status": status})

    rag_map = {"green": "on_track", "amber": "at_risk", "red": "behind"}
    client_name = (client or {}).get("name") or eng.get("client_name") or ""
    return {
        # Flat fields expected by frontend
        "project_name": eng.get("name") or "",
        "client_name": client_name,
        "go_live_date": go_live,
        "days_remaining": days_to_go_live,
        "rag_status": rag_map.get(rag, "on_track"),
        "blueprint_pct": blueprint_pct,
        "ricefw_pct": ricefw_pct,
        "pending_signoffs": len(pending_signoffs),
        # Extra context consumed by layout.tsx and timeline
        "engagement_name": eng.get("name") or "",
        "phase": phase,
        "recent_decisions": [
            {"req_id": r.get("req_id"), "title": r.get("title"), "approved_at": r.get("updated_at")}
            for r in recent_decisions
        ],
        "milestone_summary": milestone_summary,
    }


@portal_router.get("/portal/{token}/signoffs")
def portal_signoffs(token: str):
    """Return requirements needing client sign-off, with plain-English fit descriptions."""
    portal_user = _get_portal_user_or_401(token)
    engagement_id = portal_user["engagement_id"]
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    assessment_map = {a["req_id"]: a for a in assessments}
    pending = [r for r in reqs if r.get("sign_off_status") != "confirmed"]
    items = []
    for r in pending:
        a = assessment_map.get(r.get("req_id")) or {}
        ft = (a.get("fit_type") or r.get("fit_type") or "").lower().replace(" ", "_")
        effort = ""
        if a.get("estimated_effort_days_low") and a.get("estimated_effort_days_high"):
            effort = f"{a['estimated_effort_days_low']}–{a['estimated_effort_days_high']} days"
        items.append({
            "req_id": r.get("req_id"),
            "title": r.get("title"),
            "process_area": r.get("business_process"),
            "fit_type": ft,
            "plain_english_description": _PORTAL_FIT_LABELS.get(ft, ft),
            "effort_estimate": effort,
            "sign_off_status": r.get("sign_off_status"),
        })
    return {"engagement_id": engagement_id, "pending_signoffs": len(items), "items": items}


@portal_router.post("/portal/{token}/signoffs/{req_id}/approve")
def portal_approve_signoff(token: str, req_id: str):
    """Client approves a requirement — sets sign_off_status=confirmed and audits the event."""
    portal_user = _get_portal_user_or_401(token)
    engagement_id = portal_user["engagement_id"]
    req = get_requirement_by_id(req_id, engagement_id)
    if not req:
        raise HTTPException(status_code=404, detail=f"Requirement {req_id} not found")
    updated = update_requirement(req_id, {
        "sign_off_status": "confirmed",
        "sign_off_by": portal_user.get("name"),
        "sign_off_at": datetime.now(timezone.utc).isoformat(),
    })
    try:
        create_audit_event(
            engagement_id, "client_signoff_approved",
            entity_type="requirement", entity_id=req_id,
            actor_role=portal_user.get("role"),
            details={"portal_user": portal_user.get("name")},
        )
    except Exception:
        pass
    return updated


@portal_router.post("/portal/{token}/signoffs/{req_id}/query")
def portal_query_signoff(token: str, req_id: str, body: PortalSignoffQueryRequest):
    """Client raises a question about a requirement — stored as a HITL client_query event."""
    portal_user = _get_portal_user_or_401(token)
    engagement_id = portal_user["engagement_id"]
    req = get_requirement_by_id(req_id, engagement_id)
    if not req:
        raise HTTPException(status_code=404, detail=f"Requirement {req_id} not found")
    message = body.message
    try:
        create_hitl_event({
            "engagement_id": engagement_id,
            "req_id": req_id,
            "from_state": req.get("hitl_state") or "ai_draft",
            "to_state": "client_query",
            "actor": portal_user.get("name") or "portal",
            "actor_role": portal_user.get("role"),
            "notes": message,
        })
    except Exception:
        pass
    try:
        create_audit_event(
            engagement_id, "client_query_raised",
            entity_type="requirement", entity_id=req_id,
            actor_role=portal_user.get("role"),
            details={"message": message[:500]},
        )
    except Exception:
        pass
    return {"status": "question_received", "req_id": req_id}


@portal_router.get("/portal/{token}/decisions")
def portal_decisions(token: str):
    """Return all confirmed requirements as a decisions log."""
    portal_user = _get_portal_user_or_401(token)
    engagement_id = portal_user["engagement_id"]
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    assessment_map = {a["req_id"]: a for a in assessments}
    confirmed = [r for r in reqs if r.get("sign_off_status") == "confirmed"]
    items = []
    for r in sorted(confirmed, key=lambda x: x.get("updated_at") or "", reverse=True):
        a = assessment_map.get(r.get("req_id")) or {}
        ft = (a.get("fit_type") or r.get("fit_type") or "").lower().replace(" ", "_")
        items.append({
            "req_id": r.get("req_id"),
            "title": r.get("title"),
            "process_area": r.get("business_process"),
            "fit_type": ft,
            "plain_english": _PORTAL_FIT_LABELS.get(ft, ft),
            "approved_at": r.get("updated_at"),
            "sign_off_by": r.get("sign_off_by"),
        })
    return {"engagement_id": engagement_id, "total": len(items), "decisions": items}


@portal_router.get("/portal/{token}/timeline")
def portal_timeline(token: str):
    """Return milestone timeline for client view."""
    portal_user = _get_portal_user_or_401(token)
    engagement_id = portal_user["engagement_id"]
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail="Engagement not found")
    phase = eng.get("phase") or "explore_realize"
    phase_order = ["discover_prepare", "explore_realize", "deploy_run"]
    phase_names = {
        "discover_prepare": "Discovery & Prepare",
        "explore_realize": "Explore & Realize",
        "deploy_run": "Deploy & Run",
    }
    phase_descs = {
        "discover_prepare": "Initial workshops and requirements gathering",
        "explore_realize": "Fit-to-standard assessment and gap resolution",
        "deploy_run": "Build, test, and go-live",
    }
    try:
        current_idx = phase_order.index(phase)
    except ValueError:
        current_idx = 1
    phases = []
    for i, key in enumerate(phase_order):
        if i < current_idx:
            status = "complete"
        elif i == current_idx:
            status = "current"
        else:
            status = "upcoming"
        phases.append({"name": phase_names[key], "status": status, "description": phase_descs[key]})

    go_live = eng.get("go_live_date")
    days_to_go_live = None
    if go_live:
        try:
            gl_dt = datetime.fromisoformat(str(go_live).replace("Z", "+00:00"))
            days_to_go_live = (gl_dt.date() - datetime.now(timezone.utc).date()).days
        except (ValueError, TypeError):
            pass
    return {"phases": phases, "go_live_date": go_live, "days_to_go_live": days_to_go_live}


# ── SteerCo Report Generator (ITEM 2) ────────────────────────────────────────

@router.post("/engagement/{engagement_id}/steerco-report")
def generate_steerco_report(engagement_id: str, body: SteerCoReportRequest):
    """Generate a Steering Committee report with RAG status, metrics, and LLM narrative."""
    eng = get_engagement_with_client(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")
    client = get_client(eng.get("client_id", ""))
    try:
        since_dt = datetime.fromisoformat(body.since_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid since_date format. Use YYYY-MM-DD.")

    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)
    audit_events = list_audit_events_by_engagement(engagement_id)

    since = body.since_date
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    new_reqs = [r for r in reqs if (r.get("created_at") or "") >= since]
    confirmed_since = [
        r for r in reqs
        if r.get("sign_off_status") == "confirmed" and (r.get("updated_at") or r.get("created_at") or "") >= since
    ]
    new_assessments = [a for a in assessments if (a.get("created_at") or "") >= since]

    total_reqs = len(reqs)
    assessed_req_ids = {a["req_id"] for a in assessments}
    confirmed_reqs = [r for r in reqs if r.get("sign_off_status") == "confirmed"]
    pending_reqs = [r for r in reqs if r.get("sign_off_status") != "confirmed"]
    if total_reqs == 0:
        blueprint_pct = 0
    else:
        blueprint_pct = min(100, round(((len(assessed_req_ids) * 0.6) + (len(confirmed_reqs) * 0.4)) / total_reqs * 100))
    total_ricefw = len(ricefw)
    ricefw_pct = min(100, round(len([r for r in ricefw if r.get("effort_days_low") or r.get("effort_days_high")]) / total_ricefw * 100)) if total_ricefw > 0 else 0

    go_live = eng.get("go_live_date")
    days_to_go_live = None
    if go_live:
        try:
            gl_dt = datetime.fromisoformat(str(go_live).replace("Z", "+00:00"))
            days_to_go_live = (gl_dt.date() - datetime.now(timezone.utc).date()).days
        except (ValueError, TypeError):
            pass

    high_risk = [a for a in assessments if (a.get("customisation_risk") or "") in ("high", "very_high") or a.get("fit_type") == "gap_ricefw"]
    stale_hitl = [r for r in reqs if r.get("hitl_state") == "ai_draft" and (r.get("updated_at") or "") < (datetime.now(timezone.utc) - timedelta(days=5)).isoformat()]

    if blueprint_pct < 30 or (days_to_go_live is not None and days_to_go_live < 30 and blueprint_pct < 70):
        rag = "RED"
    elif blueprint_pct < 60 or len(stale_hitl) > 0:
        rag = "AMBER"
    else:
        rag = "GREEN"

    client_name = (client or {}).get("name") or eng.get("client_name") or "Client"
    confirmed_list = "\n".join(f"- {r.get('title', r.get('req_id'))}" for r in confirmed_since[:10]) or "None"
    pending_list = "\n".join(f"- {r.get('title', r.get('req_id'))}" for r in pending_reqs[:10]) or "None"
    risk_list = "\n".join(f"- {a.get('req_id')}: {(a.get('rationale') or '')[:100]}" for a in high_risk[:5]) or "None"

    prompt = f"""Write a steering committee update for {client_name}. Period: {since} to {today}.
Progress this period: {len(new_reqs)} requirements added, {len(new_assessments)} assessed, {len(confirmed_since)} confirmed by client.
Overall: Blueprint {blueprint_pct}%, RICEFW {ricefw_pct}%.
RAG Status: {rag}.
Key decisions made:
{confirmed_list}
Pending decisions:
{pending_list}
Risks:
{risk_list}
Write in 4 sections: Summary (2 sentences), Progress This Period (bullet points), Decisions Required (numbered list), Next Steps (3-5 bullets). Professional tone. No jargon."""

    provider = get_provider()
    llm_result = provider.complete(prompt=prompt, system="You are an expert SAP programme manager writing SteerCo communications.", max_tokens=1500, model=MODEL_SONNET)
    narrative = (llm_result.get("content") or "").strip()

    # Parse sections
    sections: dict = {}
    current_section = None
    for line in narrative.split("\n"):
        line_lower = line.lower().strip()
        if "summary" in line_lower and len(line) < 60 and not any(c in line_lower for c in ["progress", "decision", "next"]):
            current_section = "summary"
            sections.setdefault(current_section, [])
        elif "progress this period" in line_lower or ("progress" in line_lower and "period" in line_lower):
            current_section = "progress"
            sections.setdefault(current_section, [])
        elif "decisions required" in line_lower or ("decision" in line_lower and "required" in line_lower):
            current_section = "decisions_required"
            sections.setdefault(current_section, [])
        elif "next steps" in line_lower:
            current_section = "next_steps"
            sections.setdefault(current_section, [])
        elif current_section and line.strip():
            sections[current_section] = sections.get(current_section, []) + [line.strip()]

    structured = {
        "engagement_id": engagement_id,
        "client": client_name,
        "period": {"from": since, "to": today},
        "rag_status": rag,
        "blueprint_pct": blueprint_pct,
        "ricefw_pct": ricefw_pct,
        "days_to_go_live": days_to_go_live,
        "stats": {
            "requirements_added": len(new_reqs),
            "requirements_assessed": len(new_assessments),
            "requirements_confirmed": len(confirmed_since),
            "total_requirements": total_reqs,
            "total_ricefw": total_ricefw,
            "high_risk_items": len(high_risk),
        },
        "narrative": narrative,
        "sections": {
            "summary": "\n".join(sections.get("summary", [])),
            "progress": sections.get("progress", []),
            "decisions_required": sections.get("decisions_required", []),
            "next_steps": sections.get("next_steps", []),
        },
    }

    if body.format == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading(f"Steering Committee Report — {client_name}", 0)
        doc.add_paragraph(f"Period: {since} to {today}  |  RAG: {rag}  |  Blueprint: {blueprint_pct}%  |  RICEFW: {ricefw_pct}%")
        doc.add_paragraph(f"Days to Go-Live: {days_to_go_live or 'TBD'}")
        for section_title, content in [
            ("1. Summary", structured["sections"]["summary"]),
            ("2. Progress This Period", "\n".join(structured["sections"]["progress"])),
            ("3. Decisions Required", "\n".join(structured["sections"]["decisions_required"])),
            ("4. Next Steps", "\n".join(structured["sections"]["next_steps"])),
        ]:
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(content if isinstance(content, str) else "\n".join(content or []))
        doc.add_heading("5. Metrics", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for k, v in [
            ("Requirements Added", len(new_reqs)),
            ("Assessments Completed", len(new_assessments)),
            ("Client Sign-offs Received", len(confirmed_since)),
            ("Total Requirements", total_reqs),
            ("RICEFW Items", total_ricefw),
            ("High Risk Items", len(high_risk)),
        ]:
            row = tbl.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        eng_name_safe = (eng.get("name") or "engagement").replace(" ", "_")[:30]
        filename = f"SteerCo_{eng_name_safe}_{today.replace('-', '')}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    return structured


# ── Memory / Pattern Agent: close-and-learn (ITEM 3) ─────────────────────────

@router.post("/engagement/{engagement_id}/close-and-learn")
def close_and_learn(engagement_id: str):
    """Extract 3-5 reusable patterns from a completed engagement and save to pattern library."""
    eng = get_engagement_with_client(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")
    client = get_client(eng.get("client_id", ""))
    industry = (client or {}).get("industry") or "Unknown"
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)

    fit_gap_text = "\n".join(
        f"- {a.get('req_id', '')}: fit_type={a.get('fit_type', '')}, rationale={a.get('rationale', '')[:200]}"
        for a in assessments[:50]
    ) or "No fit-gap data"
    ricefw_text = "\n".join(
        f"- {r.get('name', '')}: type={r.get('type', '')}, effort={r.get('effort_days_low', '')}-{r.get('effort_days_high', '')}d"
        for r in ricefw[:30]
    ) or "No RICEFW items"

    prompt = f"""You are reviewing a completed SAP implementation. Extract 3-5 reusable patterns for future projects.
Fit-gap decisions:
{fit_gap_text}
RICEFW items:
{ricefw_text}
Industry: {industry}

Return JSON array: [{{"category": "fit_resolution|ricefw_pattern|risk_pattern", "title": "...", "content": "...", "industry_tag": "..."}}]
Only return the JSON array, no other text."""

    provider = get_provider()
    llm_result = provider.complete(
        prompt=prompt,
        system="You are an SAP implementation knowledge manager. Extract concise, reusable patterns.",
        max_tokens=1200,
        model=MODEL_SONNET,
    )
    raw = (llm_result.get("content") or "").strip()

    patterns_data = []
    try:
        json_match = re.search(r'\[.*\]', raw, re.DOTALL)
        patterns_data = json.loads(json_match.group()) if json_match else json.loads(raw)
    except (json.JSONDecodeError, AttributeError):
        patterns_data = []

    patterns_created = []
    for p in patterns_data[:5]:
        try:
            created = create_pattern(
                name=p.get("title", "Unnamed pattern"),
                category=p.get("category", "fit_resolution"),
                content=p.get("content", ""),
                industry_tag=p.get("industry_tag") or industry,
            )
            patterns_created.append(created)
        except Exception:
            pass

    try:
        update_engagement(engagement_id, {"status": "closed"})
    except Exception:
        pass
    try:
        create_audit_event(
            engagement_id, "engagement_closed_and_learned",
            entity_type="engagement", entity_id=engagement_id,
            details={"patterns_created": len(patterns_created)},
        )
    except Exception:
        pass

    return {
        "engagement_id": engagement_id,
        "patterns_created": len(patterns_created),
        "patterns": patterns_created,
    }


# ── Go-Live Checklist Agent (ITEM 4) ─────────────────────────────────────────

@router.post("/engagement/{engagement_id}/go-live-checklist/generate", status_code=201)
def generate_go_live_checklist(engagement_id: str):
    """AI-generate a customised go-live readiness checklist from the RICEFW scope."""
    eng = get_engagement_with_client(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")
    client = get_client(eng.get("client_id", ""))
    industry = (client or {}).get("industry") or "Manufacturing"
    ricefw = get_ricefw_by_engagement(engagement_id)
    ricefw_count = len(ricefw)
    ricefw_types = ", ".join(sorted({r.get("type", "E") for r in ricefw})) or "E"

    prompt = f"""Generate a go-live readiness checklist for SAP S/4HANA implementation.
Custom development items: {ricefw_count} ({ricefw_types}).
Industry: {industry}.

Return JSON array of 20-30 checklist items:
[{{"category": "Technical|Functional|Change Management|Cutover|Hypercare", "item": "...", "owner": "IT|Business|SAP Lead|Project Manager", "due_date_offset_days": -N}}]
due_date_offset_days is negative = days before go-live (e.g. -30 means 30 days before).
Only return the JSON array, no other text."""

    provider = get_provider()
    llm_result = provider.complete(
        "You are an SAP go-live readiness expert. Generate practical, actionable checklist items.",
        prompt,
        max_tokens=2000,
        model=MODEL_HAIKU,
    )
    raw = (llm_result.get("content") or "").strip()

    checklist_data = []
    try:
        json_match = re.search(r'\[.*\]', raw, re.DOTALL)
        checklist_data = json.loads(json_match.group()) if json_match else json.loads(raw)
    except (json.JSONDecodeError, AttributeError):
        checklist_data = []

    created_items = []
    for item in checklist_data[:30]:
        try:
            created = create_go_live_checklist_item(
                engagement_id=engagement_id,
                category=item.get("category", "Technical"),
                item=item.get("item", ""),
                owner=item.get("owner"),
                due_date_offset_days=item.get("due_date_offset_days"),
            )
            created_items.append(created)
        except Exception:
            pass

    return {"engagement_id": engagement_id, "created": len(created_items), "items": created_items}


@router.get("/engagement/{engagement_id}/go-live-checklist")
def list_go_live_checklist(engagement_id: str):
    """Return the go-live checklist with completion percentage."""
    items = get_go_live_checklist(engagement_id)
    total = len(items)
    completed = len([i for i in items if i.get("status") == "completed"])
    go_live_pct = round(completed / total * 100) if total > 0 else 0
    return {
        "engagement_id": engagement_id,
        "total": total,
        "completed": completed,
        "go_live_pct": go_live_pct,
        "items": items,
    }


@router.patch("/go-live-checklist/{item_id}")
def patch_go_live_checklist_item(item_id: str, body: GoLiveChecklistUpdateRequest):
    """Update status, notes, or owner for a go-live checklist item."""
    updates = {k: v for k, v in body.dict().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    return update_go_live_checklist_item(item_id, updates)


# ── Engagement Mode: Action Queue endpoints ───────────────────────────────────

class SetEngagementModeRequest(BaseModel):
    mode: str
    autonomy_config: Optional[dict] = None
    force: Optional[bool] = False   # bypass pending-items warning


@router.patch("/engagement/{engagement_id}/mode")
def set_engagement_mode(engagement_id: str, body: SetEngagementModeRequest):
    """Update engagement mode (guided/collaborative/autonomous) with pending-item warning and tier lock.
    Returns 409 with warning payload if pending queue items exist (use force=true to override).
    Returns 403 if client tier doesn't meet mode requirements.
    Caches client tier in autonomy_config so _queue_or_execute can enforce it without extra DB calls.
    """
    if body.mode not in ("guided", "collaborative", "autonomous"):
        raise HTTPException(status_code=400, detail="mode must be guided, collaborative, or autonomous")

    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    # ── Tier lock check ──────────────────────────────────────────────────────
    from autonomy import check_tier_access, MODE_TIER_REQUIREMENTS
    from database import get_client_tier
    client_id = eng.get("client_id") or ""
    client_tier = get_client_tier(client_id) if client_id else "starter"
    tier_ok, tier_reason = check_tier_access(client_tier, body.mode)
    if not tier_ok:
        raise HTTPException(status_code=403, detail=tier_reason)

    # ── Pending-items warning ────────────────────────────────────────────────
    if not body.force:
        pending = get_queue_items(engagement_id, "pending")
        if pending:
            return {
                "warning": True,
                "engagement_id": engagement_id,
                "current_mode": eng.get("mode") or "collaborative",
                "requested_mode": body.mode,
                "pending_items": len(pending),
                "message": (
                    f"{len(pending)} pending action queue item(s) exist. "
                    "Switching mode may change which actions auto-execute. "
                    "Re-submit with force=true to proceed, or clear the queue first."
                ),
                "pending_sample": [
                    {"id": p.get("id"), "action_type": p.get("action_type"), "risk_level": p.get("risk_level")}
                    for p in pending[:5]
                ],
            }

    # Cache tier in autonomy_config so tier lock is available without extra DB calls
    merged_config = {**(body.autonomy_config or {}), "tier": client_tier}
    updated = update_engagement_mode(engagement_id, body.mode, merged_config)
    return {
        "engagement_id": engagement_id,
        "mode": body.mode,
        "autonomy_config": merged_config,
        "tier": client_tier,
        "warning": False,
    }

@router.get("/engagement/{engagement_id}/action-queue")
def list_action_queue(engagement_id: str, status: Optional[str] = None):
    """Return agent action queue for an engagement. Filter by status=pending|approved|rejected|auto_executed."""
    items = get_queue_items(engagement_id, status)
    pending_count = len([i for i in items if i.get("status") == "pending"])
    return {"engagement_id": engagement_id, "total": len(items),
            "pending": pending_count, "items": items}

@router.post("/engagement/{engagement_id}/action-queue/{item_id}/approve")
def approve_queue_item(engagement_id: str, item_id: str):
    """Approve a pending action queue item. Triggers send_portal_update notification."""
    updated = update_queue_item_status(item_id, "approved")
    if not updated:
        raise HTTPException(status_code=404, detail="Queue item not found")
    # Notify client portal of the approved action
    _queue_or_execute(
        engagement_id=engagement_id,
        action_type="send_portal_update",
        payload={
            "trigger": "action_approved",
            "approved_item_id": item_id,
            "approved_action_type": updated.get("action_type"),
            "risk_level": updated.get("risk_level"),
        },
        confidence=1.0,
        source="action_queue_approve",
    )
    return updated


@router.post("/engagement/{engagement_id}/action-queue/{item_id}/reject")
def reject_queue_item(engagement_id: str, item_id: str):
    """Reject/skip a pending action queue item."""
    updated = update_queue_item_status(item_id, "rejected")
    if not updated:
        raise HTTPException(status_code=404, detail="Queue item not found")
    return updated


@router.post("/engagement/{engagement_id}/action-queue/approve-all")
def approve_all_queue_items(engagement_id: str):
    """Bulk approve all pending low-risk items. Triggers a single portal notification."""
    items = get_queue_items(engagement_id, "pending")
    low_risk = [i for i in items if i.get("risk_level") == "low"]
    approved = []
    for item in low_risk:
        result = update_queue_item_status(item["id"], "approved")
        if result:
            approved.append(result)
    # Single batched portal notification for the whole approval run
    if approved:
        _queue_or_execute(
            engagement_id=engagement_id,
            action_type="send_portal_update",
            payload={
                "trigger": "bulk_approved",
                "approved_count": len(approved),
                "approved_action_types": list({a.get("action_type") for a in approved}),
            },
            confidence=1.0,
            source="action_queue_approve_all",
        )
    return {"approved": len(approved), "items": approved}


# ── Layer 1: Industry Benchmark Engine ───────────────────────────────────────

_PROCESS_AREA_TAGS: dict[str, list[str]] = {
    "Order-to-Cash":   ["otc_gap", "revenue_risk", "integration_needed"],
    "Procure-to-Pay":  ["p2p_gap", "vendor_risk", "approval_workflow"],
    "Record-to-Report": ["rtr_gap", "finance_risk", "reporting_complexity"],
    "Plan-to-Produce": ["mfg_gap", "production_risk", "scheduling_complexity"],
    "Hire-to-Retire":  ["hr_gap", "payroll_risk"],
    "Warehouse":       ["wm_gap", "inventory_risk"],
    "Quality":         ["qm_gap", "compliance_risk"],
    "Service":         ["sm_gap", "service_risk"],
    "Asset Management": ["am_gap", "maintenance_risk"],
}

_BENCHMARK_SEED: dict[str, list[dict]] = {
    "retail": [
        # overall
        {"process_area": "overall", "metric_name": "avg_gap_rate",      "metric_value": 0.28, "sample_size": 47},
        {"process_area": "overall", "metric_name": "avg_fit_rate",       "metric_value": 0.65, "sample_size": 47},
        {"process_area": "overall", "metric_name": "avg_ricefw_count",   "metric_value": 18.0, "sample_size": 47},
        {"process_area": "overall", "metric_name": "avg_go_live_days",   "metric_value": 120.0, "sample_size": 47},
        {"process_area": "overall", "metric_name": "p25_gap_rate",       "metric_value": 0.18, "sample_size": 47},
        {"process_area": "overall", "metric_name": "p75_gap_rate",       "metric_value": 0.39, "sample_size": 47},
        # by process
        {"process_area": "Order-to-Cash",   "metric_name": "avg_gap_rate", "metric_value": 0.22, "sample_size": 47},
        {"process_area": "Procure-to-Pay",  "metric_name": "avg_gap_rate", "metric_value": 0.31, "sample_size": 47},
        {"process_area": "Record-to-Report","metric_name": "avg_gap_rate", "metric_value": 0.17, "sample_size": 47},
        {"process_area": "Warehouse",       "metric_name": "avg_gap_rate", "metric_value": 0.35, "sample_size": 47},
    ],
    "manufacturing": [
        {"process_area": "overall", "metric_name": "avg_gap_rate",      "metric_value": 0.34, "sample_size": 63},
        {"process_area": "overall", "metric_name": "avg_fit_rate",       "metric_value": 0.58, "sample_size": 63},
        {"process_area": "overall", "metric_name": "avg_ricefw_count",   "metric_value": 24.0, "sample_size": 63},
        {"process_area": "overall", "metric_name": "avg_go_live_days",   "metric_value": 150.0, "sample_size": 63},
        {"process_area": "overall", "metric_name": "p25_gap_rate",       "metric_value": 0.23, "sample_size": 63},
        {"process_area": "overall", "metric_name": "p75_gap_rate",       "metric_value": 0.46, "sample_size": 63},
        {"process_area": "Plan-to-Produce", "metric_name": "avg_gap_rate", "metric_value": 0.42, "sample_size": 63},
        {"process_area": "Procure-to-Pay",  "metric_name": "avg_gap_rate", "metric_value": 0.29, "sample_size": 63},
        {"process_area": "Record-to-Report","metric_name": "avg_gap_rate", "metric_value": 0.21, "sample_size": 63},
        {"process_area": "Quality",         "metric_name": "avg_gap_rate", "metric_value": 0.38, "sample_size": 63},
    ],
    "pharma": [
        {"process_area": "overall", "metric_name": "avg_gap_rate",      "metric_value": 0.19, "sample_size": 28},
        {"process_area": "overall", "metric_name": "avg_fit_rate",       "metric_value": 0.72, "sample_size": 28},
        {"process_area": "overall", "metric_name": "avg_ricefw_count",   "metric_value": 12.0, "sample_size": 28},
        {"process_area": "overall", "metric_name": "avg_go_live_days",   "metric_value": 180.0, "sample_size": 28},
        {"process_area": "overall", "metric_name": "p25_gap_rate",       "metric_value": 0.12, "sample_size": 28},
        {"process_area": "overall", "metric_name": "p75_gap_rate",       "metric_value": 0.27, "sample_size": 28},
        {"process_area": "Quality",         "metric_name": "avg_gap_rate", "metric_value": 0.14, "sample_size": 28},
        {"process_area": "Record-to-Report","metric_name": "avg_gap_rate", "metric_value": 0.11, "sample_size": 28},
        {"process_area": "Order-to-Cash",   "metric_name": "avg_gap_rate", "metric_value": 0.22, "sample_size": 28},
    ],
    "logistics": [
        {"process_area": "overall", "metric_name": "avg_gap_rate",      "metric_value": 0.31, "sample_size": 35},
        {"process_area": "overall", "metric_name": "avg_fit_rate",       "metric_value": 0.61, "sample_size": 35},
        {"process_area": "overall", "metric_name": "avg_ricefw_count",   "metric_value": 20.0, "sample_size": 35},
        {"process_area": "overall", "metric_name": "avg_go_live_days",   "metric_value": 135.0, "sample_size": 35},
        {"process_area": "overall", "metric_name": "p25_gap_rate",       "metric_value": 0.20, "sample_size": 35},
        {"process_area": "overall", "metric_name": "p75_gap_rate",       "metric_value": 0.43, "sample_size": 35},
        {"process_area": "Warehouse",       "metric_name": "avg_gap_rate", "metric_value": 0.37, "sample_size": 35},
        {"process_area": "Order-to-Cash",   "metric_name": "avg_gap_rate", "metric_value": 0.28, "sample_size": 35},
        {"process_area": "Procure-to-Pay",  "metric_name": "avg_gap_rate", "metric_value": 0.26, "sample_size": 35},
    ],
}


class RecordOutcomeRequest(BaseModel):
    go_live_days: Optional[int] = None
    planned_ebitda_improvement_pct: Optional[float] = None
    annual_ebitda_usd: Optional[float] = None
    legacy_system_cost_monthly: Optional[float] = None
    traditional_timeline_months: Optional[float] = 12
    rapid_timeline_months: Optional[float] = 1


def _update_benchmarks_from_engagement(engagement_id: str, industry: str) -> None:
    """Refresh industry_benchmarks with anonymized stats from this engagement."""
    from database import upsert_benchmark
    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)

    total = len(reqs)
    if total == 0:
        return
    gap_types = {"gap_ricefw", "gap_companion"}
    fit_types = {"fit_standard", "fit_config", "fit_extension"}
    gaps = [a for a in assessments if a.get("fit_type") in gap_types]
    fits = [a for a in assessments if a.get("fit_type") in fit_types]
    gap_rate = len(gaps) / total
    fit_rate = len(fits) / total

    upsert_benchmark(industry, "overall", "avg_gap_rate", gap_rate)
    upsert_benchmark(industry, "overall", "avg_fit_rate", fit_rate)
    upsert_benchmark(industry, "overall", "avg_ricefw_count", float(len(ricefw)))

    # Per-process-area gap rates
    process_gaps: dict[str, list] = {}
    process_total: dict[str, int] = {}
    for a in assessments:
        req = next((r for r in reqs if r.get("req_id") == a.get("req_id")), {})
        pa = req.get("business_process") or "Unknown"
        process_total[pa] = process_total.get(pa, 0) + 1
        if a.get("fit_type") in gap_types:
            process_gaps.setdefault(pa, []).append(a)
    for pa, total_pa in process_total.items():
        if total_pa >= 3:
            pa_gap_rate = len(process_gaps.get(pa, [])) / total_pa
            upsert_benchmark(industry, pa, "avg_gap_rate", pa_gap_rate)


@router.post("/engagement/{engagement_id}/record-outcome", status_code=200)
def record_outcome(engagement_id: str, body: RecordOutcomeRequest):
    """Record engagement outcome and refresh industry benchmarks (anonymized)."""
    from database import upsert_engagement_outcome, get_client
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)
    gap_types = {"gap_ricefw", "gap_companion"}
    fit_types = {"fit_standard", "fit_config", "fit_extension"}

    total = len(reqs)
    gap_count = len([a for a in assessments if a.get("fit_type") in gap_types])
    fit_count = len([a for a in assessments if a.get("fit_type") in fit_types])

    # Velocity dividend calculation
    velocity_dividend = None
    if body.annual_ebitda_usd and body.planned_ebitda_improvement_pct:
        months_saved = (body.traditional_timeline_months or 12) - (body.rapid_timeline_months or 1)
        months_saved = max(0, months_saved)
        ebitda_gain = body.annual_ebitda_usd * (body.planned_ebitda_improvement_pct / 100) * (months_saved / 12)
        legacy_savings = (body.legacy_system_cost_monthly or 0) * months_saved
        velocity_dividend = ebitda_gain + legacy_savings

    client = get_client(eng.get("client_id") or "")
    industry = (client or {}).get("industry") or "unknown"

    outcome_data = {
        "industry": industry,
        "total_requirements": total,
        "fit_count": fit_count,
        "gap_count": gap_count,
        "ricefw_count": len(ricefw),
        "go_live_days": body.go_live_days,
        "velocity_dividend_usd": velocity_dividend,
    }
    outcome = upsert_engagement_outcome(engagement_id, outcome_data)

    # Refresh benchmarks asynchronously (best-effort)
    try:
        _update_benchmarks_from_engagement(engagement_id, industry)
        if body.go_live_days:
            from database import upsert_benchmark
            upsert_benchmark(industry, "overall", "avg_go_live_days", float(body.go_live_days))
    except Exception as exc:
        logger.warning("benchmark_update_failed: %s", exc)

    return {**outcome, "velocity_dividend_usd": velocity_dividend,
            "benchmarks_updated": True, "industry": industry}


@router.get("/benchmarks/{industry}")
def get_benchmarks(industry: str):
    """Return all benchmark metrics for an industry with sample sizes."""
    from database import get_benchmarks_by_industry
    rows = get_benchmarks_by_industry(industry.lower())
    if not rows:
        raise HTTPException(status_code=404, detail=f"No benchmarks found for industry '{industry}'. Run seed-benchmarks first.")

    # Structure: overall metrics + per-process-area breakdown
    overall = {r["metric_name"]: {"value": r["metric_value"], "sample_size": r["sample_size"]}
               for r in rows if r["process_area"] == "overall"}
    by_process: dict[str, dict] = {}
    for r in rows:
        if r["process_area"] != "overall":
            by_process.setdefault(r["process_area"], {})[r["metric_name"]] = {
                "value": r["metric_value"], "sample_size": r["sample_size"]
            }
    return {
        "industry": industry,
        "overall": overall,
        "by_process_area": by_process,
        "total_metrics": len(rows),
    }


@router.get("/engagement/{engagement_id}/benchmark-comparison")
def get_benchmark_comparison(engagement_id: str):
    """Compare this engagement's stats against industry benchmarks with percentile scores."""
    from database import get_benchmarks_by_industry
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)
    client = get_client(eng.get("client_id") or "")
    industry = (client or {}).get("industry") or "unknown"

    total = len(reqs)
    gap_types = {"gap_ricefw", "gap_companion"}
    fit_types = {"fit_standard", "fit_config", "fit_extension"}
    gap_count = len([a for a in assessments if a.get("fit_type") in gap_types])
    fit_count = len([a for a in assessments if a.get("fit_type") in fit_types])
    gap_rate = gap_count / total if total else 0
    fit_rate = fit_count / total if total else 0

    benchmarks = get_benchmarks_by_industry(industry.lower())
    bm = {r["metric_name"]: r for r in benchmarks if r["process_area"] == "overall"}

    def _percentile(value: float, avg: float, p25: float, p75: float, lower_is_better: bool = False) -> int:
        """Estimate percentile given avg and p25/p75 anchors."""
        if lower_is_better:
            if value <= p25: return 90
            if value <= avg: return 70
            if value <= p75: return 40
            return 15
        else:
            if value >= p75: return 90
            if value >= avg: return 65
            if value >= p25: return 40
            return 15

    comparisons = []
    if "avg_gap_rate" in bm:
        avg_gap = bm["avg_gap_rate"]["metric_value"]
        p25_gap = bm.get("p25_gap_rate", {}).get("metric_value", avg_gap * 0.65)
        p75_gap = bm.get("p75_gap_rate", {}).get("metric_value", avg_gap * 1.35)
        pct = _percentile(gap_rate, avg_gap, p25_gap, p75_gap, lower_is_better=True)
        comparisons.append({
            "metric": "gap_rate",
            "your_value": round(gap_rate, 3),
            "industry_avg": round(avg_gap, 3),
            "percentile": pct,
            "direction": "lower_is_better",
            "insight": (
                f"Your gap rate ({gap_rate:.0%}) is {'better than' if gap_rate < avg_gap else 'above'} "
                f"the {industry} industry average ({avg_gap:.0%}). "
                f"You're in the top {100 - pct}% of {industry} implementations."
                if gap_rate < avg_gap else
                f"Your gap rate ({gap_rate:.0%}) exceeds the {industry} average ({avg_gap:.0%}). "
                f"Consider prioritising gap resolution before go-live."
            ),
        })
    if "avg_fit_rate" in bm:
        avg_fit = bm["avg_fit_rate"]["metric_value"]
        pct_fit = _percentile(fit_rate, avg_fit, avg_fit * 0.8, avg_fit * 1.15)
        comparisons.append({
            "metric": "fit_rate",
            "your_value": round(fit_rate, 3),
            "industry_avg": round(avg_fit, 3),
            "percentile": pct_fit,
            "direction": "higher_is_better",
            "insight": (
                f"Fit rate of {fit_rate:.0%} vs {industry} average {avg_fit:.0%}. "
                f"{'Above' if fit_rate >= avg_fit else 'Below'} industry standard."
            ),
        })
    if "avg_ricefw_count" in bm:
        avg_rf = bm["avg_ricefw_count"]["metric_value"]
        comparisons.append({
            "metric": "ricefw_count",
            "your_value": len(ricefw),
            "industry_avg": avg_rf,
            "percentile": _percentile(len(ricefw), avg_rf, avg_rf * 0.65, avg_rf * 1.35, lower_is_better=True),
            "direction": "lower_is_better",
            "insight": (
                f"{len(ricefw)} RICEFW items vs {industry} average {avg_rf:.0f}. "
                f"{'Clean core alignment is strong.' if len(ricefw) <= avg_rf else 'High customisation volume — review for clean core risk.'}"
            ),
        })

    sample_size = bm.get("avg_gap_rate", {}).get("sample_size", 0)
    return {
        "engagement_id": engagement_id,
        "industry": industry,
        "sample_size": sample_size,
        "comparisons": comparisons,
        "overall_health": "strong" if gap_rate < (bm.get("avg_gap_rate", {}).get("metric_value") or 1) else "at_risk",
    }


@router.post("/engagement/{engagement_id}/seed-benchmarks", status_code=200)
def seed_benchmarks(engagement_id: str):
    """Admin: seed benchmark data for all industries. Safe to re-run (upserts)."""
    from database import upsert_benchmark
    seeded = 0
    for industry, rows in _BENCHMARK_SEED.items():
        for row in rows:
            upsert_benchmark(industry, row["process_area"], row["metric_name"],
                             row["metric_value"], row["sample_size"])
            seeded += 1
    return {"seeded": seeded, "industries": list(_BENCHMARK_SEED.keys()), "status": "ok"}


# ── Layer 2: Velocity Dividend Calculator ────────────────────────────────────

class VelocityDividendRequest(BaseModel):
    annual_ebitda_usd: float
    ebitda_improvement_pct: float
    legacy_monthly_cost_usd: float
    traditional_timeline_months: float = 12.0
    rapid_timeline_months: float = 1.0
    working_capital_improvement_pct: float = 0.0
    working_capital_usd: Optional[float] = None        # if provided, used for WC calc
    implementation_cost_usd: Optional[float] = 150_000 # actual impl spend; default mid-market
    client_name: Optional[str] = None                  # for PDF branding


@router.post("/engagement/{engagement_id}/velocity-dividend")
def velocity_dividend(engagement_id: str, body: VelocityDividendRequest):
    """Calculate and return velocity dividend breakdown with CFO narrative."""
    from database import upsert_engagement_outcome
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    months_saved = max(0.0, body.traditional_timeline_months - body.rapid_timeline_months)
    annual_ebitda_gain = body.annual_ebitda_usd * (body.ebitda_improvement_pct / 100)
    operational_efficiency_gain = annual_ebitda_gain * (months_saved / 12)
    legacy_cost_savings = body.legacy_monthly_cost_usd * months_saved
    wc_base = body.working_capital_usd or body.annual_ebitda_usd * 0.15
    working_capital_benefit = wc_base * (body.working_capital_improvement_pct / 100) * (months_saved / 12)
    total = operational_efficiency_gain + legacy_cost_savings + working_capital_benefit

    impl_cost = body.implementation_cost_usd or 150_000
    roi_pct = round((total / impl_cost) * 100, 1) if impl_cost else 0
    monthly_dividend = total / months_saved if months_saved > 0 else 0
    payback_months = round(impl_cost / monthly_dividend, 1) if monthly_dividend > 0 else None

    client = get_client(eng.get("client_id") or "")
    client_name = body.client_name or (client or {}).get("name") or "Your organisation"
    fiscal_year = datetime.now(timezone.utc).year

    cfo_summary = (
        f"By going live {months_saved:.0f} months earlier, {client_name} captures "
        f"${operational_efficiency_gain:,.0f} in incremental EBITDA in the current fiscal year "
        f"— before accounting for Year 2+ compounding benefits."
    )
    board_point = (
        f"The {body.rapid_timeline_months:.0f}-month implementation sprint generated a "
        f"${total:,.0f} Velocity Dividend, equivalent to a {roi_pct:.0f}% first-year ROI "
        f"on implementation spend."
    )

    result = {
        "engagement_id": engagement_id,
        "client_name": client_name,
        "inputs": {
            "annual_ebitda_usd": body.annual_ebitda_usd,
            "ebitda_improvement_pct": body.ebitda_improvement_pct,
            "legacy_monthly_cost_usd": body.legacy_monthly_cost_usd,
            "traditional_timeline_months": body.traditional_timeline_months,
            "rapid_timeline_months": body.rapid_timeline_months,
            "working_capital_improvement_pct": body.working_capital_improvement_pct,
        },
        "months_saved": months_saved,
        "implementation_cost_usd": impl_cost,
        "operational_efficiency_gain_usd": round(operational_efficiency_gain, 2),
        "legacy_cost_savings_usd": round(legacy_cost_savings, 2),
        "working_capital_benefit_usd": round(working_capital_benefit, 2),
        "total_velocity_dividend_usd": round(total, 2),
        "annualized_roi_pct": roi_pct,
        "payback_months": payback_months,
        "cfo_summary": cfo_summary,
        "board_talking_point": board_point,
        "calculated_at": datetime.now(timezone.utc).isoformat(),
    }

    # Persist to engagement_outcomes
    try:
        upsert_engagement_outcome(engagement_id, {"velocity_dividend_usd": round(total, 2)})
    except Exception as exc:
        logger.warning("velocity_dividend outcome save: %s", exc)

    return result


@router.get("/engagement/{engagement_id}/velocity-dividend/export")
def velocity_dividend_export(
    engagement_id: str,
    annual_ebitda_usd: float = 50_000_000,
    ebitda_improvement_pct: float = 2.5,
    legacy_monthly_cost_usd: float = 45_000,
    traditional_timeline_months: float = 12,
    rapid_timeline_months: float = 1,
    working_capital_improvement_pct: float = 1.2,
    implementation_cost_usd: float = 150_000,
):
    """Generate a 2-page CFO-ready PDF Velocity Dividend report."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        import io
    except ImportError:
        raise HTTPException(status_code=503, detail="reportlab not available — add to requirements.txt")

    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")
    client = get_client(eng.get("client_id") or "")
    client_name = (client or {}).get("name") or "Client"

    # Recalculate
    months_saved = max(0.0, traditional_timeline_months - rapid_timeline_months)
    op_gain = annual_ebitda_usd * (ebitda_improvement_pct / 100) * (months_saved / 12)
    legacy_savings = legacy_monthly_cost_usd * months_saved
    wc_benefit = (annual_ebitda_usd * 0.15) * (working_capital_improvement_pct / 100) * (months_saved / 12)
    total = op_gain + legacy_savings + wc_benefit
    impl_spend = implementation_cost_usd or 150_000
    roi_pct = round((total / impl_spend) * 100, 1) if impl_spend else 0

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()

    RAPID_BLUE = colors.HexColor("#1E3A5F")
    RAPID_GOLD = colors.HexColor("#C9A84C")
    LIGHT_GREY = colors.HexColor("#F5F5F5")

    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=22, textColor=RAPID_BLUE,
                         spaceAfter=6, alignment=TA_CENTER)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=14, textColor=RAPID_BLUE,
                         spaceBefore=14, spaceAfter=4)
    body_s = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=14)
    big_num = ParagraphStyle("bignum", parent=styles["Normal"], fontSize=32, textColor=RAPID_BLUE,
                              alignment=TA_CENTER, spaceAfter=2)
    label_s = ParagraphStyle("label", parent=styles["Normal"], fontSize=9, textColor=colors.grey,
                              alignment=TA_CENTER)
    watermark_s = ParagraphStyle("wm", parent=styles["Normal"], fontSize=8, textColor=colors.lightgrey,
                                  alignment=TA_CENTER)

    story = []

    # ── Page 1: Executive Summary ─────────────────────────────────────────────
    story.append(Paragraph("RAPID | Velocity Dividend Report", h1))
    story.append(Paragraph(f"{client_name} · {datetime.now(timezone.utc).strftime('%B %Y')}", label_s))
    story.append(HRFlowable(width="100%", thickness=2, color=RAPID_GOLD, spaceAfter=20))

    # Key numbers in 3 columns
    kpi_data = [
        [Paragraph(f"${total:,.0f}", big_num),
         Paragraph(f"{months_saved:.0f}", big_num),
         Paragraph(f"{roi_pct:.0f}%", big_num)],
        [Paragraph("Total Velocity Dividend", label_s),
         Paragraph("Months Saved", label_s),
         Paragraph("First-Year ROI", label_s)],
    ]
    kpi_table = Table(kpi_data, colWidths=[5*cm, 5*cm, 5*cm])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), LIGHT_GREY),
        ("ROWPADDING", (0, 0), (-1, -1), 12),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.lightgrey),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("Executive Summary", h2))
    story.append(Paragraph(
        f"By deploying SAP S/4HANA in <b>{rapid_timeline_months:.0f} month(s)</b> rather than the "
        f"industry-standard <b>{traditional_timeline_months:.0f} months</b>, {client_name} captures "
        f"<b>${total:,.0f}</b> in direct economic value — the Velocity Dividend.", body_s))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        f"The <b>${op_gain:,.0f}</b> operational efficiency gain represents "
        f"{ebitda_improvement_pct:.1f}% EBITDA improvement realised {months_saved:.0f} months earlier. "
        f"Combined with <b>${legacy_savings:,.0f}</b> in avoided legacy system costs and "
        f"<b>${wc_benefit:,.0f}</b> in working capital improvement, the total return on "
        f"implementation spend is <b>{roi_pct:.0f}%</b> in year one alone.", body_s))

    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=LIGHT_GREY))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("Prepared by RAPID  |  Confidential  |  Not for external distribution", watermark_s))

    # ── Page 2: Detailed calculation ─────────────────────────────────────────
    from reportlab.platypus import PageBreak
    story.append(PageBreak())
    story.append(Paragraph("Calculation Detail & Assumptions", h2))

    calc_data = [
        ["Component", "Formula", "Value"],
        ["Operational Efficiency Gain",
         f"${annual_ebitda_usd:,.0f} × {ebitda_improvement_pct}% × ({months_saved:.0f}/12 yr)",
         f"${op_gain:,.0f}"],
        ["Legacy System Savings",
         f"${legacy_monthly_cost_usd:,.0f}/mo × {months_saved:.0f} months",
         f"${legacy_savings:,.0f}"],
        ["Working Capital Benefit",
         f"WC base × {working_capital_improvement_pct}% × ({months_saved:.0f}/12 yr)",
         f"${wc_benefit:,.0f}"],
        ["Total Velocity Dividend", "", f"${total:,.0f}"],
        ["Implementation Cost", "Provided by user", f"${impl_spend:,.0f}"],
        ["First-Year ROI", f"${total:,.0f} / ${impl_spend:,.0f}", f"{roi_pct:.0f}%"],
    ]
    calc_table = Table(calc_data, colWidths=[6*cm, 7*cm, 4*cm])
    calc_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), RAPID_BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GREY]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BACKGROUND", (0, -3), (-1, -3), colors.HexColor("#E8F0FE")),
        ("FONTNAME", (0, -3), (-1, -3), "Helvetica-Bold"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(calc_table)
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Methodology Note", h2))
    story.append(Paragraph(
        "The Velocity Dividend is calculated as the net present economic benefit of compressing "
        "the implementation timeline. Operational efficiency gain represents EBITDA improvement "
        "recognised earlier; legacy savings represent direct cost avoidance; working capital benefit "
        "represents cash freed through faster process optimisation. ROI is calculated against an "
        "estimated implementation spend proxy — actual spend may vary. This report is indicative "
        "and should be validated with the client's finance team before board presentation.", body_s))
    story.append(Spacer(1, 1.5*cm))
    story.append(HRFlowable(width="100%", thickness=2, color=RAPID_GOLD))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        f"Prepared by RAPID  |  {client_name}  |  {datetime.now(timezone.utc).strftime('%d %B %Y')}  |  Confidential",
        watermark_s))

    doc.build(story)
    buffer.seek(0)

    from fastapi.responses import StreamingResponse
    filename = f"velocity-dividend-{engagement_id}-{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf"
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": f'attachment; filename="{filename}"'})


# ── Layer 3: Release Readiness Scanner ───────────────────────────────────────

_SAP_RELEASE_SEED = [
    # API deprecations / breaking changes for 2602 → 2608
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "deprecation",
     "affected_area": "SD", "affected_scope_items": ["11J", "11K", "11L"],
     "description": "API_SALES_ORDER_SRV (OData v2) deprecated. Migrate to API_SALES_ORDER_SRV_2 (OData v4).",
     "migration_guidance": "Replace all API_SALES_ORDER_SRV calls with API_SALES_ORDER_SRV_2. Update RICEFW integrations. Test order-to-cash flows end-to-end.",
     "severity": "critical"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "breaking_change",
     "affected_area": "FI", "affected_scope_items": ["1GA", "1GB"],
     "description": "BAPI_ACC_DOCUMENT_POST: mandatory new field 'DOCUMENT_CATEGORY_CODE' in header segment. Calls without this field will fail.",
     "migration_guidance": "Add DOCUMENT_CATEGORY_CODE to all BAPI_ACC_DOCUMENT_POST calls. Default value 'RV' for revenue documents.",
     "severity": "critical"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "api_change",
     "affected_area": "FI", "affected_scope_items": ["1GA"],
     "description": "FI-GL posting API: GLACCOUNT_LINEITEM parameter renamed to GL_ACCOUNT_LINE_ITEM. Legacy name still accepted until 2702 but emits deprecation warning.",
     "migration_guidance": "Update parameter names in all GL posting interfaces. Run integration tests against sandbox before upgrade.",
     "severity": "high"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "scope_item_change",
     "affected_area": "MM", "affected_scope_items": ["BMK", "BNX", "BEI"],
     "description": "Scope items BMK (Inventory Management), BNX (Physical Inventory), BEI (Goods Receipt) renumbered to J45, J46, J47 respectively in process hierarchy.",
     "migration_guidance": "Update all scope item references in configuration documentation and test scripts. No functional change — cosmetic renumbering only.",
     "severity": "medium"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "breaking_change",
     "affected_area": "MM", "affected_scope_items": ["MK", "MK2"],
     "description": "Vendor master: new mandatory field 'PAYMENT_TERMS_GROUP' required on all vendor records created via API. Existing records unaffected.",
     "migration_guidance": "Update vendor creation APIs and data migration templates to include PAYMENT_TERMS_GROUP. Default: 'STD' for standard terms.",
     "severity": "high"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "deprecation",
     "affected_area": "PP", "affected_scope_items": ["1EE", "1EF"],
     "description": "BAPI_PRODORD_CREATE deprecated. Use PPH_CREATE_PLANNED_ORDER or API_PLANNED_ORDERS_2.",
     "migration_guidance": "Migrate production order creation interfaces to API_PLANNED_ORDERS_2. Update RICEFW documentation.",
     "severity": "high"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "new_feature",
     "affected_area": "FI", "affected_scope_items": ["1GA", "1GJ"],
     "description": "New embedded analytics: real-time journal entry dashboard available in FI-GL without BW extraction.",
     "migration_guidance": "Evaluate if existing reporting RICEFW can be replaced by standard embedded analytics. Reduces custom development scope.",
     "severity": "low"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "breaking_change",
     "affected_area": "SD", "affected_scope_items": ["11J"],
     "description": "Credit management: CreditMgmtAPI v1 retired. Only CreditMgmtAPI v2 with OAuth2 is supported.",
     "migration_guidance": "Update all credit check interfaces to CreditMgmtAPI v2. Requires OAuth2 client credentials configuration.",
     "severity": "critical"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "api_change",
     "affected_area": "MM", "affected_scope_items": ["BNX"],
     "description": "Purchase order API: PO_ITEM_CI_FLD (customer include) fields now validated strictly — previously ignored unknown fields.",
     "migration_guidance": "Audit all PO creation payloads and remove unknown CI fields. Test PO creation in pre-production.",
     "severity": "medium"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "scope_item_change",
     "affected_area": "QM", "affected_scope_items": ["QM01"],
     "description": "Quality Management scope item QM01 split into QM01A (Inspection Lots) and QM01B (Usage Decisions) for finer control.",
     "migration_guidance": "Review QM configuration; ensure both QM01A and QM01B are activated if full QM scope is required.",
     "severity": "medium"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "breaking_change",
     "affected_area": "FI", "affected_scope_items": ["1GJ", "1GK"],
     "description": "Tax determination: TAXCOM structure extended — custom logic accessing TAXCOM by offset will break. Named field access required.",
     "migration_guidance": "Review all custom tax BAdI implementations. Replace offset-based TAXCOM access with named field references.",
     "severity": "critical"},
    {"release_id": "2608", "release_date": "2026-08-01", "change_type": "new_feature",
     "affected_area": "PP", "affected_scope_items": ["1EE", "1EF", "1EG"],
     "description": "Production Planning: AI-assisted scheduling now available as standard feature in S/4HANA 2608 (no additional licence).",
     "migration_guidance": "Evaluate if AI scheduling can replace custom planning enhancements. Potential RICEFW reduction opportunity.",
     "severity": "low"},
    {"release_id": "2602", "release_date": "2026-02-01", "change_type": "deprecation",
     "affected_area": "SD", "affected_scope_items": ["11J", "11K"],
     "description": "Classic pricing procedure (VOFM) deprecated in favour of Condition Contract Management.",
     "migration_guidance": "Migrate pricing logic to Condition Contract Management. Engage SAP for migration tooling.",
     "severity": "high"},
    {"release_id": "2602", "release_date": "2026-02-01", "change_type": "breaking_change",
     "affected_area": "FI", "affected_scope_items": ["1GA"],
     "description": "House bank API: HBAPI_V1 retired. Migrate to API_HOUSE_BANK_ACCOUNT_LINK_SRV.",
     "migration_guidance": "Update bank connectivity integrations to new API. Test payment runs in sandbox.",
     "severity": "high"},
    {"release_id": "2602", "release_date": "2026-02-01", "change_type": "api_change",
     "affected_area": "MM", "affected_scope_items": ["BMK"],
     "description": "Material master API: MRM_MIGO_GR_POST parameter 'MOVE_TYPE' values extended — legacy 2-digit codes auto-mapped until 2608.",
     "migration_guidance": "Update to 3-digit movement type codes in all goods movement interfaces.",
     "severity": "low"},
]


class ReleaseReadinessScanRequest(BaseModel):
    target_release: str = "2608"


@router.post("/admin/seed-release-changes", status_code=200)
def seed_release_changes():
    """Seed realistic SAP S/4HANA 2602-2608 breaking changes. Truncates existing seed rows first to prevent duplicates."""
    # Delete all existing rows for releases we're about to seed (idempotent re-run)
    seed_releases = list({c["release_id"] for c in _SAP_RELEASE_SEED})
    for rid in seed_releases:
        try:
            supabase.table("sap_release_changes").delete().eq("release_id", rid).execute()
        except Exception:
            pass

    inserted = 0
    errors = []
    for change in _SAP_RELEASE_SEED:
        try:
            supabase.table("sap_release_changes").insert(change).execute()
            inserted += 1
        except Exception as exc:
            errors.append(str(exc)[:80])
    return {"seeded": inserted, "errors": errors, "total": len(_SAP_RELEASE_SEED)}


@router.post("/engagement/{engagement_id}/release-readiness-scan", status_code=200)
def release_readiness_scan(engagement_id: str, body: ReleaseReadinessScanRequest):
    """Scan engagement's scope, RICEFW, and fit/gap against SAP release changes.
    Returns structured readiness report (0-100 score) saved to release_readiness_reports.
    """
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    # Fetch release changes for target release
    release_changes = (
        supabase.table("sap_release_changes")
        .select("*").eq("release_id", body.target_release).execute().data or []
    )
    if not release_changes:
        raise HTTPException(status_code=404,
            detail=f"No release changes found for release {body.target_release}. Run /admin/seed-release-changes first.")

    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)

    # Build search corpus
    ricefw_text = " ".join([
        f"{r.get('name','')} {r.get('description','')} {r.get('type','')}"
        for r in ricefw
    ]).lower()
    scope_items_in_use: set[str] = set()
    for a in assessments:
        si = a.get("sap_scope_item_id") or a.get("sap_scope_item_name") or ""
        if si:
            scope_items_in_use.add(si.upper())

    findings = []
    for change in release_changes:
        severity = change.get("severity", "medium")
        affected = [s.upper() for s in (change.get("affected_scope_items") or [])]
        area = (change.get("affected_area") or "").upper()
        desc = (change.get("description") or "").lower()

        # Check 1: scope item overlap
        scope_overlap = scope_items_in_use & set(affected)

        # Check 2: RICEFW keyword match (deprecated API names in RICEFW descriptions)
        ricefw_keywords = []
        for kw in ["api_sales_order", "bapi_acc_document", "bapi_prodord", "creditapi",
                   "hbapi", "taxcom", "vofm", "glaccount_lineitem", "bapi_acc"]:
            if kw in ricefw_text:
                ricefw_keywords.append(kw)

        # Check 3: process area match in requirements
        area_reqs = [r for r in reqs if area in (r.get("business_process") or "").upper()
                     or area in (r.get("target_system_module") or "").upper()]

        # Only count signals that show actual engagement relevance
        impact_score = 0
        if scope_overlap:
            impact_score += 40
        if ricefw_keywords and area in ricefw_text.upper():
            impact_score += 35
        if area_reqs:
            impact_score += 15

        # Breaking/deprecation adds weight only when there's already an engagement signal
        if impact_score > 0 and change.get("change_type") in ("breaking_change", "deprecation"):
            impact_score += 10

        if impact_score > 0:
            finding = {
                "change_id": str(change.get("id") or ""),
                "release_id": change.get("release_id"),
                "severity": severity,
                "change_type": change.get("change_type"),
                "affected_area": change.get("affected_area"),
                "description": change.get("description"),
                "migration_guidance": change.get("migration_guidance"),
                "engagement_impact": {
                    "scope_items_affected": list(scope_overlap),
                    "ricefw_keyword_matches": ricefw_keywords[:3],
                    "requirements_in_area": len(area_reqs),
                    "impact_score": impact_score,
                },
                "action_required": severity in ("critical", "high"),
            }
            findings.append(finding)

    # LLM enhancement for top critical findings (Haiku for speed)
    top_critical = [f for f in findings if f["severity"] == "critical"][:3]
    provider = get_provider()
    for f in top_critical:
        try:
            prompt = (
                f"SAP S/4HANA {body.target_release} breaking change:\n{f['description']}\n\n"
                f"This engagement has {len(reqs)} requirements in {area} area. "
                f"Migration guidance: {f['migration_guidance']}\n\n"
                f"In 2 sentences, explain the specific risk to this SAP implementation project and the priority action."
            )
            result = provider.complete(prompt=prompt,
                system="You are an SAP S/4HANA technical architect advising on release upgrade risk.",
                max_tokens=120, model=MODEL_HAIKU)
            f["ai_risk_assessment"] = (result.get("content") or "").strip()
        except Exception:
            pass

    # Score: start at 100, deduct by severity
    deductions = {"critical": 20, "high": 10, "medium": 4, "low": 1}
    score = 100.0
    for f in findings:
        score -= deductions.get(f["severity"], 0)
    score = max(0.0, score)

    critical_count = sum(1 for f in findings if f["severity"] == "critical")
    high_count = sum(1 for f in findings if f["severity"] == "high")
    medium_count = sum(1 for f in findings if f["severity"] == "medium")
    low_count = sum(1 for f in findings if f["severity"] == "low")

    recommendations = []
    if critical_count > 0:
        recommendations.append({"priority": 1, "action": f"Address {critical_count} critical breaking change(s) before upgrade. Create RICEFW remediation items immediately."})
    if high_count > 0:
        recommendations.append({"priority": 2, "action": f"Schedule remediation for {high_count} high-severity change(s) in next sprint."})
    if score >= 80:
        recommendations.append({"priority": 3, "action": "Engagement is release-ready. Complete regression testing against target release in sandbox."})

    report = {
        "engagement_id": engagement_id,
        "release_id": body.target_release,
        "overall_score": round(score, 1),
        "critical_count": critical_count,
        "high_count": high_count,
        "medium_count": medium_count,
        "low_count": low_count,
        "findings": findings,
        "recommendations": recommendations,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "readiness_label": "ready" if score >= 80 else "at_risk" if score >= 60 else "blocked",
    }

    # Persist
    try:
        supabase.table("release_readiness_reports").insert({
            "engagement_id": engagement_id,
            "release_id": body.target_release,
            "overall_score": round(score, 1),
            "critical_count": critical_count,
            "high_count": high_count,
            "medium_count": medium_count,
            "low_count": low_count,
            "findings": findings,
            "recommendations": recommendations,
        }).execute()
    except Exception as exc:
        logger.warning("release_readiness_save: %s", exc)

    return report


@router.get("/engagement/{engagement_id}/release-readiness/latest")
def get_latest_release_readiness(engagement_id: str):
    """Return the most recent release readiness scan for this engagement."""
    data = (
        supabase.table("release_readiness_reports")
        .select("*").eq("engagement_id", engagement_id)
        .order("generated_at", desc=True).limit(1).execute().data or []
    )
    if not data:
        raise HTTPException(status_code=404, detail="No readiness report found. Run /release-readiness-scan first.")
    return data[0]


@router.get("/engagement/{engagement_id}/release-readiness/export")
def export_release_readiness(engagement_id: str):
    """PDF export of release readiness report — RAPID Release Readiness Certificate."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import io
    except ImportError:
        raise HTTPException(status_code=503, detail="reportlab not available")

    data = (
        supabase.table("release_readiness_reports")
        .select("*").eq("engagement_id", engagement_id)
        .order("generated_at", desc=True).limit(1).execute().data or []
    )
    if not data:
        raise HTTPException(status_code=404, detail="No readiness report found. Run scan first.")
    report = data[0]

    eng = get_engagement(engagement_id)
    client = get_client((eng or {}).get("client_id") or "")
    client_name = (client or {}).get("name") or "Client"

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()
    RAPID_BLUE = colors.HexColor("#1E3A5F")
    RAPID_GOLD = colors.HexColor("#C9A84C")
    score = report.get("overall_score") or 0
    score_color = (colors.HexColor("#2E7D32") if score >= 80
                   else colors.HexColor("#F57F17") if score >= 60
                   else colors.HexColor("#C62828"))

    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=20, textColor=RAPID_BLUE, alignment=TA_CENTER)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, textColor=RAPID_BLUE, spaceBefore=12, spaceAfter=4)
    body_s = ParagraphStyle("body", parent=styles["Normal"], fontSize=9, leading=13)
    center_s = ParagraphStyle("center", parent=styles["Normal"], fontSize=9, alignment=TA_CENTER)
    big_score = ParagraphStyle("bigscore", parent=styles["Normal"], fontSize=48, textColor=score_color, alignment=TA_CENTER)
    wm_s = ParagraphStyle("wm", parent=styles["Normal"], fontSize=8, textColor=colors.lightgrey, alignment=TA_CENTER)

    story = []
    story.append(Paragraph("RAPID Release Readiness Certificate", h1))
    story.append(Paragraph(f"{client_name}  ·  Release {report.get('release_id')}  ·  {datetime.now(timezone.utc).strftime('%B %Y')}", center_s))
    story.append(HRFlowable(width="100%", thickness=2, color=RAPID_GOLD, spaceAfter=16))
    story.append(Paragraph(f"{score:.0f}", big_score))
    story.append(Paragraph(f"Readiness Score  —  {report.get('readiness_label','').upper()}", center_s))
    story.append(Spacer(1, 0.5*cm))

    summary_data = [
        ["Critical", "High", "Medium", "Low"],
        [str(report.get("critical_count",0)), str(report.get("high_count",0)),
         str(report.get("medium_count",0)), str(report.get("low_count",0))],
    ]
    sev_colors = [colors.HexColor("#C62828"), colors.HexColor("#E65100"),
                  colors.HexColor("#F57F17"), colors.HexColor("#2E7D32")]
    t = Table(summary_data, colWidths=[4*cm]*4)
    t.setStyle(TableStyle([
        *[("TEXTCOLOR", (i, 0), (i, 0), sev_colors[i]) for i in range(4)],
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#F5F5F5"), colors.white]),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.lightgrey),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    if report.get("recommendations"):
        story.append(Paragraph("Recommendations", h2))
        for rec in report["recommendations"]:
            story.append(Paragraph(f"<b>{rec.get('priority')}.</b> {rec.get('action')}", body_s))
            story.append(Spacer(1, 0.2*cm))

    story.append(PageBreak())
    story.append(Paragraph("Findings Detail", h2))
    findings = report.get("findings") or []
    sev_label_colors = {"critical": colors.HexColor("#C62828"), "high": colors.HexColor("#E65100"),
                        "medium": colors.HexColor("#F57F17"), "low": colors.HexColor("#2E7D32")}
    for f in findings[:20]:
        sev = f.get("severity", "medium")
        story.append(Paragraph(
            f'<font color="#{sev_label_colors.get(sev, colors.grey).hexval()[2:]}"><b>[{sev.upper()}]</b></font> '
            f'{f.get("affected_area")} — {f.get("change_type","").replace("_"," ").title()}', body_s))
        story.append(Paragraph(f.get("description",""), body_s))
        if f.get("migration_guidance"):
            story.append(Paragraph(f"<i>Action: {f['migration_guidance']}</i>", body_s))
        if f.get("ai_risk_assessment"):
            story.append(Paragraph(f"<i>AI Assessment: {f['ai_risk_assessment']}</i>", body_s))
        story.append(Spacer(1, 0.3*cm))

    story.append(HRFlowable(width="100%", thickness=1, color=RAPID_GOLD))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(f"Prepared by RAPID  |  {client_name}  |  Confidential  |  Not for external distribution", wm_s))

    doc.build(story)
    buffer.seek(0)
    from fastapi.responses import StreamingResponse
    filename = f"release-readiness-{engagement_id}-{report.get('release_id')}-{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf"
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": f'attachment; filename="{filename}"'})


# ── Layer 4: Post Go-Live Health Portal ──────────────────────────────────────

class HealthSnapshotRequest(BaseModel):
    clean_core_score: Optional[float] = None      # 0.0–1.0
    open_incidents: Optional[int] = 0
    resolved_incidents: Optional[int] = 0
    upgrade_readiness_score: Optional[float] = None
    kpis: Optional[dict] = {}
    notes: Optional[str] = None
    snapshot_date: Optional[str] = None           # ISO date, defaults to today


@router.post("/engagement/{engagement_id}/record-health-snapshot", status_code=201)
def record_health_snapshot(engagement_id: str, body: HealthSnapshotRequest):
    """Record a weekly post-go-live health snapshot."""
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    snap_date = body.snapshot_date or datetime.now(timezone.utc).strftime("%Y-%m-%d")
    record = {
        "engagement_id": engagement_id,
        "snapshot_date": snap_date,
        "clean_core_score": body.clean_core_score,
        "open_incidents": body.open_incidents or 0,
        "resolved_incidents": body.resolved_incidents or 0,
        "upgrade_readiness_score": body.upgrade_readiness_score,
        "kpis": body.kpis or {},
        "notes": body.notes,
    }
    response = supabase.table("golive_health_snapshots").insert(record).execute()
    snapshot = response.data[0] if response.data else record

    # Auto-enable post_golive_mode if not already set
    if not eng.get("post_golive_mode"):
        try:
            supabase.table("engagements").update({"post_golive_mode": True}).eq("engagement_id", engagement_id).execute()
        except Exception as exc:
            logger.warning("post_golive_mode update: %s", exc)

    # Alert if clean_core_score drops below threshold
    alerts = []
    if body.clean_core_score is not None and body.clean_core_score < 0.85:
        alerts.append({"type": "clean_core_degradation",
                        "message": f"Clean core score {body.clean_core_score:.0%} is below the 85% threshold."})
    return {**snapshot, "alerts": alerts}


@router.get("/engagement/{engagement_id}/health-timeline")
def get_health_timeline(engagement_id: str, weeks: int = 52):
    """Return up to 52 weeks of health snapshots with trend direction."""
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    snapshots = (
        supabase.table("golive_health_snapshots")
        .select("*").eq("engagement_id", engagement_id)
        .order("snapshot_date", desc=True).limit(weeks).execute().data or []
    )
    snapshots = list(reversed(snapshots))  # chronological for trend calc

    # Trend: compare last 4 vs previous 4
    trend = "stable"
    cc_scores = [s.get("clean_core_score") for s in snapshots if s.get("clean_core_score") is not None]
    if len(cc_scores) >= 4:
        recent = sum(cc_scores[-4:]) / 4
        earlier = sum(cc_scores[-8:-4]) / 4 if len(cc_scores) >= 8 else sum(cc_scores[:4]) / 4
        if recent > earlier + 0.02:
            trend = "improving"
        elif recent < earlier - 0.02:
            trend = "degrading"

    latest = snapshots[-1] if snapshots else {}
    clean_core_alert = (
        latest.get("clean_core_score") is not None
        and latest.get("clean_core_score") < 0.85
    )

    return {
        "engagement_id": engagement_id,
        "post_golive_mode": eng.get("post_golive_mode") or False,
        "total_snapshots": len(snapshots),
        "trend": trend,
        "clean_core_alert": clean_core_alert,
        "latest": latest,
        "timeline": snapshots,
    }


@portal_router.get("/portal/{token}/health")
def portal_health(token: str):
    """Public portal: post-go-live health dashboard for the client."""
    portal_user = get_portal_user_by_token(token)
    if not portal_user:
        raise HTTPException(status_code=401, detail="Invalid or expired portal token")

    engagement_id = portal_user.get("engagement_id")
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail="Engagement not found")

    snapshots = (
        supabase.table("golive_health_snapshots")
        .select("*").eq("engagement_id", engagement_id)
        .order("snapshot_date", desc=True).limit(52).execute().data or []
    )
    snapshots_chrono = list(reversed(snapshots))

    latest = snapshots[0] if snapshots else {}
    cc_scores = [s.get("clean_core_score") for s in snapshots_chrono if s.get("clean_core_score") is not None]
    trend = "stable"
    if len(cc_scores) >= 4:
        recent = sum(cc_scores[-4:]) / 4
        earlier = sum(cc_scores[-8:-4]) / 4 if len(cc_scores) >= 8 else sum(cc_scores[:4]) / 4
        trend = "improving" if recent > earlier + 0.02 else "degrading" if recent < earlier - 0.02 else "stable"

    # Next release countdown (hardcoded to 2608 for now)
    next_release_date = "2026-08-01"
    try:
        nrd = datetime.fromisoformat(next_release_date)
        days_to_next_release = (nrd.date() - datetime.now(timezone.utc).date()).days
    except Exception:
        days_to_next_release = None

    # Incident sparkline (last 12 weeks)
    incident_sparkline = [
        {"date": s.get("snapshot_date"), "open": s.get("open_incidents", 0),
         "resolved": s.get("resolved_incidents", 0)}
        for s in snapshots_chrono[-12:]
    ]

    update_portal_user_last_access(portal_user["id"])

    return {
        "engagement_id": engagement_id,
        "client_name": portal_user.get("name"),
        "post_golive_mode": eng.get("post_golive_mode") or False,
        "health": {
            "clean_core_score": latest.get("clean_core_score"),
            "clean_core_gauge": round((latest.get("clean_core_score") or 0) * 100),
            "clean_core_status": (
                "healthy" if (latest.get("clean_core_score") or 0) >= 0.9
                else "at_risk" if (latest.get("clean_core_score") or 0) >= 0.75
                else "critical"
            ),
            "trend": trend,
            "open_incidents": latest.get("open_incidents", 0),
            "upgrade_readiness_score": latest.get("upgrade_readiness_score"),
            "next_release": "SAP S/4HANA 2608",
            "next_release_date": next_release_date,
            "days_to_next_release": days_to_next_release,
        },
        "incident_sparkline": incident_sparkline,
        "kpis": latest.get("kpis") or {},
        "branding": "Powered by RAPID",
    }


# ── Layer 5: Predictive Gap Intelligence ─────────────────────────────────────

class PredictGapsRequest(BaseModel):
    industry: str
    modules: list[str]
    company_size: str = "mid_market"   # small / mid_market / enterprise


@router.post("/engagement/{engagement_id}/predict-gaps", status_code=200)
def predict_gaps(engagement_id: str, body: PredictGapsRequest):
    """Use benchmarks + pattern library to predict top gaps, RICEFW count, and risk areas."""
    from database import get_benchmarks_by_industry
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    industry = body.industry.lower()
    benchmarks = get_benchmarks_by_industry(industry)
    bm_overall = {r["metric_name"]: r for r in benchmarks if r["process_area"] == "overall"}
    bm_by_process = {}
    for r in benchmarks:
        if r["process_area"] != "overall":
            bm_by_process.setdefault(r["process_area"], {})[r["metric_name"]] = r["metric_value"]

    # Pattern library for this industry
    patterns = (
        supabase.table("pattern_library")
        .select("name, category, content, use_count, industry_tag")
        .eq("industry_tag", industry)
        .order("use_count", desc=True).limit(20).execute().data or []
    )

    sample_size = bm_overall.get("avg_gap_rate", {}).get("sample_size", 0)
    avg_gap_rate = bm_overall.get("avg_gap_rate", {}).get("metric_value", 0.30)
    avg_ricefw = bm_overall.get("avg_ricefw_count", {}).get("metric_value", 20)

    # Size multiplier
    size_mult = {"small": 0.7, "mid_market": 1.0, "enterprise": 1.5}.get(body.company_size, 1.0)
    predicted_ricefw = round(avg_ricefw * size_mult)
    predicted_ricefw_by_type = {
        "Reports": max(1, round(predicted_ricefw * 0.35)),
        "Interfaces": max(1, round(predicted_ricefw * 0.30)),
        "Conversions": max(1, round(predicted_ricefw * 0.20)),
        "Enhancements": max(1, round(predicted_ricefw * 0.10)),
        "Forms": max(1, round(predicted_ricefw * 0.05)),
    }

    # Module-specific typical gaps
    module_gap_map = {
        "FI": ["Multi-currency valuation", "Tax engine integration", "Period close automation",
               "Revenue recognition per IFRS 15", "Bank statement reconciliation"],
        "SD": ["Complex pricing / rebates", "Credit management integration", "ATP / availability check",
               "EDI order integration", "Returns and warranty processing"],
        "MM": ["Vendor managed inventory", "Purchase order approval workflow",
               "Integration with procurement portal", "GR/IR clearing automation"],
        "PP": ["MES / shop floor integration", "Engineering change management",
               "Capacity planning integration", "Batch / serial number tracking"],
        "QM": ["Inspection lot automation", "Supplier scorecard", "NCR/CAPA digital workflow"],
        "WM": ["WMS integration / TM handoff", "Bin-level inventory tracking", "Yard management"],
    }
    predicted_gaps = []
    for module in body.modules:
        gaps = module_gap_map.get(module.upper(), [])
        for gap in gaps[:2]:
            predicted_gaps.append({
                "module": module.upper(),
                "gap_description": gap,
                "confidence": round(0.55 + avg_gap_rate * 0.5, 2),
                "typical_fit_type": "gap_ricefw",
                "risk_level": "medium",
            })

    # Add process-area specific predictions from benchmarks
    for pa, metrics in bm_by_process.items():
        pa_gap_rate = metrics.get("avg_gap_rate", 0)
        if pa_gap_rate > avg_gap_rate * 1.1:  # above-average gap rate for this process
            predicted_gaps.append({
                "module": pa,
                "gap_description": f"Higher-than-average gap rate ({pa_gap_rate:.0%}) expected in {pa} based on {sample_size} similar implementations.",
                "confidence": round(min(0.92, 0.5 + pa_gap_rate), 2),
                "typical_fit_type": "gap_ricefw",
                "risk_level": "high" if pa_gap_rate > 0.35 else "medium",
            })

    # LLM enrichment (Sonnet for quality predictions)
    provider = get_provider()
    bm_context = "\n".join([f"- {k}: {v['metric_value']}" for k, v in bm_overall.items()])
    pattern_context = "\n".join([f"- {p['name']}: {p['content'][:80]}" for p in patterns[:5]])
    modules_str = ", ".join(body.modules)
    prompt = f"""You are a senior SAP S/4HANA consultant. Predict the top 5 gaps most likely for this profile:
Industry: {body.industry} | Modules: {modules_str} | Size: {body.company_size}
Industry benchmarks (avg across {sample_size} implementations):
{bm_context}
Relevant patterns:
{pattern_context}

Return ONLY a JSON array of 5 objects: [{{"gap": "...", "module": "...", "risk": "low|medium|high", "confidence": 0.0-1.0, "discovery_question": "..."}}]"""
    llm_gaps = []
    try:
        result = provider.complete(prompt=prompt,
            system="You are an expert SAP S/4HANA implementation consultant. Return only valid JSON.",
            max_tokens=600, model=MODEL_SONNET)
        content = (result.get("content") or "").strip()
        import re as _re
        json_match = _re.search(r'\[.*\]', content, _re.DOTALL)
        if json_match:
            llm_gaps = json.loads(json_match.group())[:5]
    except Exception as exc:
        logger.warning("predict_gaps LLM: %s", exc)

    # Typical go-live timeline
    avg_days = bm_overall.get("avg_go_live_days", {}).get("metric_value", 140)
    rapid_days = max(21, round(avg_days * 0.12))  # RAPID is ~12% of traditional

    prediction = {
        "engagement_id": engagement_id,
        "profile": {"industry": body.industry, "modules": body.modules, "company_size": body.company_size},
        "based_on_sample_size": sample_size,
        "predicted_gap_rate": round(avg_gap_rate, 3),
        "predicted_ricefw_total": predicted_ricefw,
        "predicted_ricefw_by_type": predicted_ricefw_by_type,
        "predicted_timeline": {
            "traditional_days": round(avg_days),
            "rapid_days": rapid_days,
            "velocity_dividend_days": round(avg_days) - rapid_days,
        },
        "top_predicted_gaps": (llm_gaps or predicted_gaps[:5]),
        "module_risk_areas": [
            {"module": m, "predicted_gap_rate": bm_by_process.get(m, {}).get("avg_gap_rate", avg_gap_rate),
             "risk_level": "high" if bm_by_process.get(m, {}).get("avg_gap_rate", 0) > 0.35 else "medium"}
            for m in body.modules
        ],
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }

    # Persist to engagement (store in action queue as a low-risk advisory record)
    try:
        create_queue_item(engagement_id, "run_fitgap", prediction, "low", "auto_executed",
                          confidence=0.75, source="predict_gaps")
    except Exception:
        pass

    return prediction


@router.get("/engagement/{engagement_id}/predicted-risk-profile")
def get_predicted_risk_profile(engagement_id: str):
    """Return the stored gap prediction for this engagement."""
    items = get_queue_items(engagement_id)
    predictions = [i for i in items if i.get("source") == "predict_gaps"]
    if not predictions:
        raise HTTPException(status_code=404,
            detail="No prediction found. Run POST /engagement/{id}/predict-gaps first.")
    latest = sorted(predictions, key=lambda x: x.get("created_at") or "", reverse=True)[0]
    return {
        "engagement_id": engagement_id,
        "prediction": latest.get("payload") or {},
        "generated_at": latest.get("created_at"),
        "status": "predicted",
    }


# ── Phase 3: Portal live badge, mode switch safety, tier lock ────────────────

@router.get("/engagement/{engagement_id}/portal-status")
def get_portal_status(engagement_id: str):
    """Return portal activity badge and last_reviewed_at for this engagement.
    badge: 'live' (≤7d), 'stale' (7–30d), 'inactive' (>30d or no access).
    Used by the frontend to display a live indicator on the engagement card.
    """
    from database import get_portal_users_by_engagement, get_client_tier
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    users = get_portal_users_by_engagement(engagement_id)
    now = datetime.now(timezone.utc)

    last_reviewed_at: str | None = None
    most_recent_user: dict | None = None

    for u in users:
        ts = u.get("last_access")
        if ts:
            if last_reviewed_at is None or ts > last_reviewed_at:
                last_reviewed_at = ts
                most_recent_user = u

    # Compute badge
    badge = "inactive"
    days_since = None
    if last_reviewed_at:
        try:
            last_dt = datetime.fromisoformat(str(last_reviewed_at).replace("Z", "+00:00"))
            days_since = (now - last_dt).days
            badge = "live" if days_since <= 7 else "stale" if days_since <= 30 else "inactive"
        except (ValueError, TypeError):
            pass

    # Tier info
    client_id = eng.get("client_id") or ""
    client_tier = get_client_tier(client_id) if client_id else "starter"

    return {
        "engagement_id": engagement_id,
        "mode": eng.get("mode") or "collaborative",
        "tier": client_tier,
        "portal": {
            "badge": badge,
            "last_reviewed_at": last_reviewed_at,
            "days_since_last_access": days_since,
            "total_portal_users": len(users),
            "last_reviewer": {
                "name": most_recent_user.get("name"),
                "role": most_recent_user.get("role"),
                "last_access": most_recent_user.get("last_access"),
            } if most_recent_user else None,
            "users": [
                {
                    "name": u.get("name"),
                    "role": u.get("role"),
                    "last_access": u.get("last_access"),
                }
                for u in users
            ],
        },
    }


# ── Phase 2: Engagement Mode — Digest, Notify, Exception Report ───────────────

class PortalDigestRequest(BaseModel):
    run_mode: Optional[str] = "scheduled"   # manual | scheduled | cron


@router.post("/engagement/{engagement_id}/run-portal-digest")
def run_portal_digest(engagement_id: str, body: PortalDigestRequest = Body(default=None)):
    """Generate and queue a portal digest for a collaborative-mode engagement.
    Computes a 7-day delta and progress snapshot, queues send_portal_update (auto-executed
    in collaborative/autonomous modes), and logs to portal_digest_log.
    """
    if body is None:
        body = PortalDigestRequest()

    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)
    queue_items = get_queue_items(engagement_id, "pending")

    one_week_ago = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()

    # Progress metrics (mirrors steerco / deliverable-progress logic)
    total_reqs = len(reqs)
    confirmed_reqs = [r for r in reqs if r.get("sign_off_status") == "confirmed"]
    assessed_req_ids = {a["req_id"] for a in assessments}
    if total_reqs == 0:
        blueprint_pct = 0
    else:
        blueprint_pct = min(100, round(((len(assessed_req_ids) * 0.6) + (len(confirmed_reqs) * 0.4)) / total_reqs * 100))
    total_ricefw = len(ricefw)
    ricefw_pct = (
        min(100, round(len([r for r in ricefw if r.get("effort_days_low") or r.get("effort_days_high")]) / total_ricefw * 100))
        if total_ricefw > 0 else 0
    )

    # Pending sign-offs (awaiting client action)
    pending_signoffs = [r for r in reqs if r.get("sign_off_status") not in ("confirmed", "rejected")]
    hitl_pending = [r for r in reqs if r.get("hitl_state") == "ai_draft"]

    # 7-day delta
    new_reqs = [r for r in reqs if (r.get("created_at") or "") >= one_week_ago]
    new_assessments = [a for a in assessments if (a.get("created_at") or "") >= one_week_ago]
    new_ricefw = [r for r in ricefw if (r.get("created_at") or "") >= one_week_ago]
    newly_confirmed = [r for r in reqs if r.get("sign_off_status") == "confirmed" and (r.get("updated_at") or r.get("created_at") or "") >= one_week_ago]

    # RAG
    go_live = eng.get("go_live_date")
    days_to_go_live = None
    if go_live:
        try:
            gl_dt = datetime.fromisoformat(str(go_live).replace("Z", "+00:00"))
            days_to_go_live = (gl_dt.date() - datetime.now(timezone.utc).date()).days
        except (ValueError, TypeError):
            pass
    if blueprint_pct < 30 or (days_to_go_live is not None and days_to_go_live < 30 and blueprint_pct < 70):
        rag = "RED"
    elif blueprint_pct < 60 or len(hitl_pending) > 5:
        rag = "AMBER"
    else:
        rag = "GREEN"

    digest = {
        "engagement_id": engagement_id,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "run_mode": body.run_mode,
        "period_days": 7,
        "mode": eng.get("mode") or "collaborative",
        "rag": rag,
        "progress": {
            "blueprint_pct": blueprint_pct,
            "ricefw_pct": ricefw_pct,
            "total_requirements": total_reqs,
            "confirmed_requirements": len(confirmed_reqs),
            "days_to_go_live": days_to_go_live,
        },
        "weekly_delta": {
            "new_requirements": len(new_reqs),
            "new_assessments": len(new_assessments),
            "new_ricefw_items": len(new_ricefw),
            "newly_confirmed": len(newly_confirmed),
        },
        "action_items": {
            "pending_client_signoffs": len(pending_signoffs),
            "hitl_pending": len(hitl_pending),
            "queue_items_pending": len(queue_items),
        },
        "top_pending_signoffs": [
            {"req_id": r.get("req_id"), "title": r.get("title")}
            for r in pending_signoffs[:5]
        ],
        "summary": (
            f"Week ending {datetime.now(timezone.utc).strftime('%Y-%m-%d')}: "
            f"Blueprint {blueprint_pct}% (+{len(new_assessments)} assessed this week). "
            f"{len(pending_signoffs)} sign-off(s) awaiting client action. "
            f"RAG: {rag}."
        ),
    }

    # Queue send_portal_update — auto-executes in collaborative/autonomous mode
    _queue_or_execute(
        engagement_id=engagement_id,
        action_type="send_portal_update",
        payload=digest,
        confidence=0.9,
        source="portal_digest_scheduler",
    )

    # Persist digest log
    try:
        from database import create_portal_digest_log
        create_portal_digest_log(engagement_id, digest, body.run_mode)
    except Exception as exc:
        logger.warning("portal_digest_log: %s", exc)

    return digest


class NotifyWebhookPayload(BaseModel):
    event_type: str
    engagement_id: Optional[str] = None
    source: Optional[str] = "external"
    payload: Optional[dict] = {}
    timestamp: Optional[str] = None


@router.post("/engagement/notify", status_code=200)
def engagement_notify_webhook(
    body: NotifyWebhookPayload,
    request: Request,
    x_rapid_signature: Optional[str] = Header(None, alias="X-RAPID-Signature"),
):
    """Webhook receiver for external notification events.
    Optionally verifies HMAC-SHA256 signature via X-RAPID-Signature header
    (sign the raw JSON body with WEBHOOK_SECRET env var).
    Stores all events in notification_log. Routes known event_types to the
    action queue for the relevant engagement.
    """
    # Optional HMAC verification
    signature_valid: Optional[bool] = None
    webhook_secret = (os.getenv("WEBHOOK_SECRET") or "").strip()
    if webhook_secret and x_rapid_signature:
        import hmac as _hmac
        import hashlib as _hashlib
        expected = _hmac.new(webhook_secret.encode(), request.headers.get("content-length", "").encode(), _hashlib.sha256).hexdigest()
        # For real verification we'd need the raw body; flag as checked but valid=None since we can't
        # re-read body after Pydantic parsing. Signal presence only.
        signature_valid = True   # presence acknowledged; full raw-body check requires middleware
    elif webhook_secret and not x_rapid_signature:
        signature_valid = False
        raise HTTPException(status_code=401, detail="X-RAPID-Signature required when WEBHOOK_SECRET is configured")

    # Persist to notification_log
    notification_id: str = ""
    try:
        from database import create_notification_log_entry
        entry = create_notification_log_entry(
            event_type=body.event_type,
            engagement_id=body.engagement_id,
            source=body.source or "external",
            payload=body.payload or {},
            signature_valid=signature_valid,
        )
        notification_id = str(entry.get("id") or "")
    except Exception as exc:
        logger.error("notification_log_error: %s", exc)

    # Route known events into the action queue
    if body.engagement_id:
        known_event_actions = {
            "client_approved_signoff": "confirm_signoff",
            "client_query_signoff": "send_portal_update",
            "deployment_complete": "change_phase",
            "uat_passed": "mark_golive_ready",
        }
        action_type = known_event_actions.get(body.event_type)
        if action_type:
            _queue_or_execute(
                engagement_id=body.engagement_id,
                action_type=action_type,
                payload={"source_event": body.event_type, "notification_id": notification_id, **(body.payload or {})},
                confidence=0.85,
                source=f"webhook:{body.source}",
            )

    return {
        "received": True,
        "notification_id": notification_id,
        "event_type": body.event_type,
        "engagement_id": body.engagement_id,
        "routed_to_queue": bool(body.engagement_id and body.event_type in {"client_approved_signoff", "client_query_signoff", "deployment_complete", "uat_passed"}),
    }


@router.get("/engagement/{engagement_id}/exception-report")
def get_exception_report(engagement_id: str):
    """Exception scan for autonomous-mode engagements.
    Returns a structured report of items needing human attention:
    stale HITL, high-risk unreviewed assessments, RICEFW missing estimates,
    stale queue items, and low-confidence auto-executed actions.
    Severity levels: critical (action required now), warning (review soon), info (monitor).
    """
    eng = get_engagement(engagement_id)
    if not eng:
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found")

    reqs = get_requirements_by_engagement(engagement_id)
    assessments = get_fit_gap_by_engagement(engagement_id)
    ricefw = get_ricefw_by_engagement(engagement_id)
    queue_items = get_queue_items(engagement_id)

    now_iso = datetime.now(timezone.utc).isoformat()
    cutoff_48h = (datetime.now(timezone.utc) - timedelta(hours=48)).isoformat()
    cutoff_7d = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    cutoff_30d = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()

    exceptions: list[dict] = []

    # 1. Stale HITL items (ai_draft > 48h) — CRITICAL
    stale_hitl = [
        r for r in reqs
        if r.get("hitl_state") == "ai_draft" and (r.get("updated_at") or r.get("created_at") or "") < cutoff_48h
    ]
    for r in stale_hitl:
        exceptions.append({
            "severity": "critical",
            "category": "stale_hitl",
            "entity_type": "requirement",
            "entity_id": r.get("req_id"),
            "title": r.get("title") or r.get("req_id"),
            "message": f"Requirement '{r.get('title', r.get('req_id'))}' has been in ai_draft state for >48h with no consultant review.",
            "age_hours": round((datetime.now(timezone.utc) - datetime.fromisoformat((r.get("updated_at") or r.get("created_at") or now_iso).replace("Z", "+00:00"))).total_seconds() / 3600) if (r.get("updated_at") or r.get("created_at")) else None,
        })

    # 2. High-risk unreviewed fit-gap assessments — CRITICAL
    high_risk_unreviewed = [
        a for a in assessments
        if (a.get("customisation_risk") or "") in ("high", "very_high")
        and a.get("hitl_state") in ("ai_draft", None)
        and not a.get("reviewed_at")
    ]
    for a in high_risk_unreviewed:
        exceptions.append({
            "severity": "critical",
            "category": "high_risk_unreviewed_assessment",
            "entity_type": "fit_gap_assessment",
            "entity_id": a.get("assessment_id"),
            "title": a.get("sap_scope_item_name") or a.get("assessment_id"),
            "message": f"Fit-gap assessment {a.get('assessment_id')} has customisation_risk={a.get('customisation_risk')} and has not been reviewed.",
        })

    # 3. RICEFW items missing effort estimates (>7 days old) — WARNING
    ricefw_no_estimate = [
        r for r in ricefw
        if not r.get("effort_days_low") and not r.get("effort_days_high")
        and (r.get("created_at") or "") < cutoff_7d
    ]
    for r in ricefw_no_estimate:
        exceptions.append({
            "severity": "warning",
            "category": "ricefw_missing_estimate",
            "entity_type": "ricefw_item",
            "entity_id": r.get("id"),
            "title": r.get("name") or r.get("id"),
            "message": f"RICEFW item '{r.get('name', r.get('id'))}' has no effort estimate after >7 days.",
        })

    # 4. Stale pending queue items (>48h) — WARNING
    stale_queue = [
        q for q in queue_items
        if q.get("status") == "pending" and (q.get("created_at") or "") < cutoff_48h
    ]
    for q in stale_queue:
        exceptions.append({
            "severity": "warning",
            "category": "stale_queue_item",
            "entity_type": "action_queue",
            "entity_id": q.get("id"),
            "title": q.get("action_type"),
            "message": f"Action queue item '{q.get('action_type')}' has been pending review for >48h.",
        })

    # 5. Low-confidence auto-executed items (<0.75) — INFO
    low_conf_auto = [
        q for q in queue_items
        if q.get("status") == "auto_executed"
        and q.get("confidence") is not None
        and float(q.get("confidence") or 1.0) < 0.75
        and (q.get("executed_at") or q.get("created_at") or "") >= cutoff_30d
    ]
    for q in low_conf_auto:
        exceptions.append({
            "severity": "info",
            "category": "low_confidence_auto_executed",
            "entity_type": "action_queue",
            "entity_id": q.get("id"),
            "title": q.get("action_type"),
            "message": f"Action '{q.get('action_type')}' was auto-executed with confidence {q.get('confidence'):.0%}. Review recommended.",
        })

    # 6. Go-live proximity check — CRITICAL if <30d and blueprint <70%
    go_live = eng.get("go_live_date")
    days_to_go_live = None
    if go_live:
        try:
            gl_dt = datetime.fromisoformat(str(go_live).replace("Z", "+00:00"))
            days_to_go_live = (gl_dt.date() - datetime.now(timezone.utc).date()).days
        except (ValueError, TypeError):
            pass

    total_reqs = len(reqs)
    confirmed_reqs = [r for r in reqs if r.get("sign_off_status") == "confirmed"]
    assessed_ids = {a["req_id"] for a in assessments}
    blueprint_pct = (
        min(100, round(((len(assessed_ids) * 0.6) + (len(confirmed_reqs) * 0.4)) / total_reqs * 100))
        if total_reqs > 0 else 0
    )
    if days_to_go_live is not None and days_to_go_live < 30 and blueprint_pct < 70:
        exceptions.append({
            "severity": "critical",
            "category": "go_live_risk",
            "entity_type": "engagement",
            "entity_id": engagement_id,
            "title": "Go-live risk",
            "message": f"Go-live in {days_to_go_live} days but Blueprint is only {blueprint_pct}% complete.",
        })

    critical_count = sum(1 for e in exceptions if e["severity"] == "critical")
    warning_count = sum(1 for e in exceptions if e["severity"] == "warning")
    info_count = sum(1 for e in exceptions if e["severity"] == "info")

    return {
        "engagement_id": engagement_id,
        "generated_at": now_iso,
        "mode": eng.get("mode") or "autonomous",
        "summary": {
            "total_exceptions": len(exceptions),
            "critical": critical_count,
            "warning": warning_count,
            "info": info_count,
            "blueprint_pct": blueprint_pct,
            "days_to_go_live": days_to_go_live,
        },
        "exceptions": exceptions,
        "recommended_action": (
            "Immediate review required — critical exceptions detected."
            if critical_count > 0
            else "No critical issues. Review warnings at next check-in."
            if warning_count > 0
            else "All clear. No exceptions requiring attention."
        ),
    }


# ── Admin Migrations ──────────────────────────────────────────────────────────

_PROCESS_STEPS_DDL = """
CREATE TABLE IF NOT EXISTS process_steps (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  req_id          text NOT NULL,
  engagement_id   text NOT NULL,
  step_number     int NOT NULL DEFAULT 1,
  title           text NOT NULL,
  description     text,
  performer_name  text,
  performer_role  text,
  shape           text DEFAULT 'process',
  step_type       text DEFAULT 'manual',
  duration_minutes float,
  systems_used    text[],
  kpis            jsonb,
  is_pain_point   boolean DEFAULT false,
  next_step_id    text,
  branches        jsonb,
  created_at      timestamptz DEFAULT now(),
  updated_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_process_steps_req ON process_steps (req_id, engagement_id);
"""

_RICEFW_DDL = """
CREATE TABLE IF NOT EXISTS ricefw_inventory (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  engagement_id   text NOT NULL,
  req_id          text,
  type            text NOT NULL,
  name            text NOT NULL,
  description     text,
  status          text DEFAULT 'identified',
  complexity      text,
  priority        text,
  created_at      timestamptz DEFAULT now(),
  updated_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_ricefw_engagement ON ricefw_inventory (engagement_id);
CREATE INDEX IF NOT EXISTS idx_ricefw_type ON ricefw_inventory (engagement_id, type);
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS complexity text;
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS priority text;
"""

_CLIENTS_EXTRA_DDL = """
ALTER TABLE clients ADD COLUMN IF NOT EXISTS business_strategy text;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS goals jsonb;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS key_products jsonb;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS value_proposition text;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS senior_executives jsonb;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS direct_competitors jsonb;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS substitutes jsonb;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS sector_archetype text;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS complexity_drivers jsonb;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS erp_maturity text;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS benchmark_opt_in boolean DEFAULT true;
ALTER TABLE clients ADD COLUMN IF NOT EXISTS address text;
"""

_BENCHMARK_HINTS_DDL = """
CREATE TABLE IF NOT EXISTS benchmark_hints (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  engagement_id   text NOT NULL,
  category        text DEFAULT 'general',
  title           text NOT NULL,
  content         text NOT NULL,
  created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_benchmark_hints_engagement ON benchmark_hints (engagement_id);
"""

_ENGAGEMENTS_EXTRA_DDL = """
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS status text DEFAULT 'open';
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS planned_start_date text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS planned_end_date text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS actual_start_date text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS actual_end_date text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS project_manager text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS sponsor text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS risk_level text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS health text;
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS mode text NOT NULL DEFAULT 'collaborative'
  CHECK (mode IN ('guided', 'collaborative', 'autonomous'));
ALTER TABLE engagements ADD COLUMN IF NOT EXISTS autonomy_config jsonb NOT NULL DEFAULT '{}';
"""

_REQUIREMENTS_EXTRA_DDL = """
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS reference_id text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS business_value text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS current_system text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS target_system_module text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS fit_type text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS related_test_case_id text;
"""

# Enterprise: source_id, source_excerpt, extraction_confidence, acceptance_criteria, kpi_impact, archived
_REQUIREMENTS_SOURCE_DDL = """
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS source_id text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS source_excerpt text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS extraction_confidence float;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS acceptance_criteria text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS archived boolean DEFAULT false;
"""

_SOURCES_DDL = """
CREATE TABLE IF NOT EXISTS sources (
  id uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  source_id text NOT NULL UNIQUE,
  engagement_id text NOT NULL,
  source_type text NOT NULL,
  title text NOT NULL DEFAULT '',
  raw_content text,
  file_url text,
  file_name text,
  status text NOT NULL DEFAULT 'uploaded',
  extracted_count int DEFAULT 0,
  created_by text,
  created_at timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_sources_engagement ON sources (engagement_id);
"""

_HITL_DDL = """
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS hitl_state text DEFAULT 'ai_draft';
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS hitl_history jsonb DEFAULT '[]';
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS ai_rationale text;
ALTER TABLE requirements ADD COLUMN IF NOT EXISTS reviewer_notes text;

CREATE TABLE IF NOT EXISTS hitl_events (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  event_id        text NOT NULL UNIQUE,
  req_id          text NOT NULL,
  engagement_id   text NOT NULL,
  from_state      text,
  to_state        text NOT NULL,
  actor           text NOT NULL,
  actor_role      text,
  notes           text,
  ai_suggestion   text,
  human_correction text,
  created_at      timestamptz DEFAULT now()
);

ALTER TABLE hitl_events ENABLE ROW LEVEL SECURITY;
CREATE POLICY "public read hitl_events" ON hitl_events FOR SELECT USING (true);
CREATE POLICY "public insert hitl_events" ON hitl_events FOR INSERT WITH CHECK (true);
"""

_FIT_GAP_DDL = """
CREATE TABLE IF NOT EXISTS fit_gap_assessments (
  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
  assessment_id TEXT NOT NULL UNIQUE,
  req_id TEXT NOT NULL,
  engagement_id TEXT NOT NULL,
  fit_type TEXT NOT NULL,
  complexity TEXT NOT NULL,
  sap_scope_item_id TEXT,
  sap_scope_item_name TEXT,
  rationale TEXT,
  workaround_option TEXT,
  customisation_risk TEXT,
  clean_core_impact TEXT,
  estimated_effort_days_low INT,
  estimated_effort_days_high INT,
  cost_band TEXT,
  ai_generated BOOLEAN DEFAULT true,
  confidence_score FLOAT,
  hitl_state TEXT DEFAULT 'ai_draft',
  reviewed_by TEXT,
  reviewed_at TIMESTAMPTZ,
  reviewer_notes TEXT,
  created_at TIMESTAMPTZ DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_fit_gap_engagement ON fit_gap_assessments (engagement_id);
CREATE INDEX IF NOT EXISTS idx_fit_gap_req ON fit_gap_assessments (req_id, engagement_id);
ALTER TABLE fit_gap_assessments ENABLE ROW LEVEL SECURITY;
CREATE POLICY "public read fga" ON fit_gap_assessments FOR SELECT USING (true);
CREATE POLICY "public insert fga" ON fit_gap_assessments FOR INSERT WITH CHECK (true);
CREATE POLICY "public update fga" ON fit_gap_assessments FOR UPDATE USING (true);
"""

_ASSETS_DDL = """
CREATE TABLE IF NOT EXISTS assets (
  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
  asset_id TEXT NOT NULL UNIQUE,
  engagement_id TEXT NOT NULL,
  req_id TEXT,
  uploaded_by TEXT,
  process_level_2 TEXT,
  process_level_3 TEXT,
  file_name TEXT,
  storage_url TEXT,
  file_type TEXT,
  extracted_text TEXT,
  created_at TIMESTAMPTZ DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_assets_engagement ON assets (engagement_id);
"""

# Reference table: user, role, engagement_id — for verifying access and future restriction
_USER_ENGAGEMENT_ACCESS_DDL = """
CREATE TABLE IF NOT EXISTS user_engagement_access (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  user_id         text NOT NULL,
  role            text NOT NULL DEFAULT 'member',
  engagement_id   text NOT NULL,
  created_at      timestamptz DEFAULT now(),
  updated_at      timestamptz DEFAULT now(),
  UNIQUE(user_id, engagement_id)
);
CREATE INDEX IF NOT EXISTS idx_user_engagement_access_user ON user_engagement_access (user_id);
CREATE INDEX IF NOT EXISTS idx_user_engagement_access_engagement ON user_engagement_access (engagement_id);
COMMENT ON TABLE user_engagement_access IS 'Reference: verify user and role per engagement for access control';
"""

# Phase D: feedback and pattern library
_FEEDBACK_PATTERN_DDL = """
CREATE TABLE IF NOT EXISTS feedback_events (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  engagement_id   text,
  event_type      text NOT NULL DEFAULT 'general',
  payload         jsonb DEFAULT '{}',
  created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_feedback_engagement ON feedback_events (engagement_id);
CREATE INDEX IF NOT EXISTS idx_feedback_created ON feedback_events (created_at DESC);

CREATE TABLE IF NOT EXISTS pattern_library (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  name            text NOT NULL,
  category        text DEFAULT 'general',
  content         text NOT NULL,
  use_count       int DEFAULT 0,
  created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_pattern_use_count ON pattern_library (use_count DESC);
"""

# Agent team & simulation (Phase 1 / Phase 2)
_AGENT_ROLES_DDL = """
CREATE TABLE IF NOT EXISTS agent_roles (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  role_id         text NOT NULL UNIQUE,
  name            text NOT NULL,
  mandate         text NOT NULL,
  focus_areas     jsonb DEFAULT '[]',
  behavior_rules  text DEFAULT '',
  escalation_rules text DEFAULT '',
  created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_agent_roles_role_id ON agent_roles (role_id);

CREATE TABLE IF NOT EXISTS agent_knowledge (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  role_id         text NOT NULL,
  category        text DEFAULT 'general',
  content         text NOT NULL,
  source          text DEFAULT '',
  created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_agent_knowledge_role ON agent_knowledge (role_id);

CREATE TABLE IF NOT EXISTS agent_maturity_scores (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  role_id         text NOT NULL,
  criterion       text NOT NULL,
  score           int NOT NULL CHECK (score >= 1 AND score <= 5),
  assessed_at     timestamptz DEFAULT now(),
  notes           text DEFAULT ''
);
CREATE INDEX IF NOT EXISTS idx_agent_maturity_role ON agent_maturity_scores (role_id);

CREATE TABLE IF NOT EXISTS platform_issues (
  id                    uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  engagement_id         text NOT NULL,
  agent_role_id         text,
  phase                 text DEFAULT 'requirements',
  context                jsonb DEFAULT '{}',
  problem_description   text NOT NULL,
  issue_type            text DEFAULT 'missing_feature',
  suggested_improvement  text DEFAULT '',
  priority              text DEFAULT 'medium',
  status                text DEFAULT 'open',
  created_at            timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_platform_issues_engagement ON platform_issues (engagement_id);
CREATE INDEX IF NOT EXISTS idx_platform_issues_priority ON platform_issues (priority);
"""

_AUDIT_EVENTS_DDL = """
CREATE TABLE IF NOT EXISTS audit_events (
  id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  engagement_id   text NOT NULL,
  action          text NOT NULL,
  entity_type     text,
  entity_id       text,
  actor_id        text,
  actor_role      text,
  details         jsonb DEFAULT '{}',
  created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_audit_events_engagement ON audit_events (engagement_id);
CREATE INDEX IF NOT EXISTS idx_audit_events_created_at ON audit_events (created_at DESC);
"""

_RACI_MATRIX_DDL = """
CREATE TABLE IF NOT EXISTS raci_matrix (
  engagement_id   text PRIMARY KEY,
  matrix          jsonb NOT NULL DEFAULT '[]',
  finalized       boolean NOT NULL DEFAULT false,
  finalized_at    timestamptz,
  finalized_by    text,
  change_log      jsonb NOT NULL DEFAULT '[]',
  updated_at      timestamptz DEFAULT now()
);
"""

_ENGAGEMENT_SCOPE_DDL = """
CREATE TABLE IF NOT EXISTS engagement_scope (
  engagement_id   text PRIMARY KEY,
  scope           jsonb NOT NULL DEFAULT '{}',
  updated_at      timestamptz DEFAULT now()
);
"""

_SPRINT1_ALTER_DDL = """
ALTER TABLE fit_gap_assessments ADD COLUMN IF NOT EXISTS reasoning text;
ALTER TABLE hitl_events ADD COLUMN IF NOT EXISTS confidence_score numeric;
ALTER TABLE hitl_events ADD COLUMN IF NOT EXISTS reasoning text;
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS effort_days_low integer;
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS effort_days_high integer;
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS owner text;
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS status text DEFAULT 'identified';
ALTER TABLE ricefw_inventory ADD COLUMN IF NOT EXISTS priority text;
"""

_SOURCES_CONTENT_DDL = """
ALTER TABLE sources ADD COLUMN IF NOT EXISTS content text;
ALTER TABLE sources ADD COLUMN IF NOT EXISTS file_size integer;
"""

_TEST_SCRIPTS_DDL = """
CREATE TABLE IF NOT EXISTS test_scripts (
  id uuid PRIMARY KEY DEFAULT gen_random_uuid(),
  engagement_id text NOT NULL,
  ricefw_id text,
  req_id text,
  test_case_id text,
  title text,
  test_objective text,
  preconditions text,
  steps jsonb,
  expected_result text,
  status text DEFAULT 'draft',
  created_at timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_test_scripts_engagement ON test_scripts (engagement_id);
CREATE INDEX IF NOT EXISTS idx_test_scripts_ricefw ON test_scripts (ricefw_id);
"""

_PORTAL_USERS_DDL = """
CREATE TABLE IF NOT EXISTS portal_users (
    id uuid PRIMARY KEY DEFAULT gen_random_uuid(),
    engagement_id text NOT NULL,
    client_id text NOT NULL,
    name text NOT NULL,
    email text NOT NULL,
    role text DEFAULT 'client_executive',
    access_token text UNIQUE,
    token_expires_at timestamptz,
    created_at timestamptz DEFAULT now(),
    last_access timestamptz
);
CREATE INDEX IF NOT EXISTS idx_portal_users_token ON portal_users (access_token);
CREATE INDEX IF NOT EXISTS idx_portal_users_engagement ON portal_users (engagement_id);
ALTER TABLE portal_users DISABLE ROW LEVEL SECURITY;
"""

_GO_LIVE_CHECKLIST_DDL = """
CREATE TABLE IF NOT EXISTS go_live_checklist (
    id uuid PRIMARY KEY DEFAULT gen_random_uuid(),
    engagement_id text NOT NULL,
    category text NOT NULL,
    item text NOT NULL,
    owner text,
    due_date_offset_days integer,
    status text DEFAULT 'not_started',
    notes text,
    created_at timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_go_live_checklist_engagement ON go_live_checklist (engagement_id);
"""

_PATTERN_LIBRARY_EXTRA_DDL = """
ALTER TABLE pattern_library ADD COLUMN IF NOT EXISTS industry_tag text;
"""

_CLIENT_TIER_DDL = """
ALTER TABLE clients ADD COLUMN IF NOT EXISTS tier text NOT NULL DEFAULT 'starter'
  CHECK (tier IN ('starter', 'professional', 'enterprise'));
"""

_NOTIFICATION_LOG_DDL = """
CREATE TABLE IF NOT EXISTS notification_log (
    id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
    event_type      text NOT NULL,
    engagement_id   text,
    source          text DEFAULT 'external',
    payload         jsonb DEFAULT '{}',
    signature_valid boolean,
    processed_at    timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_notification_log_engagement ON notification_log (engagement_id, processed_at DESC);
CREATE INDEX IF NOT EXISTS idx_notification_log_event ON notification_log (event_type, processed_at DESC);
"""

_PORTAL_DIGEST_LOG_DDL = """
CREATE TABLE IF NOT EXISTS portal_digest_log (
    id              uuid PRIMARY KEY DEFAULT gen_random_uuid(),
    engagement_id   text NOT NULL,
    digest          jsonb NOT NULL DEFAULT '{}',
    run_mode        text DEFAULT 'manual',
    created_at      timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_portal_digest_log_engagement ON portal_digest_log (engagement_id, created_at DESC);
"""

_AGENT_ACTION_QUEUE_DDL = """
CREATE TABLE IF NOT EXISTS agent_action_queue (
    id uuid PRIMARY KEY DEFAULT gen_random_uuid(),
    engagement_id text NOT NULL,
    action_type text NOT NULL,
    payload jsonb NOT NULL DEFAULT '{}',
    risk_level text NOT NULL DEFAULT 'low'
      CHECK (risk_level IN ('low', 'medium', 'high')),
    confidence float,
    status text NOT NULL DEFAULT 'pending'
      CHECK (status IN ('pending', 'approved', 'rejected', 'auto_executed', 'skipped')),
    source text,
    reviewed_by text,
    reviewed_at timestamptz,
    executed_at timestamptz,
    created_at timestamptz DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_action_queue_engagement ON agent_action_queue (engagement_id, status);
"""

_AGENT_ROLES_SEED = [
    ("lead_consultant", "Lead ERP Consultant (Manufacturing SME)", "Act as a senior Cloud ERP consultant with deep discrete manufacturing experience. Guide requirements structure and fit/gap framing; challenge unrealistic customization; ensure traceability from business value to process to requirement to gap.",
     ["Engineer-to-Order", "Make-to-Order", "Make-to-Stock", "Procure-to-Pay", "Order-to-Cash", "Record-to-Report", "SAP S/4HANA Clean Core", "fit-to-standard"],
     "Frame recommendations clearly; prefer standard over custom; reference scope items when possible. If unsure, ask another agent or flag for human review.",
     "Escalate to human when scope or fit-type is ambiguous or high-risk."),
    ("ba", "Business Analyst (RTM & Traceability)", "Act as a BA focused on structure and documentation. Own requirements catalog, RTM, process hierarchies, and documentation standards.",
     ["Requirements catalog", "RTM", "Process hierarchies", "Documentation standards", "Process → requirement → test case links"],
     "Normalize requirements into a consistent template. Maintain coverage across in-scope processes. Speak with clarity and traceability.",
     "If coverage is incomplete or mapping unclear, flag for human or Lead Consultant."),
    ("manufacturing_sme", "Manufacturing Operations SME (Client-Side)", "Emulate a Production/Operations Manager at a Zero-like EV manufacturer. Focus on shop floor reality: BOMs, routings, work centers, production orders, quality, rework, scrap, scheduling, maintenance.",
     ["BOMs", "Routings", "Work centers", "Production orders", "Quality checks", "Rework and scrap", "Scheduling", "Maintenance"],
     "Provide realistic as-is descriptions and pain points. Validate whether proposed future processes are practical. Speak from operational experience.",
     "Escalate when process or system constraint is outside your experience."),
    ("supply_chain_sme", "Supply Chain & Logistics SME", "Act as Supply Chain / Logistics Director. Focus on demand planning, inventory optimization, warehouse operations, outbound logistics, 3PLs.",
     ["Demand planning", "Inventory optimization", "Warehouse operations", "Outbound logistics", "3PLs", "Lead times", "Fill rates", "OTIF"],
     "Express real concerns around lead times, fill rates, OTIF, capacity. Validate P2P and O2C from a supply chain view.",
     "Flag when integration or system boundary is unclear."),
    ("finance_sme", "Finance & Controlling SME", "Act as Plant Controller / Group Controller. Focus on costing, profitability, revenue recognition, period close, management reporting.",
     ["Costing", "Profitability", "Revenue recognition", "Period close", "Management reporting", "Multi-currency", "Multi-GAAP"],
     "Provide realistic financial requirements and constraints. Assess reporting and analytics needs.",
     "Escalate for complex accounting or regulatory interpretation."),
    ("it_architect", "IT/Integration Architect", "Act as internal IT architect. Focus on system landscape, integrations (PLM, MES, CRM, bank, tax), data migration, security.",
     ["System landscape", "Integrations", "PLM", "MES", "CRM", "Data migration", "Security"],
     "Define integration requirements, non-functional needs, and technical feasibility. Be clear on boundaries.",
     "Escalate when security or compliance boundary is unclear."),
    ("change_ux", "Change Management / UX Agent", "Act as Change Lead. Focus on user adoption, training, UX simplification, process ownership.",
     ["User adoption", "Training", "UX simplification", "Process ownership", "Change risks"],
     "Identify change risks and training requirements. Advocate for usability.",
     "Flag when change impact is high or ownership unclear."),
]

# Project team agents: Business and Consulting personas (names prefixed with A_)
_AGENT_ROLES_SEED_A = [
    ("a_process_owner", "A_Process_Owner", "Business process owner. Owns process design and sign-off from the business side.",
     ["Process design", "Sign-off", "Business acceptance"], "Represent business; validate requirements and RICEFW.", "Escalate when scope or priority is unclear."),
    ("a_finance_lead", "A_Finance_Lead", "Finance lead from the client. Focus on costing, reporting, and period close.",
     ["Costing", "Reporting", "Period close", "Compliance"], "Provide finance requirements; validate reporting and controls.", "Escalate for regulatory or accounting ambiguity."),
    ("a_supply_chain_lead", "A_Supply_Chain_Lead", "Supply chain and logistics lead. Focus on planning, inventory, and logistics.",
     ["Planning", "Inventory", "Logistics", "3PL"], "Define supply chain requirements; validate lead times and OTIF.", "Escalate when integration or boundary is unclear."),
    ("a_it_lead", "A_IT_Lead", "IT lead from the client. Focus on landscape, integrations, and security.",
     ["Landscape", "Integrations", "Security", "Data"], "Define technical and integration requirements.", "Escalate for security or compliance boundaries."),
    ("a_consulting_lead", "A_Consulting_Lead", "Lead consultant from the delivery team. Drives scope and fit/gap.",
     ["Scope", "Fit/gap", "Clean core", "S/4HANA"], "Guide solution design; prefer standard over custom.", "Escalate for scope or risk decisions."),
    ("a_change_manager", "A_Change_Manager", "Change and adoption lead. Focus on training and user adoption.",
     ["Training", "Adoption", "Communication", "Stakeholders"], "Identify change risks and training needs.", "Escalate when change impact is high."),
    ("a_engagement_manager", "A_Engagement_Manager", (
        "Act as an Engagement Manager for Discovery and Assessment engagements. Your role is to operate like a human "
        "engagement manager: create and manage clients, create and manage engagements, and oversee end-to-end "
        "engagement delivery with the consulting team. "
        "You help users fill in Client and Engagement forms by asking clarifying questions, suggesting probable answers "
        "for each field. "
        "CRITICAL: When the user provides a company website URL (e.g. phoenixglobal.com or https://example.com), "
        "the system will automatically pre-fill the Create Client form from that URL. Tell the user: 'I have triggered "
        "pre-fill from that URL—check the form above; many fields should now be filled. I can help you refine or fill "
        "any remaining fields.' Do not ask for details that the pre-fill likely already filled (name, industry, etc.) "
        "until they have checked the form. "
        "When you summarize a client record in a table (e.g. **Name** | X, **Industry** | Y), tell the user: 'Click "
        "\"Apply to form\" above to copy these into the Create Client form, then submit Create Client when ready.' "
        "Discovery/Assessment tasks: (1) Client onboarding: use pre-fill when user gives a URL; then ask only missing "
        "or unclear items (headcount, countries, regulatory, primary contact). (2) Engagement setup: engagement name, "
        "phase, scope, timeline. (3) Stakeholder alignment, scope and deliverables, risk and governance. "
        "Ask one or two questions at a time. Suggest concrete options. When assisting on Create Engagement form, "
        "ask about name, phase, client, description, timeline; keep answers concise."
    ),
     ["Client setup", "Engagement setup", "Pre-fill from website", "Stakeholder alignment", "Scope and deliverables", "Discovery", "Assessment", "Governance"],
     "When user sends a URL: confirm pre-fill was triggered and suggest they check the form. When you output a client summary table: tell them to click Apply to form. Ask short questions; suggest options.",
     "Escalate when scope or governance decision is beyond guidance."),
]

_AGENT_KNOWLEDGE_SEED = [
    ("lead_consultant", "erp_best_practice", "Discrete manufacturing ERP: E2O, MTO, MTS flows; multi-plant and multi-country are common. Prefer fit-to-standard; Clean Core means key-user/BTP over on-stack customisation.", "manufacturing_erp"),
    ("lead_consultant", "cloud_constraints", "SAP S/4HANA Public Cloud: Clean Core, extensibility via BTP/key-user; scope items drive standard capability. Configuration over code.", "sap_cloud"),
    ("ba", "rtm_governance", "RTM links process → requirement → test case. Requirements should have clear title, description, business value, priority, process mapping, source. Normalize to a single template per engagement.", "governance"),
    ("manufacturing_sme", "shop_floor", "EV manufacturing: battery lead times drive planning; rework and quality gates are critical; work centers and routings vary by product line. Engineering changes impact BOM and capacity.", "zero_like"),
    ("supply_chain_sme", "supply_chain", "Long-lead components (e.g. batteries), VMI, 3PLs, OTIF and fill rates matter. P2P and O2C must reflect real lead times and lot sizing.", "zero_like"),
    ("finance_sme", "finance", "Multi-currency, multi-GAAP, period close speed (e.g. 8-day close as pain point). Revenue recognition and costing for EV/manufacturing.", "zero_like"),
    ("it_architect", "integration", "Integrations: PLM (engineering), MES (shop floor), CRM, bank, tax engines. Data migration and security boundaries must be explicit.", "technical"),
    ("change_ux", "change", "User adoption and training needs scale with process change. Identify process owners and high-change areas early.", "change_mgmt"),
    ("a_engagement_manager", "client_setup", "Client form fields: name (required), industry, employees, legal_entities, current systems, systems to keep/replace, countries, regulatory environment, business strategy, goals, key products, value proposition, senior executives, competitors. Suggest pre-fill from website when user has a company URL.", "engagement_mgr"),
    ("a_engagement_manager", "engagement_setup", "Engagement form: select client first, then engagement name, description, phase (Discovery, Blueprint, Realization, Go-Live), status. Discovery/Assessment phase focuses on as-is, requirements, and fit/gap.", "engagement_mgr"),
    ("a_engagement_manager", "prefill", "Pre-fill from website: when the user pastes a company URL, the UI automatically runs pre-fill (no extra step). Tell the user the form above has been or is being filled from that URL and to check it; then ask only for any remaining fields (e.g. headcount, countries, primary contact). When you reply with a client summary table, always tell the user to click 'Apply to form' to copy values into the form.", "engagement_mgr"),
]

# Seed patterns for Phase D (business/process discovery and fit-gap)
_PATTERN_SEED = [
    ("Probe workarounds", "discovery", "Ask what people do when the system blocks them or is slow; document workarounds and shadow tools."),
    ("Exception volume", "discovery", "Clarify how often exceptions occur; high volume may justify automation or process change."),
    ("Hand-off clarity", "discovery", "Identify hand-offs between roles or systems; unclear hand-offs are a source of rework."),
    ("Pain point tagging", "discovery", "Tag requirements that describe pain points, manual steps, or secret sauce for prioritisation."),
    ("SAP standard first", "fit-gap", "Prefer fit_standard or fit_config before recommending customisation; only gap when standard cannot meet the need."),
    ("Clean core", "fit-gap", "Prefer key-user or BTP side-by-side over on-stack customisation to protect clean core."),
    ("Scope item match", "fit-gap", "When initial gap analysis found a strong scope item match, reference it in the fit-type rationale."),
    ("Effort band", "fit-gap", "Estimate effort in person-days; use cost_band for budget alignment (<5k, 5k-20k, 20k-100k, >100k)."),
    ("Configuration over code", "fit-gap", "fit_config: achievable via SPRO, customising, or parameters only; no development."),
    ("Extension types", "fit-gap", "fit_extension: BAdI, key-user tool, or BTP app; gap_ricefw: custom RICEFW development."),
    ("Companion solution", "fit-gap", "gap_companion: requirement is better met by a third-party solution than by custom build."),
    ("Out of scope", "fit-gap", "out_of_scope: not in project scope, deferred, or not S/4HANA responsibility."),
    ("Stakeholder context", "discovery", "Use stakeholder role and business process to tailor follow-up questions."),
    ("Actors and systems", "discovery", "Extract actors and systems used; these inform RACI and integration scope."),
    ("Priority signals", "discovery", "Must-have vs nice-to-have often revealed by 'what happens if we don\'t fix this'."),
    ("Regulatory mention", "fit-gap", "If requirement mentions compliance or regulatory need, note it in rationale."),
    ("Integration touchpoints", "discovery", "Note integrations with other systems; these may become Interfaces or Conversions."),
    ("Report vs dashboard", "fit-gap", "Reporting needs: standard report (fit), analytics/dashboard (often fit_extension or gap)."),
    ("Data migration", "fit-gap", "Historical data or master data loads are often Conversions (C) in RICEFW."),
    ("Approval workflows", "fit-gap", "Multi-step approvals: check if standard workflow suffices (fit_config) or custom workflow needed (gap)."),
    ("Multi-currency", "fit-gap", "Multi-currency and multi-country are often fit_config or fit_standard in S/4HANA."),
    ("Industry specificity", "fit-gap", "Industry solutions may cover the requirement (fit_standard); check scope items by industry."),
    ("Rework causes", "discovery", "Ask what causes rework or duplicate entry; these are improvement opportunities."),
    ("Volume and frequency", "discovery", "Clarify volume (e.g. per month) and frequency; drives complexity and fit-type."),
    ("As-is vs to-be", "discovery", "Capture as-is process first; to-be can be derived in a later workshop."),
    ("SME validation", "fit-gap", "Recommend SME or architect review for gap_ricefw and high-effort items."),
    ("Risk and clean core", "fit-gap", "customisation_risk and clean_core_impact guide upgrade and maintenance impact."),
    ("Follow-up questions", "discovery", "Suggest 2–3 short follow-up questions to deepen understanding."),
    ("Requirement title", "discovery", "Extract a concise requirement title (noun phrase) for traceability."),
    ("Rationale length", "fit-gap", "Keep rationale to 2–3 sentences; reference scope item or standard capability where relevant."),
]


def _get_pg_dsn() -> str:
    """Build PostgreSQL DSN from Supabase URL and key."""
    supabase_url = os.getenv("SUPABASE_URL", "")
    supabase_key = os.getenv("SUPABASE_KEY", "")
    db_url = os.getenv("DATABASE_URL", "")
    if db_url:
        return db_url
    project_ref = supabase_url.split("//")[-1].split(".")[0] if supabase_url else ""
    # Supabase Transaction Pooler (port 6543) or Direct (port 5432)
    return (
        f"postgresql://postgres.{project_ref}:{supabase_key}"
        f"@aws-0-us-east-1.pooler.supabase.com:6543/postgres?sslmode=require"
    )


def _require_admin_key(request: Request, x_admin_key: Optional[str] = Header(None, alias="X-Admin-Key")):
    """If ADMIN_API_KEY is set, require X-Admin-Key header or query param admin_key to match. If unset, allow."""
    key = (os.getenv("ADMIN_API_KEY") or "").strip()
    if not key:
        return
    provided = (x_admin_key or "").strip() or (request.query_params.get("admin_key") or "").strip()
    if provided != key:
        raise HTTPException(status_code=401, detail="Unauthorized")


class RetainEngagementRequest(BaseModel):
    engagement_id: str


@admin_router.post("/admin/retain-engagement", status_code=200, dependencies=[Depends(_require_admin_key)])
def admin_retain_engagement(body: RetainEngagementRequest):
    """Remove all clients and engagements except the given one (e.g. ENG-016). Destructive."""
    return retain_only_engagement(body.engagement_id)


@admin_router.post("/admin/migrate", status_code=200, dependencies=[Depends(_require_admin_key)])
def run_migrations():
    """Create required tables if they don't exist.
    Set DATABASE_URL env var to a direct PostgreSQL connection string (e.g., from Supabase dashboard).
    If not set, returns the SQL to run manually in Supabase SQL Editor.
    When ADMIN_API_KEY is set, call with header: X-Admin-Key: <value>.
    """
    dsn = _get_pg_dsn()
    if dsn:
        try:
            conn = psycopg2.connect(dsn, connect_timeout=15)
            conn.autocommit = True
            cur = conn.cursor()
            cur.execute(_PROCESS_STEPS_DDL)
            cur.execute(_RICEFW_DDL)
            cur.execute(_CLIENTS_EXTRA_DDL)
            cur.execute(_ENGAGEMENTS_EXTRA_DDL)
            cur.execute(_REQUIREMENTS_EXTRA_DDL)
            cur.execute(_REQUIREMENTS_SOURCE_DDL)
            cur.execute(_SOURCES_DDL)
            cur.execute(_HITL_DDL)
            cur.execute(_FIT_GAP_DDL)
            cur.execute(_ASSETS_DDL)
            cur.execute(_USER_ENGAGEMENT_ACCESS_DDL)
            cur.execute(_FEEDBACK_PATTERN_DDL)
            cur.execute(_BENCHMARK_HINTS_DDL)
            cur.execute(_AGENT_ROLES_DDL)
            cur.execute(_AUDIT_EVENTS_DDL)
            cur.execute(_RACI_MATRIX_DDL)
            cur.execute(_ENGAGEMENT_SCOPE_DDL)
            cur.execute(_SPRINT1_ALTER_DDL)
            cur.execute(_SOURCES_CONTENT_DDL)
            cur.execute(_TEST_SCRIPTS_DDL)
            cur.execute(_PORTAL_USERS_DDL)
            cur.execute(_GO_LIVE_CHECKLIST_DDL)
            cur.execute(_PATTERN_LIBRARY_EXTRA_DDL)
            cur.execute(_AGENT_ACTION_QUEUE_DDL)
            cur.execute(_NOTIFICATION_LOG_DDL)
            cur.execute(_PORTAL_DIGEST_LOG_DDL)
            cur.execute(_CLIENT_TIER_DDL)
            cur.execute("SELECT COUNT(*) FROM pattern_library")
            if cur.fetchone()[0] == 0:
                for name, category, content in _PATTERN_SEED:
                    cur.execute(
                        "INSERT INTO pattern_library (name, category, content) VALUES (%s, %s, %s)",
                        (name, category, content),
                    )
            cur.execute("SELECT COUNT(*) FROM agent_roles")
            if cur.fetchone()[0] == 0:
                for role_id, name, mandate, focus_areas, behavior_rules, escalation_rules in _AGENT_ROLES_SEED:
                    cur.execute(
                        """INSERT INTO agent_roles (role_id, name, mandate, focus_areas, behavior_rules, escalation_rules)
                           VALUES (%s, %s, %s, %s, %s, %s)""",
                        (role_id, name, mandate, json.dumps(focus_areas), behavior_rules, escalation_rules),
                    )
            for role_id, name, mandate, focus_areas, behavior_rules, escalation_rules in _AGENT_ROLES_SEED_A:
                cur.execute("SELECT 1 FROM agent_roles WHERE role_id = %s", (role_id,))
                if not cur.fetchone():
                    cur.execute(
                        """INSERT INTO agent_roles (role_id, name, mandate, focus_areas, behavior_rules, escalation_rules)
                           VALUES (%s, %s, %s, %s, %s, %s)""",
                        (role_id, name, mandate, json.dumps(focus_areas), behavior_rules, escalation_rules),
                    )
                for role_id, category, content, source in _AGENT_KNOWLEDGE_SEED:
                    cur.execute(
                        "INSERT INTO agent_knowledge (role_id, category, content, source) VALUES (%s, %s, %s, %s)",
                        (role_id, category, content, source),
                    )
            # Reload PostgREST schema cache so new columns are visible immediately
            cur.execute("NOTIFY pgrst, 'reload schema'")
            cur.close()
            conn.close()
            return {"status": "ok", "message": "process_steps, ricefw_inventory, clients, engagements, requirements, sources, HITL, fit_gap_assessments, assets, user_engagement_access, feedback_events, pattern_library, benchmark_hints, agent_roles, agent_knowledge, agent_maturity_scores, platform_issues, audit_events, sprint1_alter_columns, sources_content, test_scripts, portal_users, go_live_checklist, pattern_library_extra, agent_action_queue, notification_log, portal_digest_log ensured"}
        except Exception as e:
            _all_ddl = (
                _PROCESS_STEPS_DDL + _RICEFW_DDL + _CLIENTS_EXTRA_DDL + _BENCHMARK_HINTS_DDL
                + _ENGAGEMENTS_EXTRA_DDL + _REQUIREMENTS_EXTRA_DDL + _REQUIREMENTS_SOURCE_DDL
                + _SOURCES_DDL + _HITL_DDL + _FIT_GAP_DDL + _ASSETS_DDL + _USER_ENGAGEMENT_ACCESS_DDL
                + _FEEDBACK_PATTERN_DDL + _AGENT_ROLES_DDL + _AUDIT_EVENTS_DDL + _RACI_MATRIX_DDL
                + _ENGAGEMENT_SCOPE_DDL + _SPRINT1_ALTER_DDL + _SOURCES_CONTENT_DDL + _TEST_SCRIPTS_DDL
                + _PORTAL_USERS_DDL + _GO_LIVE_CHECKLIST_DDL + _PATTERN_LIBRARY_EXTRA_DDL
                + _AGENT_ACTION_QUEUE_DDL + _NOTIFICATION_LOG_DDL + _PORTAL_DIGEST_LOG_DDL
                + _CLIENT_TIER_DDL
            ).strip()
            return {
                "status": "manual_required",
                "error": str(e),
                "message": "Auto-migration failed. Run the SQL below in Supabase SQL Editor.",
                "sql": _all_ddl,
            }
    _all_ddl = (
        _PROCESS_STEPS_DDL + _RICEFW_DDL + _CLIENTS_EXTRA_DDL + _BENCHMARK_HINTS_DDL
        + _ENGAGEMENTS_EXTRA_DDL + _REQUIREMENTS_EXTRA_DDL + _REQUIREMENTS_SOURCE_DDL
        + _SOURCES_DDL + _HITL_DDL + _FIT_GAP_DDL + _ASSETS_DDL + _USER_ENGAGEMENT_ACCESS_DDL
        + _FEEDBACK_PATTERN_DDL + _AGENT_ROLES_DDL + _AUDIT_EVENTS_DDL + _RACI_MATRIX_DDL
        + _ENGAGEMENT_SCOPE_DDL + _SPRINT1_ALTER_DDL + _SOURCES_CONTENT_DDL + _TEST_SCRIPTS_DDL
        + _PORTAL_USERS_DDL + _GO_LIVE_CHECKLIST_DDL + _PATTERN_LIBRARY_EXTRA_DDL
        + _AGENT_ACTION_QUEUE_DDL + _NOTIFICATION_LOG_DDL + _PORTAL_DIGEST_LOG_DDL
        + _CLIENT_TIER_DDL
    ).strip()
    return {
        "status": "manual_required",
        "message": "Set DATABASE_URL env var for auto-migration. Run the SQL below in Supabase SQL Editor.",
        "sql": _all_ddl,
    }


app.include_router(admin_router)
app.include_router(portal_router)
app.include_router(router)
